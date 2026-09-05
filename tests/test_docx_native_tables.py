import copy
import tempfile
import unittest
from dataclasses import replace
from pathlib import Path
from zipfile import ZipFile

from docx import Document as open_docx
from docx.oxml.ns import qn

from aiteqno.adapters import PythonDocxObserver, PythonDocxRenderer
from aiteqno.application import (
    build_docx_structure_relationships,
    build_evaluation_reference,
    evaluate_restoration,
    render_docx,
)
from aiteqno.domain import (
    TABLE_TOPOLOGY_COORDINATE_SPACE,
    TABLE_TOPOLOGY_EXTENSION_KEY,
    TABLE_TOPOLOGY_PROVIDER,
    TABLE_TOPOLOGY_PROVIDER_VERSION,
    TABLE_TOPOLOGY_SCHEMA_VERSION,
    BoundingBox,
    DocumentIR,
    PageSize,
    RectangleElement,
    TextElement,
    read_page_table_topology,
)
from aiteqno.ports import RenderPolicy


FIXTURE_PATH = (
    Path(__file__).parent / "fixtures" / "document_ir" / "canonical.document.ir.json"
)
TABLE_ID = "page-001-table-0000"
PARAMETERS_DIGEST = "0" * 64


def _provenance(source_ids: list[str]) -> dict[str, object]:
    return {
        "stage": "derived",
        "provider": TABLE_TOPOLOGY_PROVIDER,
        "provider_version": TABLE_TOPOLOGY_PROVIDER_VERSION,
        "source_element_ids": source_ids,
        "parameters_digest": PARAMETERS_DIGEST,
    }


def _topology_document(
    merge: str | None = None,
    *,
    include_outside_text: bool = True,
) -> DocumentIR:
    canonical = DocumentIR.from_json(FIXTURE_PATH.read_bytes())
    base_text = canonical.pages[0].elements[0]
    base_rectangle = canonical.pages[0].elements[2]
    assert isinstance(base_text, TextElement)
    assert isinstance(base_rectangle, RectangleElement)

    if merge == "horizontal":
        cell_specs = (
            (0, 0, 1, 2),
            (1, 0, 1, 1),
            (1, 1, 1, 1),
        )
    elif merge == "vertical":
        cell_specs = (
            (0, 0, 2, 1),
            (0, 1, 1, 1),
            (1, 1, 1, 1),
        )
    else:
        cell_specs = (
            (0, 0, 1, 1),
            (0, 1, 1, 1),
            (1, 0, 1, 1),
            (1, 1, 1, 1),
        )

    text_elements: list[TextElement] = []
    reading_order = 0
    if include_outside_text:
        text_elements.append(
            replace(
                base_text,
                id="p001-text-0000",
                bbox=BoundingBox(x=20, y=10, width=70, height=10),
                text="Before table",
                reading_order=reading_order,
                style=replace(
                    base_text.style,
                    font_family="Arial",
                    font_size_pt=8,
                    font_weight=400,
                ),
            )
        )
        reading_order += 1

    cell_text_ids: list[str] = []
    cell_bboxes: list[BoundingBox] = []
    for cell_number, (row, column, rowspan, colspan) in enumerate(cell_specs):
        bbox = BoundingBox(
            x=20 + column * 80,
            y=40 + row * 40,
            width=80 * colspan,
            height=40 * rowspan,
        )
        cell_bboxes.append(bbox)
        text_id = f"p001-text-{reading_order:04d}"
        cell_text_ids.append(text_id)
        text_elements.append(
            replace(
                base_text,
                id=text_id,
                bbox=BoundingBox(
                    x=bbox.x + 5,
                    y=bbox.y + 14,
                    width=min(55, bbox.width - 10),
                    height=10,
                ),
                text=f"Cell {cell_number + 1}",
                reading_order=reading_order,
                style=replace(
                    base_text.style,
                    font_family="Arial",
                    font_size_pt=8,
                    font_weight=400,
                ),
            )
        )
        reading_order += 1

    outside_ids: list[str] = []
    if include_outside_text:
        outside_id = f"p001-text-{reading_order:04d}"
        outside_ids.extend(("p001-text-0000", outside_id))
        text_elements.append(
            replace(
                base_text,
                id=outside_id,
                bbox=BoundingBox(x=20, y=135, width=65, height=10),
                text="After table",
                reading_order=reading_order,
                style=replace(
                    base_text.style,
                    font_family="Arial",
                    font_size_pt=8,
                    font_weight=400,
                ),
            )
        )

    outer_id = "p001-rectangle-0000"
    outer = replace(
        base_rectangle,
        id=outer_id,
        bbox=BoundingBox(x=20, y=40, width=160, height=80),
    )
    rectangles: list[RectangleElement] = [outer]
    rectangle_ids: list[str] = []
    cells: list[dict[str, object]] = []
    roles: list[dict[str, object]] = [
        {
            "element_id": outer_id,
            "role": "table_outer_border",
            "table_id": TABLE_ID,
        }
    ]
    for cell_number, ((row, column, rowspan, colspan), bbox, text_id) in enumerate(
        zip(cell_specs, cell_bboxes, cell_text_ids, strict=True),
        start=1,
    ):
        rectangle_id = f"p001-rectangle-{cell_number:04d}"
        rectangle_ids.append(rectangle_id)
        rectangles.append(replace(base_rectangle, id=rectangle_id, bbox=bbox))
        cell_id = f"{TABLE_ID}-cell-r{row:03d}-c{column:03d}"
        cells.append(
            {
                "id": cell_id,
                "row_index": row,
                "column_index": column,
                "rowspan": rowspan,
                "colspan": colspan,
                "bbox": {
                    "x": bbox.x,
                    "y": bbox.y,
                    "width": bbox.width,
                    "height": bbox.height,
                },
                "supporting_element_ids": [rectangle_id],
                "text_element_ids": [text_id],
                "confidence": 1.0,
                "provenance": _provenance([rectangle_id]),
            }
        )
        roles.append(
            {
                "element_id": rectangle_id,
                "role": "cell_rectangle",
                "table_id": TABLE_ID,
                "cell_id": cell_id,
            }
        )

    support_ids = [outer_id, *rectangle_ids]
    extension = {
        "schema_version": TABLE_TOPOLOGY_SCHEMA_VERSION,
        "provider": TABLE_TOPOLOGY_PROVIDER,
        "provider_version": TABLE_TOPOLOGY_PROVIDER_VERSION,
        "coordinate_space": TABLE_TOPOLOGY_COORDINATE_SPACE,
        "boundary_tolerance_pt": 1.0,
        "tables": [
            {
                "id": TABLE_ID,
                "bbox": {"x": 20, "y": 40, "width": 160, "height": 80},
                "logical_grid": {"rows": 2, "columns": 2},
                "rows": [
                    {
                        "id": f"{TABLE_ID}-row-000",
                        "index": 0,
                        "start": 40,
                        "end": 80,
                    },
                    {
                        "id": f"{TABLE_ID}-row-001",
                        "index": 1,
                        "start": 80,
                        "end": 120,
                    },
                ],
                "columns": [
                    {
                        "id": f"{TABLE_ID}-column-000",
                        "index": 0,
                        "start": 20,
                        "end": 100,
                    },
                    {
                        "id": f"{TABLE_ID}-column-001",
                        "index": 1,
                        "start": 100,
                        "end": 180,
                    },
                ],
                "cells": cells,
                "supporting_element_ids": support_ids,
                "confidence": 1.0,
                "provenance": _provenance(support_ids),
            }
        ],
        "primitive_roles": roles,
        "diagnostics": {
            "ambiguous_text_element_ids": [],
            "unassigned_text_element_ids": outside_ids,
            "ambiguous_primitive_element_ids": [],
            "unassigned_primitive_element_ids": [],
            "rejected_table_outer_element_ids": [],
        },
    }
    page = replace(
        canonical.pages[0],
        size=PageSize(width=200, height=160),
        source=None,
        elements=(*text_elements, *rectangles),
        extensions={TABLE_TOPOLOGY_EXTENSION_KEY: extension},
    )
    return replace(canonical, pages=(page,), assets=())


def _fragmented_topology_document() -> DocumentIR:
    payload = _topology_document().to_dict()
    page = payload["pages"][0]
    elements = page["elements"]
    topology = page["extensions"][TABLE_TOPOLOGY_EXTENSION_KEY]

    def replace_text(
        source_id: str,
        fragments: tuple[tuple[str, str, float, float, float, float], ...],
    ) -> list[str]:
        source_index = next(
            index for index, element in enumerate(elements) if element["id"] == source_id
        )
        source = elements[source_index]
        replacements = []
        for element_id, text, x, y, width, height in fragments:
            replacement = copy.deepcopy(source)
            replacement.update(
                {
                    "id": element_id,
                    "text": text,
                    "bbox": {
                        "x": x,
                        "y": y,
                        "width": width,
                        "height": height,
                    },
                }
            )
            replacements.append(replacement)
        elements[source_index : source_index + 1] = replacements
        return [element["id"] for element in replacements]

    outside_ids = replace_text(
        "p001-text-0000",
        (
            ("p001-text-0000", "文", 20, 10, 8, 10),
            ("p001-text-0900", "書", 27, 10, 8, 10),
            ("p001-text-0901", "解析", 36, 10, 20, 10),
            ("p001-text-0902", "2026", 150, 10, 20, 10),
            ("p001-text-0903", "1.", 20, 25, 8, 10),
            ("p001-text-0904", "概要", 35, 25, 18, 10),
        ),
    )
    first_cell_ids = replace_text(
        "p001-text-0001",
        (
            ("p001-text-0001", "依頼", 25, 54, 18, 10),
            ("p001-text-0910", "する", 42, 54, 14, 10),
            ("p001-text-0911", "処理", 59, 54, 18, 10),
        ),
    )
    second_cell_ids = replace_text(
        "p001-text-0002",
        (
            ("p001-text-0002", "Alpha", 105, 54, 22, 10),
            ("p001-text-0920", "Beta", 132, 54, 20, 10),
        ),
    )
    third_cell_ids = replace_text(
        "p001-text-0003",
        (
            ("p001-text-0003", "PNG", 25, 94, 15, 10),
            ("p001-text-0930", "・", 42, 94, 5, 10),
            ("p001-text-0931", "PDF", 48, 94, 15, 10),
            ("p001-text-0932", "(", 65, 94, 3, 10),
            ("p001-text-0933", "26", 69, 94, 10, 10),
            ("p001-text-0934", ")", 80, 94, 3, 10),
        ),
    )
    fourth_cell_ids = replace_text(
        "p001-text-0004",
        (("p001-text-0004", "項目", 105, 94, 18, 10),),
    )

    fragment_font_sizes = {
        "p001-text-0000": 6,
        "p001-text-0900": 12,
        "p001-text-0901": 8,
        "p001-text-0902": 10,
        "p001-text-0001": 7,
        "p001-text-0910": 12,
        "p001-text-0911": 8,
        "p001-text-0004": 6,
    }
    for element in elements:
        if element["id"] in fragment_font_sizes:
            element["style"]["font_size_pt"] = fragment_font_sizes[element["id"]]

    cells = topology["tables"][0]["cells"]
    cells[0]["text_element_ids"] = first_cell_ids
    cells[1]["text_element_ids"] = second_cell_ids
    cells[2]["text_element_ids"] = third_cell_ids
    cells[3]["text_element_ids"] = fourth_cell_ids
    diagnostics = topology["diagnostics"]
    diagnostics["unassigned_text_element_ids"] = [
        *outside_ids,
        diagnostics["unassigned_text_element_ids"][-1],
    ]

    text_elements = [element for element in elements if element["type"] == "text"]
    for reading_order, element in enumerate(text_elements):
        element["reading_order"] = reading_order
    return DocumentIR.from_dict(payload)


class NativeWordTableRendererTest(unittest.TestCase):
    def test_native_table_preserves_grid_geometry_text_and_report_accounting(self):
        document = _topology_document()
        topology = read_page_table_topology(document.pages[0])
        assert topology is not None
        with tempfile.TemporaryDirectory() as temporary_directory:
            output = Path(temporary_directory) / "native.docx"
            result = render_docx(
                document,
                output,
                renderer=PythonDocxRenderer(),
                policy=RenderPolicy.STRICT,
            )
            reopened = open_docx(output)
            observation = PythonDocxObserver().observe(output)

        self.assertEqual(len(reopened.tables), 1)
        table = reopened.tables[0]
        caption = table._tbl.tblPr.find(qn("w:tblCaption"))
        self.assertIsNotNone(caption)
        assert caption is not None
        self.assertEqual(caption.get(qn("w:val")), f"aiteqno-table:{TABLE_ID}")
        self.assertEqual(
            [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid],
            [1600, 1600],
        )
        self.assertEqual(
            [round(row.height.pt) for row in table.rows],
            [34, 34],
        )
        self.assertEqual(result.report.native_table_ids, (TABLE_ID,))
        self.assertEqual(
            len(result.report.native_table_consumed_element_ids),
            9,
        )
        self.assertEqual(
            len(result.report.native_table_consumed_element_ids),
            len(set(result.report.native_table_consumed_element_ids)),
        )
        self.assertEqual(result.report.fallback_element_ids, ())
        self.assertEqual(result.report.omitted_element_ids, ())
        self.assertEqual(observation.errors, ())
        self.assertEqual(
            [
                element.source_element_id
                for element in observation.elements
                if element.source_element_id is not None
            ],
            [
                element.id
                for element in document.pages[0].elements
                if isinstance(element, TextElement)
            ],
        )
        self.assertEqual(
            sum(
                relationship.source == TABLE_ID
                and relationship.kind.value == "containment"
                for relationship in observation.relationships
            ),
            4,
        )

    def test_fragmented_text_uses_script_aware_spacing_and_layout_tabs(self):
        document = _fragmented_topology_document()
        with tempfile.TemporaryDirectory() as temporary_directory:
            output = Path(temporary_directory) / "fragmented.docx"
            result = render_docx(
                document,
                output,
                renderer=PythonDocxRenderer(),
                policy=RenderPolicy.STRICT,
            )
            reopened = open_docx(output)
            observation = PythonDocxObserver().observe(output)

        def source_tags(node) -> set[str]:
            return {
                tag.get(qn("w:val"))
                for tag in node.xpath(".//w:sdtPr/w:tag")
            }

        def text_content(node) -> str:
            return "".join(text.text or "" for text in node.xpath(".//w:t"))

        outside = next(
            paragraph
            for paragraph in reopened._element.body.xpath("./w:p")
            if "aiteqno-source:p001-text-0000" in source_tags(paragraph)
        )
        self.assertEqual(text_content(outside), "文書解析2026")
        self.assertEqual(len(outside.xpath("./w:r/w:tab")), 1)
        outside_font_sizes = {
            size.get(qn("w:val"))
            for size in outside.xpath(".//w:sdtContent/w:r/w:rPr/w:sz")
        }
        self.assertEqual(len(outside_font_sizes), 1)

        numbered = next(
            paragraph
            for paragraph in reopened._element.body.xpath("./w:p")
            if "aiteqno-source:p001-text-0903" in source_tags(paragraph)
        )
        self.assertEqual(text_content(numbered), "1. 概要")

        table = reopened.tables[0]
        self.assertEqual(text_content(table.cell(0, 0)._tc), "依頼する処理")
        self.assertEqual(text_content(table.cell(0, 1)._tc), "Alpha Beta")
        latin_spaces = table.cell(0, 1)._tc.xpath('.//w:r/w:t[text()=" "]')
        self.assertEqual(len(latin_spaces), 1)
        latin_space_fonts = latin_spaces[0].getparent().xpath("./w:rPr/w:rFonts")
        self.assertEqual(len(latin_space_fonts), 1)
        for channel in ("ascii", "hAnsi", "eastAsia", "cs"):
            self.assertEqual(
                latin_space_fonts[0].get(qn(f"w:{channel}")),
                "Arial",
            )
        self.assertEqual(text_content(table.cell(1, 0)._tc), "PNG・PDF(26)")
        self.assertEqual(text_content(table.cell(1, 1)._tc), "項目")
        self.assertFalse(table.cell(0, 0)._tc.xpath('.//w:t[text()=" "]'))
        self.assertFalse(table.cell(1, 0)._tc.xpath('.//w:t[text()=" "]'))
        cell_font_sizes = {
            size.get(qn("w:val"))
            for size in table.cell(0, 0)._tc.xpath(
                ".//w:sdtContent/w:r/w:rPr/w:sz"
            )
        }
        self.assertEqual(len(cell_font_sizes), 1)
        short_label_font_sizes = {
            size.get(qn("w:val"))
            for size in table.cell(1, 1)._tc.xpath(
                ".//w:sdtContent/w:r/w:rPr/w:sz"
            )
        }
        self.assertEqual(short_label_font_sizes, {"21"})

        expected_source_ids = {
            "p001-text-0000",
            "p001-text-0900",
            "p001-text-0901",
            "p001-text-0902",
            "p001-text-0903",
            "p001-text-0904",
            "p001-text-0001",
            "p001-text-0910",
            "p001-text-0911",
            "p001-text-0002",
            "p001-text-0920",
            "p001-text-0003",
            "p001-text-0930",
            "p001-text-0931",
            "p001-text-0932",
            "p001-text-0933",
            "p001-text-0934",
            "p001-text-0004",
        }
        observed_source_ids = [
            element.source_element_id
            for element in observation.elements
            if element.source_element_id in expected_source_ids
        ]
        self.assertEqual(set(observed_source_ids), expected_source_ids)
        self.assertEqual(len(observed_source_ids), len(expected_source_ids))
        self.assertEqual(result.report.fallback_element_ids, ())
        self.assertEqual(result.report.omitted_element_ids, ())

    def test_default_japanese_font_is_applied_to_table_and_outside_runs(self):
        document = _topology_document()
        page = document.pages[0]
        elements = tuple(
            replace(
                element,
                style=replace(element.style, font_family="Noto Sans CJK JP"),
            )
            if isinstance(element, TextElement)
            else element
            for element in page.elements
        )
        document = replace(document, pages=(replace(page, elements=elements),))

        with tempfile.TemporaryDirectory() as temporary_directory:
            output = Path(temporary_directory) / "noto-native.docx"
            result = render_docx(
                document,
                output,
                renderer=PythonDocxRenderer(),
                policy=RenderPolicy.STRICT,
            )
            reopened = open_docx(output)

        outside_contents = reopened._element.body.xpath("./w:p//w:sdtContent")
        table_contents = reopened._element.body.xpath("./w:tbl//w:sdtContent")
        self.assertTrue(outside_contents)
        self.assertTrue(table_contents)
        self.assertEqual(
            len(outside_contents) + len(table_contents),
            sum(isinstance(element, TextElement) for element in elements),
        )
        for content in (*outside_contents, *table_contents):
            run = content.find(qn("w:r"))
            self.assertIsNotNone(run)
            run_properties = run.find(qn("w:rPr"))
            self.assertIsNotNone(run_properties)
            run_fonts = run_properties.find(qn("w:rFonts"))
            self.assertIsNotNone(run_fonts)
            for channel in ("ascii", "hAnsi", "eastAsia", "cs"):
                self.assertEqual(
                    run_fonts.get(qn(f"w:{channel}")),
                    "Noto Sans CJK JP",
                )
        self.assertEqual(result.report.font_substitutions, ())
        self.assertNotIn(
            "font_substituted",
            {warning.code for warning in result.report.warnings},
        )

    def test_fragmented_text_plan_is_deterministic_for_shuffled_planner_input(self):
        document = _fragmented_topology_document()
        text_elements = tuple(
            element
            for element in document.pages[0].elements
            if isinstance(element, TextElement)
        )
        ordered_plan = PythonDocxRenderer._plan_text_lines(text_elements)
        shuffled_plan = PythonDocxRenderer._plan_text_lines(
            tuple(reversed(text_elements))
        )
        self.assertEqual(
            tuple(tuple(element.id for element in line.elements) for line in ordered_plan),
            tuple(
                tuple(element.id for element in line.elements) for line in shuffled_plan
            ),
        )

        topology = read_page_table_topology(document.pages[0])
        assert topology is not None
        outside_ids = set(topology.diagnostics.unassigned_text_element_ids)
        outside_elements = tuple(
            element for element in text_elements if element.id in outside_ids
        )
        outside_plan = PythonDocxRenderer._plan_text_lines(
            tuple(reversed(outside_elements))
        )
        self.assertEqual(
            tuple(tuple(element.id for element in line.elements) for line in outside_plan),
            (
                (
                    "p001-text-0000",
                    "p001-text-0900",
                    "p001-text-0901",
                    "p001-text-0902",
                ),
                ("p001-text-0903", "p001-text-0904"),
                ("p001-text-0005",),
            ),
        )

        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            outputs = (root / "first.docx", root / "second.docx")
            results = []
            for output in outputs:
                results.append(
                    render_docx(
                        document,
                        output,
                        renderer=PythonDocxRenderer(),
                        policy=RenderPolicy.STRICT,
                    )
                )
            with ZipFile(outputs[0]) as first, ZipFile(outputs[1]) as second:
                first_xml = first.read("word/document.xml")
                second_xml = second.read("word/document.xml")

        self.assertEqual(first_xml, second_xml)
        self.assertEqual(
            results[0].report.rendered_element_ids,
            results[1].report.rendered_element_ids,
        )
        self.assertEqual(
            results[0].report.native_table_consumed_element_ids,
            results[1].report.native_table_consumed_element_ids,
        )

    def test_horizontal_and_vertical_merges_are_native_and_observable(self):
        for merge, expected_xpath, expected_values in (
            ("horizontal", ".//w:gridSpan", ["2"]),
            ("vertical", ".//w:vMerge", ["restart", None]),
        ):
            with self.subTest(merge=merge):
                document = _topology_document(merge)
                with tempfile.TemporaryDirectory() as temporary_directory:
                    output = Path(temporary_directory) / f"{merge}.docx"
                    render_docx(
                        document,
                        output,
                        renderer=PythonDocxRenderer(),
                        policy=RenderPolicy.STRICT,
                    )
                    reopened = open_docx(output)
                    observation = PythonDocxObserver().observe(output)

                values = [
                    node.get(qn("w:val"))
                    for node in reopened.tables[0]._tbl.xpath(expected_xpath)
                ]
                self.assertEqual(values, expected_values)
                table_cells = {
                    relationship.target
                    for relationship in observation.relationships
                    if relationship.kind.value == "containment"
                    and relationship.source == TABLE_ID
                }
                self.assertEqual(len(table_cells), 3)
                self.assertEqual(observation.errors, ())

    def test_structure_relationships_raise_restoration_above_threshold(self):
        document = _topology_document()
        relationships = build_docx_structure_relationships(document)
        reference = build_evaluation_reference(
            document,
            reference_id="native-table-unit",
            reviewed=True,
            relationships=relationships,
        )
        with tempfile.TemporaryDirectory() as temporary_directory:
            output = Path(temporary_directory) / "evaluated.docx"
            render_result = render_docx(
                document,
                output,
                renderer=PythonDocxRenderer(),
                policy=RenderPolicy.STRICT,
            )
            evaluation = evaluate_restoration(
                document,
                reference,
                output,
                render_result.report,
                observer=PythonDocxObserver(),
            )

        components = {item.name: item.score for item in evaluation.components}
        self.assertEqual(components["text_similarity"], 100.0)
        self.assertGreater(components["structure_similarity"], 0.0)
        self.assertGreaterEqual(evaluation.overall_score, 70.0)

    def test_normalized_document_xml_is_deterministic(self):
        document = _topology_document("horizontal")
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            paths = (root / "first.docx", root / "second.docx")
            for path in paths:
                render_docx(
                    document,
                    path,
                    renderer=PythonDocxRenderer(),
                    policy=RenderPolicy.STRICT,
                )
            with ZipFile(paths[0]) as first, ZipFile(paths[1]) as second:
                first_xml = first.read("word/document.xml")
                second_xml = second.read("word/document.xml")

        self.assertEqual(first_xml, second_xml)


if __name__ == "__main__":
    unittest.main()
