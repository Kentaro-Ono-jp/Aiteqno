import base64
import hashlib
import json
import math
import os
import shutil
import subprocess
import sys
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import Mock, patch

from docx import Document as open_docx
from docx.oxml.ns import qn
from PIL import Image

from aiteqno.domain import (
    TABLE_TOPOLOGY_EXTENSION_KEY,
    DocumentIR,
    TablePrimitiveRole,
    read_page_table_topology,
)
from aiteqno.ports import (
    OcrRuntimeEvidence,
    OcrTrainedDataEvidence,
    SourceBaselineReference,
)
from scripts.run_real_baseline import (
    _copytree_new_atomic,
    _read_fixture,
    _select_language_profile,
    _select_ocr_input,
    _select_padding_input,
    _select_region_grouping,
    run as run_real_baseline,
)


REPOSITORY_ROOT = Path(__file__).resolve().parents[1]
FIXTURE_DIRECTORY = (
    Path(__file__).resolve().parent
    / "fixtures"
    / "baseline"
    / "synthetic-dense-japanese-form-v1"
)
LEGACY_UNLICENSED_SOURCE = REPOSITORY_ROOT / "input" / "form_blank_testClinic_v1.png"
EXPECTED_SOURCE_SHA256 = (
    "df0b724d8fcc1b5d5e0483a60401c2cb3882675f71d1e37ecdbcff9e687ffc25"
)
EXPECTED_REFERENCE_SHA256 = (
    "45d3322ee7eea3d86fe981d93dba5cc9ac83b27ca638259051a62868c8f15a31"
)


def _mock_padding_evidence(*, candidate: bool):
    evidence = Mock(effective_ocr_dpi=96)
    evidence.to_dict.return_value = {
        "schema_version": "1.0",
        "padding_version": "tesseract-crop-padding-v1",
        "enabled": candidate,
        "configured_padding_pixels": 2 if candidate else 0,
        "effective_ocr_dpi": 96,
        "crops": [],
    }
    return evidence


def _mock_comparison(*, decision: str, reason: str):
    comparison = Mock(reasons=(reason,))
    comparison.decision.value = decision
    report = {
        "schema_version": "1.0",
        "decision": decision,
        "reasons": [reason],
        "recovery": {
            "protected_literals": {"items": [], "lost": []},
        },
        "multilingual_smoke": {"status": "pass"},
        "singleton_observations": {"status": "pass"},
    }
    comparison.to_dict.return_value = report
    comparison.to_json.return_value = json.dumps(report, sort_keys=True)
    return comparison


def _mock_invocation(label: str):
    evidence = Mock(parameters_digest="3" * 64)
    evidence.to_dict.return_value = {
        "schema_version": "1.0",
        "label": label,
        "parameters_digest": "3" * 64,
    }
    return evidence


def _mock_language_runtime():
    return Mock(
        control_backend=Mock(),
        candidate_backend=Mock(),
        smoke_backend=Mock(),
        control_invocations=[_mock_invocation("language-control")],
        candidate_invocations=[_mock_invocation("language-candidate")],
        smoke_invocations=[_mock_invocation("multilingual-smoke")],
    )


def _mock_grouping_plan(label: str):
    evidence = Mock()
    evidence.to_dict.return_value = {
        "schema_version": "1.0",
        "label": label,
        "plan_digest": ("4" if label == "control" else "5") * 64,
    }
    return evidence


def _mock_grouping_runtime():
    return Mock(
        control_backend=Mock(),
        candidate_backend=Mock(),
        control_invocations=[_mock_invocation("grouping-control")],
        candidate_invocations=[_mock_invocation("grouping-candidate")],
        control_plans=[_mock_grouping_plan("control")],
        candidate_plans=[_mock_grouping_plan("candidate")],
    )


def _mock_smoke_run():
    return Mock(source_sha256="c" * 64, observed_text="AITEQNO 2026 患者番号")


def _observation_bundle_kind(bundle_directory: Path) -> str:
    if bundle_directory.name in {"control-bundle", "candidate-bundle"}:
        return bundle_directory.name
    if bundle_directory.name != "bundle":
        raise AssertionError(f"unexpected observation bundle: {bundle_directory}")
    return f"{bundle_directory.parent.parent.name}/{bundle_directory.parent.name}"


class RealDocumentBaselineContractTest(unittest.TestCase):
    def test_original_mit_source_and_reference_are_reviewed_and_hash_pinned(self):
        manifest = json.loads(
            (FIXTURE_DIRECTORY / "manifest.json").read_text(encoding="utf-8")
        )
        source = base64.b64decode(
            (FIXTURE_DIRECTORY / "source.png.b64").read_bytes().strip(),
            validate=True,
        )
        reference = SourceBaselineReference.from_json(
            (FIXTURE_DIRECTORY / "reference.json").read_bytes()
        )

        self.assertEqual(hashlib.sha256(source).hexdigest(), EXPECTED_SOURCE_SHA256)
        self.assertEqual(
            hashlib.sha256(
                (FIXTURE_DIRECTORY / "reference.json").read_bytes()
            ).hexdigest(),
            EXPECTED_REFERENCE_SHA256,
        )
        self.assertEqual(manifest["source"]["sha256"], EXPECTED_SOURCE_SHA256)
        self.assertEqual(reference.source_sha256, EXPECTED_SOURCE_SHA256)
        self.assertEqual(manifest["license"], "MIT")
        self.assertTrue(manifest["redistribution_allowed"])
        self.assertFalse(manifest["contains_personal_data"])
        self.assertEqual(manifest["review"]["status"], "reviewed")
        self.assertTrue(reference.reviewed)
        self.assertEqual(
            manifest["quality_contract"]["expected_current_state"],
            "fail",
        )
        with Image.open(BytesIO(source)) as image:
            self.assertEqual(image.size, (700, 991))
            self.assertEqual(image.format, "PNG")

    def test_reference_contract_has_component_floors_and_manual_gates(self):
        manifest = json.loads(
            (FIXTURE_DIRECTORY / "manifest.json").read_text(encoding="utf-8")
        )
        reference = SourceBaselineReference.from_json(
            (FIXTURE_DIRECTORY / "reference.json").read_bytes()
        )

        quality = manifest["quality_contract"]
        self.assertEqual(quality["overall_minimum"], 70.0)
        self.assertEqual(
            quality["component_minimums"],
            {
                "text_character_accuracy": 70.0,
                "logical_block_coverage": 60.0,
                "structure_similarity": 60.0,
                "geometry_similarity": 50.0,
            },
        )
        self.assertGreaterEqual(len(reference.text_regions), 40)
        self.assertGreaterEqual(len(reference.structural_items), 25)
        self.assertGreaterEqual(len(reference.relationships), 12)
        self.assertEqual(
            {item.kind.value for item in reference.relationships},
            {"reading_order", "containment", "adjacency"},
        )
        self.assertTrue(all(item.essential for item in reference.relationships))
        self.assertGreaterEqual(len(reference.essential_text_anchors), 10)
        self.assertEqual(reference.expected_page_count, 1)
        self.assertEqual(
            set(reference.required_manual_checks),
            {
                "no_fatal_text_overlap",
                "no_text_clipping",
                "layout_human_usable",
                "word_open_edit_save",
            },
        )

    def test_unverified_legacy_source_is_not_retained_in_current_tree(self):
        excluded = json.loads(
            (FIXTURE_DIRECTORY.parent / "excluded-sources.json").read_text(
                encoding="utf-8"
            )
        )
        record = excluded["sources"][0]

        self.assertFalse(LEGACY_UNLICENSED_SOURCE.exists())
        self.assertEqual(record["fixture_status"], "excluded")
        self.assertEqual(record["redistribution_status"], "unverified")
        self.assertEqual(
            record["sha256"],
            "f57c5246c248cf6a2f7abe548225d65ae9ab376d83a64d4183224a3e0f395369",
        )

    def test_runner_rejects_fixture_contract_drift_before_runtime_execution(self):
        cases = (
            ("fixture id", "manifest", ("fixture_id",), "other-fixture", "fixture_id"),
            (
                "source traversal",
                "manifest",
                ("source", "path"),
                "../source.png.b64",
                "source.path",
            ),
            (
                "source encoding",
                "manifest",
                ("source", "encoding"),
                "hex",
                "source.encoding",
            ),
            (
                "source hash",
                "manifest",
                ("source", "sha256"),
                "0" * 64,
                "source.sha256",
            ),
            (
                "manifest source width",
                "manifest",
                ("source", "pixel_width"),
                701,
                "manifest source dimensions",
            ),
            (
                "reference source height",
                "reference",
                ("source_dimensions", "pixel_height"),
                992,
                "reference JSON SHA-256",
            ),
            (
                "reference text",
                "reference",
                ("essential_text_anchors", 0),
                "都合のよい正解",
                "reference JSON SHA-256",
            ),
            (
                "ocr provider",
                "manifest",
                ("ocr_contract", "provider"),
                "other",
                "ocr provider",
            ),
            (
                "ocr languages",
                "manifest",
                ("ocr_contract", "languages"),
                ["eng", "jpn"],
                "ocr languages",
            ),
            (
                "ocr psm",
                "manifest",
                ("ocr_contract", "page_segmentation_mode"),
                3,
                "page segmentation mode",
            ),
            (
                "ocr oem",
                "manifest",
                ("ocr_contract", "engine_mode"),
                1,
                "engine mode",
            ),
            (
                "manifest source dpi",
                "manifest",
                ("source", "dpi"),
                300,
                "manifest source DPI",
            ),
            (
                "ocr source dpi",
                "manifest",
                ("ocr_contract", "source_dpi"),
                300,
                "ocr source DPI",
            ),
            (
                "quality page count",
                "manifest",
                ("quality_contract", "expected_page_count"),
                2,
                "expected_page_count",
            ),
            (
                "ocr text threshold",
                "manifest",
                (
                    "quality_contract",
                    "component_minimums",
                    "text_character_accuracy",
                ),
                60.0,
                "OCR text character minimum",
            ),
            (
                "ocr block threshold",
                "manifest",
                (
                    "quality_contract",
                    "component_minimums",
                    "logical_block_coverage",
                ),
                50.0,
                "OCR logical block coverage minimum",
            ),
            (
                "ocr anchor threshold",
                "manifest",
                ("quality_contract", "essential_anchor_recall"),
                90.0,
                "OCR essential anchor recall",
            ),
        )

        with tempfile.TemporaryDirectory() as temporary_directory:
            temporary_root = Path(temporary_directory)
            for name, document_name, key_path, replacement, expected_message in cases:
                with self.subTest(name=name):
                    fixture = temporary_root / name.replace(" ", "-")
                    shutil.copytree(FIXTURE_DIRECTORY, fixture)
                    document_path = fixture / f"{document_name}.json"
                    document = json.loads(document_path.read_text(encoding="utf-8"))
                    target = document
                    for key in key_path[:-1]:
                        target = target[key]
                    target[key_path[-1]] = replacement
                    document_path.write_text(
                        json.dumps(document, ensure_ascii=False, indent=2) + "\n",
                        encoding="utf-8",
                    )

                    with self.assertRaisesRegex(RuntimeError, expected_message):
                        _read_fixture(fixture)

    def test_existing_output_directory_is_not_modified_by_cli_failure_handling(self):
        with tempfile.TemporaryDirectory() as temporary_directory:
            output = Path(temporary_directory) / "already-exists"
            output.mkdir()
            sentinel = output / "sentinel.txt"
            sentinel.write_text("retain exactly\n", encoding="utf-8")
            before = {
                item.relative_to(output): item.read_bytes()
                for item in output.rglob("*")
                if item.is_file()
            }

            completed = subprocess.run(
                [
                    sys.executable,
                    str(REPOSITORY_ROOT / "scripts" / "run_real_baseline.py"),
                    "--output",
                    str(output),
                ],
                cwd=REPOSITORY_ROOT,
                check=False,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=30,
            )

            self.assertEqual(completed.returncode, 1)
            self.assertIn("baseline output already exists", completed.stderr)
            after = {
                item.relative_to(output): item.read_bytes()
                for item in output.rglob("*")
                if item.is_file()
            }
            self.assertEqual(after, before)
            self.assertFalse((output / "operational-error.json").exists())

    def test_ocr_input_selection_policy_covers_all_decision_classes(self):
        self.assertEqual(_select_ocr_input("supported"), "control")
        self.assertEqual(_select_ocr_input("inconclusive"), "control")
        self.assertEqual(_select_ocr_input("regressed"), "control")
        with self.assertRaisesRegex(RuntimeError, "comparison is invalid"):
            _select_ocr_input("invalid")
        with self.assertRaisesRegex(RuntimeError, "unknown OCR input"):
            _select_ocr_input("unexpected")
        self.assertEqual(_select_padding_input("supported"), "two-pixel-padding")
        self.assertEqual(_select_padding_input("inconclusive"), "control")
        self.assertEqual(_select_padding_input("regressed"), "control")
        with self.assertRaisesRegex(RuntimeError, "padding comparison is invalid"):
            _select_padding_input("invalid")
        with self.assertRaisesRegex(RuntimeError, "unknown OCR padding"):
            _select_padding_input("unexpected")
        self.assertEqual(_select_language_profile("supported"), "jpn")
        self.assertEqual(_select_language_profile("inconclusive"), "jpn-eng")
        self.assertEqual(_select_language_profile("regressed"), "jpn-eng")
        with self.assertRaisesRegex(
            RuntimeError, "language-profile comparison is invalid"
        ):
            _select_language_profile("invalid")
        with self.assertRaisesRegex(RuntimeError, "unknown OCR language-profile"):
            _select_language_profile("unexpected")
        self.assertEqual(
            _select_region_grouping("supported"),
            "geometry-line-groups",
        )
        self.assertEqual(_select_region_grouping("inconclusive"), "single-regions")
        self.assertEqual(_select_region_grouping("regressed"), "single-regions")
        with self.assertRaisesRegex(RuntimeError, "grouping comparison is invalid"):
            _select_region_grouping("invalid")
        with self.assertRaisesRegex(RuntimeError, "unknown OCR region-grouping"):
            _select_region_grouping("unexpected")

    def test_selected_bundle_publication_never_exposes_a_partial_bundle(self):
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            source = root / "control-bundle"
            source.mkdir()
            (source / "document.ir.json").write_text("{}\n", encoding="utf-8")
            destination = root / "bundle"
            staging = root / ".bundle.staging"

            def partial_copy(_source, target):
                target.mkdir()
                (target / "partial.txt").write_text("partial", encoding="utf-8")
                raise OSError("forced copy failure")

            with patch(
                "scripts.run_real_baseline.shutil.copytree",
                side_effect=partial_copy,
            ):
                with self.assertRaisesRegex(OSError, "forced copy failure"):
                    _copytree_new_atomic(source, destination)

            self.assertFalse(destination.exists())
            self.assertFalse(staging.exists())

            destination.mkdir()
            sentinel = destination / "sentinel.txt"
            sentinel.write_text("retain exactly\n", encoding="utf-8")
            with self.assertRaisesRegex(FileExistsError, "already exists"):
                _copytree_new_atomic(source, destination)
            self.assertEqual(sentinel.read_text(encoding="utf-8"), "retain exactly\n")

    def test_invalid_comparison_retains_ab_evidence_and_stops_before_publication(self):
        document_payload = (
            Path(__file__).resolve().parent
            / "fixtures"
            / "document_ir"
            / "canonical.document.ir.json"
        ).read_bytes()
        documents = iter(
            (
                DocumentIR.from_json(document_payload),
                DocumentIR.from_json(document_payload),
                DocumentIR.from_json(document_payload),
                DocumentIR.from_json(document_payload),
                DocumentIR.from_json(document_payload),
                DocumentIR.from_json(document_payload),
                DocumentIR.from_json(document_payload),
            )
        )

        def transform(*, enabled, target_dpi, effective_ocr_dpi):
            evidence = Mock(effective_ocr_dpi=effective_ocr_dpi)
            evidence.to_dict.return_value = {
                "schema_version": "1.0",
                "enabled": enabled,
                "target_dpi": target_dpi,
                "effective_ocr_dpi": effective_ocr_dpi,
            }
            return evidence

        control_transform = transform(
            enabled=False,
            target_dpi=None,
            effective_ocr_dpi=96,
        )
        candidate_transform = transform(
            enabled=True,
            target_dpi=300,
            effective_ocr_dpi=300,
        )
        control_padding = _mock_padding_evidence(candidate=False)
        candidate_padding = _mock_padding_evidence(candidate=True)
        quality_results = []
        for label in (
            "control",
            "resolution-candidate",
            "padding-candidate",
            "language-control",
            "language-candidate",
            "grouping-control",
            "grouping-candidate",
        ):
            result = Mock()
            result.to_json.return_value = json.dumps(
                {"schema_version": "1.0", "observation": label},
                sort_keys=True,
            )
            quality_results.append(result)
        comparison = Mock(reasons=("comparison_invalid:transform_integrity",))
        comparison.decision.value = "invalid"
        comparison.to_json.return_value = json.dumps(
            {
                "schema_version": "1.0",
                "decision": "invalid",
                "reasons": ["comparison_invalid:transform_integrity"],
            },
            sort_keys=True,
        )
        padding_comparison = _mock_comparison(
            decision="supported",
            reason="all_ocr_crop_padding_adoption_conditions_pass",
        )
        language_comparison = _mock_comparison(
            decision="supported",
            reason="all_japanese_only_language_profile_adoption_conditions_pass",
        )
        grouping_comparison = _mock_comparison(
            decision="supported",
            reason="all_geometry_only_region_grouping_adoption_conditions_pass",
        )
        render_docx_mock = Mock()

        with tempfile.TemporaryDirectory() as temporary_directory:
            output = Path(temporary_directory) / "invalid-comparison"

            def completed_extraction(_source_data, bundle_directory, **_kwargs):
                bundle_directory.mkdir(parents=True)
                (bundle_directory / "document.ir.json").write_bytes(document_payload)
                return Mock(document=next(documents), diagnostics=())

            with (
                patch(
                    "scripts.run_real_baseline._runtime",
                    return_value=(
                        Mock(),
                        Mock(),
                        Mock(),
                        Mock(),
                        Mock(),
                        [control_transform],
                        [candidate_transform],
                        [control_padding],
                        [candidate_padding],
                    ),
                ),
                patch(
                    "scripts.run_real_baseline._environment_record",
                    return_value={"phase": "deterministic-test"},
                ),
                patch(
                    "scripts.run_real_baseline._language_runtime",
                    return_value=_mock_language_runtime(),
                ),
                patch(
                    "scripts.run_real_baseline._grouping_runtime",
                    return_value=_mock_grouping_runtime(),
                ),
                patch(
                    "scripts.run_real_baseline._run_language_smoke",
                    return_value=(_mock_smoke_run(), []),
                ),
                patch(
                    "scripts.run_real_baseline.extract_png",
                    side_effect=completed_extraction,
                ),
                patch(
                    "scripts.run_real_baseline._ocr_runtime_evidence",
                    return_value=Mock(),
                ),
                patch(
                    "scripts.run_real_baseline._runtime_evidence_from_invocation",
                    return_value=Mock(),
                ),
                patch(
                    "scripts.run_real_baseline._evaluate_ocr_run",
                    side_effect=quality_results,
                ),
                patch(
                    "scripts.run_real_baseline.OcrResolutionRun",
                    side_effect=lambda **values: values,
                ),
                patch(
                    "scripts.run_real_baseline.OcrExperimentRun",
                    side_effect=lambda **values: values,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_resolution",
                    return_value=comparison,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_padding",
                    return_value=padding_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_language_profile",
                    return_value=language_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_region_grouping",
                    return_value=grouping_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.render_docx",
                    render_docx_mock,
                ),
            ):
                with self.assertRaisesRegex(RuntimeError, "comparison is invalid"):
                    run_real_baseline(FIXTURE_DIRECTORY, output, expect_state=None)

            render_docx_mock.assert_not_called()
            for relative_path in (
                "ocr-quality-control-evaluation.json",
                "ocr-quality-evaluation.json",
                "ocr-input-transform.json",
                "ocr-resolution-comparison.json",
                "ocr-padding/control/ocr-quality-evaluation.json",
                "ocr-padding/candidate/ocr-quality-evaluation.json",
                "ocr-padding/crop-padding-evidence.json",
                "ocr-padding/comparison.json",
                "ocr-language/control/ocr-quality-evaluation.json",
                "ocr-language/candidate/ocr-quality-evaluation.json",
                "ocr-language/control/runtime-config-evidence.json",
                "ocr-language/candidate/runtime-config-evidence.json",
                "ocr-language/comparison.json",
                "ocr-language/protected-literal-diagnostics.json",
                "ocr-language/multilingual-smoke.json",
                "ocr-region-grouping/control/ocr-quality-evaluation.json",
                "ocr-region-grouping/candidate/ocr-quality-evaluation.json",
                "ocr-region-grouping/control/runtime-config-evidence.json",
                "ocr-region-grouping/candidate/runtime-config-evidence.json",
                "ocr-region-grouping/region-plan-evidence.json",
                "ocr-region-grouping/comparison.json",
                "ocr-region-grouping/protected-literal-diagnostics.json",
                "ocr-region-grouping/singleton-observations.json",
                "ocr-region-grouping/environment-evidence.json",
            ):
                self.assertTrue((output / relative_path).is_file(), relative_path)
            self.assertTrue((output / "control-bundle" / "document.ir.json").is_file())
            self.assertTrue(
                (output / "candidate-bundle" / "document.ir.json").is_file()
            )
            self.assertTrue(
                (
                    output / "ocr-padding" / "candidate" / "bundle" / "document.ir.json"
                ).is_file()
            )
            self.assertFalse((output / "bundle").exists())
            self.assertFalse((output / ".bundle.staging").exists())
            failure = json.loads(
                (output / "operational-error.json").read_text(encoding="utf-8")
            )
            self.assertEqual(failure["error_type"], "RuntimeError")
            self.assertIn("comparison is invalid", failure["message"])

    def test_new_output_owned_by_runner_retains_preflight_failure_evidence(self):
        with tempfile.TemporaryDirectory() as temporary_directory:
            temporary_root = Path(temporary_directory)
            fixture = temporary_root / "invalid-fixture"
            shutil.copytree(FIXTURE_DIRECTORY, fixture)
            manifest_path = fixture / "manifest.json"
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            manifest["source"]["encoding"] = "hex"
            manifest_path.write_text(
                json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )
            output = temporary_root / "new-output"

            with self.assertRaisesRegex(RuntimeError, "source.encoding"):
                run_real_baseline(fixture, output, expect_state=None)

            evidence = json.loads(
                (output / "operational-error.json").read_text(encoding="utf-8")
            )
            self.assertEqual(evidence["error_type"], "RuntimeError")
            self.assertIn("source.encoding", evidence["message"])
            self.assertEqual(
                {item.name for item in output.iterdir()},
                {"operational-error.json"},
            )

    def test_supported_padding_candidate_is_adopted_for_downstream(self):
        document_payload = (
            Path(__file__).resolve().parent
            / "fixtures"
            / "document_ir"
            / "canonical.document.ir.json"
        ).read_bytes()
        control_ir = DocumentIR.from_json(document_payload)
        candidate_ir = DocumentIR.from_json(document_payload)
        language_control_ir = DocumentIR.from_json(document_payload)
        language_candidate_ir = DocumentIR.from_json(document_payload)
        grouping_control_ir = DocumentIR.from_json(document_payload)
        grouping_candidate_ir = DocumentIR.from_json(document_payload)
        padding_ir = DocumentIR.from_json(document_payload)

        def runtime(effective_ocr_dpi):
            return OcrRuntimeEvidence(
                provider="tesseract",
                provider_version="5.0.0-test",
                executable="test-tesseract",
                languages=("jpn", "eng"),
                page_segmentation_mode=6,
                engine_mode=3,
                effective_ocr_dpi=effective_ocr_dpi,
                source_dpi_x=96.012,
                source_dpi_y=96.012,
                traineddata=(
                    OcrTrainedDataEvidence(
                        language="jpn",
                        size_bytes=1,
                        sha256="1" * 64,
                    ),
                    OcrTrainedDataEvidence(
                        language="eng",
                        size_bytes=1,
                        sha256="2" * 64,
                    ),
                ),
                operating_system="deterministic-test-os",
                python_version="3.test",
            )

        control_transform = Mock(effective_ocr_dpi=96)
        control_transform.to_dict.return_value = {
            "schema_version": "1.0",
            "enabled": False,
            "target_dpi": None,
            "effective_ocr_dpi": 96,
        }
        candidate_transform = Mock(effective_ocr_dpi=300)
        candidate_transform.to_dict.return_value = {
            "schema_version": "1.0",
            "enabled": True,
            "target_dpi": 300,
            "effective_ocr_dpi": 300,
        }
        comparison = Mock(reasons=())
        comparison.decision.value = "supported"
        comparison.to_json.return_value = json.dumps(
            {
                "schema_version": "1.0",
                "decision": "supported",
            },
            sort_keys=True,
        )
        control_padding = _mock_padding_evidence(candidate=False)
        candidate_padding = _mock_padding_evidence(candidate=True)
        padding_comparison = _mock_comparison(
            decision="supported",
            reason="all_ocr_crop_padding_adoption_conditions_pass",
        )
        language_comparison = _mock_comparison(
            decision="supported",
            reason="all_japanese_only_language_profile_adoption_conditions_pass",
        )
        grouping_comparison = _mock_comparison(
            decision="supported",
            reason="all_geometry_only_region_grouping_adoption_conditions_pass",
        )

        with tempfile.TemporaryDirectory() as temporary_directory:
            output = Path(temporary_directory) / "downstream-failure"
            events = []

            def completed_extraction(_source_data, bundle_directory, **_kwargs):
                bundle_directory.mkdir(parents=True)
                kind = _observation_bundle_kind(bundle_directory)
                if kind == "control-bundle":
                    label = "control-bundle"
                    document = control_ir
                elif kind == "candidate-bundle":
                    label = "resolution-candidate-bundle"
                    document = candidate_ir
                elif kind == "ocr-padding/candidate":
                    label = "padding-candidate-bundle"
                    document = padding_ir
                elif kind == "ocr-language/control":
                    label = "language-control-bundle"
                    document = language_control_ir
                elif kind == "ocr-language/candidate":
                    label = "language-candidate-bundle"
                    document = language_candidate_ir
                elif kind == "ocr-region-grouping/control":
                    label = "grouping-control-bundle"
                    document = grouping_control_ir
                else:
                    label = "grouping-candidate-bundle"
                    document = grouping_candidate_ir
                events.append(f"extract:{label}")
                (bundle_directory / "selection.txt").write_text(
                    label,
                    encoding="utf-8",
                )
                return Mock(document=document, diagnostics=())

            def completed_comparison(*_args, **_kwargs):
                events.append("compare-resolution")
                return comparison

            def completed_padding_comparison(*_args, **_kwargs):
                events.append("compare-padding")
                return padding_comparison

            def completed_language_comparison(*_args, **_kwargs):
                events.append("compare-language")
                return language_comparison

            def completed_grouping_comparison(*_args, **_kwargs):
                events.append("compare-grouping")
                return grouping_comparison

            def completed_topology(document):
                events.append("topology")
                self.assertIs(document, grouping_candidate_ir)
                return document

            def failed_render(document, path, **_kwargs):
                events.append("render")
                self.assertIs(document, grouping_candidate_ir)
                self.assertTrue(path.parent.samefile(output / "bundle"))
                self.assertEqual(path.name, "reconstructed.docx")
                raise RuntimeError("forced downstream render failure")

            with (
                patch(
                    "scripts.run_real_baseline._runtime",
                    return_value=(
                        Mock(),
                        Mock(),
                        Mock(),
                        Mock(),
                        Mock(),
                        [control_transform],
                        [candidate_transform],
                        [control_padding],
                        [candidate_padding],
                    ),
                ),
                patch(
                    "scripts.run_real_baseline._environment_record",
                    return_value={"phase": "deterministic-test"},
                ),
                patch(
                    "scripts.run_real_baseline._language_runtime",
                    return_value=_mock_language_runtime(),
                ),
                patch(
                    "scripts.run_real_baseline._grouping_runtime",
                    return_value=_mock_grouping_runtime(),
                ),
                patch(
                    "scripts.run_real_baseline._run_language_smoke",
                    return_value=(_mock_smoke_run(), []),
                ),
                patch(
                    "scripts.run_real_baseline.extract_png",
                    side_effect=completed_extraction,
                ),
                patch(
                    "scripts.run_real_baseline._ocr_runtime_evidence",
                    side_effect=(runtime(96), runtime(300), runtime(96)),
                ),
                patch(
                    "scripts.run_real_baseline._runtime_evidence_from_invocation",
                    side_effect=(runtime(96), runtime(96), runtime(96), runtime(96)),
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_resolution",
                    side_effect=completed_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_padding",
                    side_effect=completed_padding_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_language_profile",
                    side_effect=completed_language_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_region_grouping",
                    side_effect=completed_grouping_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.infer_table_topology",
                    side_effect=completed_topology,
                ),
                patch(
                    "scripts.run_real_baseline.render_docx",
                    side_effect=failed_render,
                ),
            ):
                with self.assertRaisesRegex(
                    RuntimeError,
                    "forced downstream render failure",
                ):
                    run_real_baseline(FIXTURE_DIRECTORY, output, expect_state=None)

            control_report = json.loads(
                (output / "ocr-quality-control-evaluation.json").read_text(
                    encoding="utf-8"
                )
            )
            report = json.loads(
                (output / "ocr-quality-evaluation.json").read_text(encoding="utf-8")
            )
            self.assertEqual(
                control_report["runtime"]["configuration"]["effective_ocr_dpi"],
                96,
            )
            self.assertEqual(report["schema_version"], "1.0")
            self.assertEqual(report["scope"]["text_source"], "candidate_ir")
            self.assertEqual(
                report["runtime"]["configuration"]["effective_ocr_dpi"],
                300,
            )
            self.assertEqual(report["state"], "fail")
            padding_report = json.loads(
                (
                    output / "ocr-padding" / "candidate" / "ocr-quality-evaluation.json"
                ).read_text(encoding="utf-8")
            )
            self.assertEqual(
                padding_report["runtime"]["configuration"]["effective_ocr_dpi"],
                96,
            )
            transform = json.loads(
                (output / "ocr-input-transform.json").read_text(encoding="utf-8")
            )
            self.assertFalse(transform["control"]["enabled"])
            self.assertTrue(transform["candidate"]["enabled"])
            resolution = json.loads(
                (output / "ocr-resolution-comparison.json").read_text(encoding="utf-8")
            )
            self.assertEqual(resolution["decision"], "supported")
            padding = json.loads(
                (output / "ocr-padding" / "comparison.json").read_text(encoding="utf-8")
            )
            self.assertEqual(padding["decision"], "supported")
            language = json.loads(
                (output / "ocr-language" / "comparison.json").read_text(
                    encoding="utf-8"
                )
            )
            self.assertEqual(language["decision"], "supported")
            self.assertTrue(
                (
                    output / "ocr-language" / "control" / "ocr-quality-evaluation.json"
                ).is_file()
            )
            self.assertTrue(
                (
                    output
                    / "ocr-language"
                    / "candidate"
                    / "ocr-quality-evaluation.json"
                ).is_file()
            )
            self.assertEqual(
                (output / "bundle" / "selection.txt").read_text(encoding="utf-8"),
                "grouping-candidate-bundle",
            )
            self.assertEqual(
                events,
                [
                    "extract:control-bundle",
                    "extract:resolution-candidate-bundle",
                    "extract:padding-candidate-bundle",
                    "extract:language-control-bundle",
                    "extract:language-candidate-bundle",
                    "extract:grouping-control-bundle",
                    "extract:grouping-candidate-bundle",
                    "compare-resolution",
                    "compare-padding",
                    "compare-language",
                    "compare-grouping",
                    "topology",
                    "render",
                ],
            )
            self.assertFalse((output / "source-quality-evaluation.json").exists())
            self.assertFalse((output / "bundle" / "reconstructed.docx").exists())
            failure = json.loads(
                (output / "operational-error.json").read_text(encoding="utf-8")
            )
            self.assertEqual(failure["error_type"], "RuntimeError")
            self.assertIn("forced downstream render failure", failure["message"])

    def test_inconclusive_ocr_comparison_selects_control_for_downstream(self):
        document_payload = (
            Path(__file__).resolve().parent
            / "fixtures"
            / "document_ir"
            / "canonical.document.ir.json"
        ).read_bytes()
        control_ir = DocumentIR.from_json(document_payload)
        candidate_ir = DocumentIR.from_json(document_payload)
        language_control_ir = DocumentIR.from_json(document_payload)
        language_candidate_ir = DocumentIR.from_json(document_payload)
        grouping_control_ir = DocumentIR.from_json(document_payload)
        grouping_candidate_ir = DocumentIR.from_json(document_payload)

        def runtime(effective_ocr_dpi):
            return OcrRuntimeEvidence(
                provider="tesseract",
                provider_version="5.0.0-test",
                executable="test-tesseract",
                languages=("jpn", "eng"),
                page_segmentation_mode=6,
                engine_mode=3,
                effective_ocr_dpi=effective_ocr_dpi,
                source_dpi_x=96.012,
                source_dpi_y=96.012,
                traineddata=(
                    OcrTrainedDataEvidence(
                        language="jpn",
                        size_bytes=1,
                        sha256="1" * 64,
                    ),
                    OcrTrainedDataEvidence(
                        language="eng",
                        size_bytes=1,
                        sha256="2" * 64,
                    ),
                ),
                operating_system="deterministic-test-os",
                python_version="3.test",
            )

        def transform(*, enabled, target_dpi, effective_ocr_dpi):
            evidence = Mock(effective_ocr_dpi=effective_ocr_dpi)
            evidence.to_dict.return_value = {
                "schema_version": "1.0",
                "enabled": enabled,
                "target_dpi": target_dpi,
                "effective_ocr_dpi": effective_ocr_dpi,
            }
            return evidence

        control_transform = transform(
            enabled=False,
            target_dpi=None,
            effective_ocr_dpi=96,
        )
        candidate_transform = transform(
            enabled=True,
            target_dpi=300,
            effective_ocr_dpi=300,
        )
        comparison = Mock(reasons=("text_accuracy_delta_below_minimum:0<1",))
        comparison.decision.value = "inconclusive"
        comparison.to_json.return_value = json.dumps(
            {
                "schema_version": "1.0",
                "decision": "inconclusive",
                "reasons": ["text_accuracy_delta_below_minimum:0<1"],
            },
            sort_keys=True,
        )
        control_padding = _mock_padding_evidence(candidate=False)
        candidate_padding = _mock_padding_evidence(candidate=True)
        padding_comparison = _mock_comparison(
            decision="inconclusive",
            reason="text_accuracy_delta_below_minimum:0<1",
        )
        language_comparison = _mock_comparison(
            decision="regressed",
            reason="regression:protected_literal:PNG",
        )
        grouping_comparison = _mock_comparison(
            decision="inconclusive",
            reason="required_grouping_target_not_newly_recovered",
        )
        with tempfile.TemporaryDirectory() as temporary_directory:
            output = Path(temporary_directory) / "unsupported-comparison"
            events = []

            def completed_extraction(_source_data, bundle_directory, **_kwargs):
                bundle_directory.mkdir(parents=True)
                kind = _observation_bundle_kind(bundle_directory)
                if kind == "control-bundle":
                    label = "control-bundle"
                    document = control_ir
                elif kind == "candidate-bundle":
                    label = "resolution-candidate-bundle"
                    document = candidate_ir
                elif kind == "ocr-padding/candidate":
                    label = "padding-candidate-bundle"
                    document = candidate_ir
                elif kind == "ocr-language/control":
                    label = "language-control-bundle"
                    document = language_control_ir
                elif kind == "ocr-language/candidate":
                    label = "language-candidate-bundle"
                    document = language_candidate_ir
                elif kind == "ocr-region-grouping/control":
                    label = "grouping-control-bundle"
                    document = grouping_control_ir
                else:
                    label = "grouping-candidate-bundle"
                    document = grouping_candidate_ir
                events.append(f"extract:{label}")
                (bundle_directory / "selection.txt").write_text(
                    label,
                    encoding="utf-8",
                )
                return Mock(document=document, diagnostics=())

            def completed_comparison(*_args, **_kwargs):
                events.append("compare-resolution")
                return comparison

            def completed_padding_comparison(*_args, **_kwargs):
                events.append("compare-padding")
                return padding_comparison

            def completed_language_comparison(*_args, **_kwargs):
                events.append("compare-language")
                return language_comparison

            def completed_grouping_comparison(*_args, **_kwargs):
                events.append("compare-grouping")
                return grouping_comparison

            def failed_render(document, path, **_kwargs):
                events.append("render")
                self.assertIs(document, language_control_ir)
                self.assertTrue(path.parent.samefile(output / "bundle"))
                self.assertEqual(path.name, "reconstructed.docx")
                raise RuntimeError("forced control downstream render failure")

            with (
                patch(
                    "scripts.run_real_baseline._runtime",
                    return_value=(
                        Mock(),
                        Mock(),
                        Mock(),
                        Mock(),
                        Mock(),
                        [control_transform],
                        [candidate_transform],
                        [control_padding],
                        [candidate_padding],
                    ),
                ),
                patch(
                    "scripts.run_real_baseline._environment_record",
                    return_value={"phase": "deterministic-test"},
                ),
                patch(
                    "scripts.run_real_baseline._language_runtime",
                    return_value=_mock_language_runtime(),
                ),
                patch(
                    "scripts.run_real_baseline._grouping_runtime",
                    return_value=_mock_grouping_runtime(),
                ),
                patch(
                    "scripts.run_real_baseline._run_language_smoke",
                    return_value=(_mock_smoke_run(), []),
                ),
                patch(
                    "scripts.run_real_baseline.extract_png",
                    side_effect=completed_extraction,
                ),
                patch(
                    "scripts.run_real_baseline._ocr_runtime_evidence",
                    side_effect=(runtime(96), runtime(300), runtime(96)),
                ),
                patch(
                    "scripts.run_real_baseline._runtime_evidence_from_invocation",
                    side_effect=(runtime(96), runtime(96), runtime(96), runtime(96)),
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_resolution",
                    side_effect=completed_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_padding",
                    side_effect=completed_padding_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_language_profile",
                    side_effect=completed_language_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.compare_ocr_region_grouping",
                    side_effect=completed_grouping_comparison,
                ),
                patch(
                    "scripts.run_real_baseline.render_docx",
                    side_effect=failed_render,
                ),
            ):
                with self.assertRaisesRegex(
                    RuntimeError,
                    "forced control downstream render failure",
                ):
                    run_real_baseline(FIXTURE_DIRECTORY, output, expect_state=None)

            self.assertTrue((output / "ocr-quality-control-evaluation.json").is_file())
            self.assertTrue((output / "ocr-quality-evaluation.json").is_file())
            self.assertTrue((output / "ocr-input-transform.json").is_file())
            comparison_report = json.loads(
                (output / "ocr-resolution-comparison.json").read_text(encoding="utf-8")
            )
            self.assertEqual(comparison_report["decision"], "inconclusive")
            padding_report = json.loads(
                (output / "ocr-padding" / "comparison.json").read_text(encoding="utf-8")
            )
            self.assertEqual(padding_report["decision"], "inconclusive")
            self.assertEqual(
                (output / "bundle" / "selection.txt").read_text(encoding="utf-8"),
                "language-control-bundle",
            )
            self.assertEqual(
                events,
                [
                    "extract:control-bundle",
                    "extract:resolution-candidate-bundle",
                    "extract:padding-candidate-bundle",
                    "extract:language-control-bundle",
                    "extract:language-candidate-bundle",
                    "extract:grouping-control-bundle",
                    "extract:grouping-candidate-bundle",
                    "compare-resolution",
                    "compare-padding",
                    "compare-language",
                    "compare-grouping",
                    "render",
                ],
            )
            self.assertFalse((output / "bundle" / "reconstructed.docx").exists())
            self.assertFalse((output / "source-quality-evaluation.json").exists())
            failure = json.loads(
                (output / "operational-error.json").read_text(encoding="utf-8")
            )
            self.assertIn(
                "forced control downstream render failure", failure["message"]
            )


@unittest.skipUnless(
    os.environ.get("AITEQNO_RUN_REAL_IMAGE_BASELINE") == "1",
    "set AITEQNO_RUN_REAL_IMAGE_BASELINE=1 for the real runtime baseline",
)
class RealDocumentBaselineIntegrationTest(unittest.TestCase):
    def test_real_tesseract_docx_libreoffice_baseline_remains_an_explicit_failure(self):
        raw_output = os.environ.get("AITEQNO_REAL_BASELINE_OUTPUT")
        self.assertTrue(raw_output, "AITEQNO_REAL_BASELINE_OUTPUT must be set")
        output = Path(raw_output).resolve(strict=False)
        self.assertFalse(output.exists(), "real baseline output must be create-only")

        completed = subprocess.run(
            [
                sys.executable,
                str(REPOSITORY_ROOT / "scripts" / "run_real_baseline.py"),
                "--output",
                str(output),
                "--expect-state",
                "fail",
            ],
            cwd=REPOSITORY_ROOT,
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=600,
        )
        self.assertEqual(
            completed.returncode,
            0,
            msg=f"stdout:\n{completed.stdout}\nstderr:\n{completed.stderr}",
        )
        summary = json.loads(
            (output / "baseline-summary.json").read_text(encoding="utf-8")
        )
        grouping_comparison = json.loads(
            (output / "ocr-region-grouping" / "comparison.json").read_text(
                encoding="utf-8"
            )
        )
        self.assertEqual(grouping_comparison["decision"], "inconclusive")
        self.assertEqual(
            grouping_comparison["reasons"],
            ["text_accuracy_delta_below_minimum:0.757575<1"],
        )
        grouping_supported = False
        expected_grouping = "single-regions"
        expected_grouping_side = "control"
        self.assertEqual(summary["final_state"], "fail")
        self.assertEqual(summary["expected_current_state"], "fail")
        self.assertEqual(
            summary["layers"]["source_to_candidate_ir_ocr"]["state"],
            "fail",
        )
        self.assertEqual(
            summary["layers"]["source_to_candidate_ir_ocr"]["text_evidence"],
            f"two-pixel-padding_jpn_{expected_grouping}_ir",
        )
        self.assertEqual(
            summary["layers"]["source_to_candidate_ir_ocr"]["selected_input"],
            "two-pixel-padding",
        )
        self.assertEqual(
            summary["layers"]["source_to_candidate_ir_ocr"]["report"],
            f"ocr-region-grouping/{expected_grouping_side}/ocr-quality-evaluation.json",
        )
        self.assertEqual(
            summary["layers"]["source_to_candidate_ir_ocr"]["selected_profile"],
            "jpn",
        )
        self.assertEqual(
            summary["layers"]["source_to_candidate_ir_ocr"]["selected_grouping"],
            expected_grouping,
        )
        self.assertEqual(
            summary["layers"]["source_to_actual_docx"]["state"],
            "fail",
        )
        self.assertEqual(
            summary["layers"]["source_to_actual_docx"]["text_evidence"],
            "rendered_visible",
        )
        source_evaluation = json.loads(
            (output / "source-quality-evaluation.json").read_text(encoding="utf-8")
        )
        self.assertEqual(source_evaluation["state"], "fail")
        self.assertEqual(source_evaluation["text_evidence"], "rendered_visible")
        self.assertEqual(
            set(source_evaluation["components"]),
            {
                "text_accuracy",
                "logical_block_coverage",
                "structure_similarity",
                "geometry_similarity",
            },
        )
        self.assertEqual(
            {gate["name"] for gate in source_evaluation["hard_gates"]},
            {
                "source_digest_matches",
                "reference_reviewed",
                "candidate_ir_page_count",
                "rendered_docx_page_count",
                "essential_text_anchors",
                "essential_logical_blocks",
                "essential_structures",
                "essential_relationships",
                "manual_checks",
            },
        )
        self.assertTrue((output / "preflight-environment.json").is_file())
        self.assertTrue((output / "environment.json").is_file())
        control_ocr_evaluation = json.loads(
            (output / "ocr-quality-control-evaluation.json").read_text(encoding="utf-8")
        )
        ocr_evaluation = json.loads(
            (output / "ocr-quality-evaluation.json").read_text(encoding="utf-8")
        )
        self.assertEqual(ocr_evaluation["schema_version"], "1.0")
        self.assertEqual(ocr_evaluation["scope"]["text_source"], "candidate_ir")
        self.assertEqual(
            ocr_evaluation["scope"]["ends_before"],
            ["docx", "preview", "libreoffice", "poppler", "rendered_page_ocr"],
        )
        self.assertEqual(ocr_evaluation["state"], "fail")
        self.assertEqual(
            ocr_evaluation["runtime"]["configuration"]["languages"],
            ["jpn", "eng"],
        )
        self.assertEqual(
            ocr_evaluation["runtime"]["configuration"]["page_segmentation_mode"],
            6,
        )
        self.assertEqual(
            ocr_evaluation["runtime"]["configuration"]["engine_mode"],
            3,
        )
        self.assertEqual(
            control_ocr_evaluation["runtime"]["configuration"]["effective_ocr_dpi"],
            96,
        )
        self.assertEqual(
            ocr_evaluation["runtime"]["configuration"]["effective_ocr_dpi"],
            300,
        )
        for report in (control_ocr_evaluation, ocr_evaluation):
            source_metadata_dpi = report["runtime"]["configuration"][
                "source_metadata_dpi"
            ]
            self.assertAlmostEqual(source_metadata_dpi["x"], 96.012, places=3)
            self.assertAlmostEqual(source_metadata_dpi["y"], 96.012, places=3)
        self.assertEqual(
            control_ocr_evaluation["source_digest"],
            ocr_evaluation["source_digest"],
        )
        self.assertEqual(
            control_ocr_evaluation["reference_id"],
            ocr_evaluation["reference_id"],
        )
        self.assertEqual(
            control_ocr_evaluation["runtime"]["provider"],
            ocr_evaluation["runtime"]["provider"],
        )
        self.assertEqual(
            control_ocr_evaluation["runtime"]["provider_version"],
            ocr_evaluation["runtime"]["provider_version"],
        )
        self.assertEqual(
            control_ocr_evaluation["runtime"]["configuration"]["languages"],
            ocr_evaluation["runtime"]["configuration"]["languages"],
        )
        self.assertEqual(
            control_ocr_evaluation["runtime"]["configuration"][
                "page_segmentation_mode"
            ],
            ocr_evaluation["runtime"]["configuration"]["page_segmentation_mode"],
        )
        self.assertEqual(
            control_ocr_evaluation["runtime"]["configuration"]["engine_mode"],
            ocr_evaluation["runtime"]["configuration"]["engine_mode"],
        )
        self.assertEqual(
            control_ocr_evaluation["runtime"]["traineddata"],
            ocr_evaluation["runtime"]["traineddata"],
        )
        self.assertEqual(
            ocr_evaluation["thresholds"]["text_character_accuracy"],
            70.0,
        )
        self.assertEqual(
            ocr_evaluation["thresholds"]["logical_block_coverage"],
            60.0,
        )
        self.assertEqual(
            ocr_evaluation["thresholds"]["essential_anchor_recall"],
            100.0,
        )
        transform = json.loads(
            (output / "ocr-input-transform.json").read_text(encoding="utf-8")
        )
        self.assertEqual(transform["schema_version"], "1.0")
        self.assertFalse(transform["control"]["enabled"])
        self.assertIsNone(transform["control"]["target_dpi"])
        self.assertEqual(transform["control"]["effective_ocr_dpi"], 96)
        self.assertTrue(transform["candidate"]["enabled"])
        self.assertEqual(transform["candidate"]["target_dpi"], 300)
        self.assertEqual(transform["candidate"]["effective_ocr_dpi"], 300)
        self.assertAlmostEqual(
            transform["candidate"]["source_effective_dpi"],
            96.012,
            places=3,
        )
        self.assertTrue(transform["candidate"]["crops"])
        self.assertEqual(
            len(transform["control"]["crops"]),
            len(transform["candidate"]["crops"]),
        )
        candidate_scale = 300 / transform["candidate"]["source_effective_dpi"]
        for control_crop, candidate_crop in zip(
            transform["control"]["crops"],
            transform["candidate"]["crops"],
        ):
            self.assertEqual(
                control_crop["region_ref"],
                candidate_crop["region_ref"],
            )
            self.assertEqual(
                control_crop["source_bbox"],
                candidate_crop["source_bbox"],
            )
            self.assertEqual(
                control_crop["source_dimensions"],
                candidate_crop["source_dimensions"],
            )
            self.assertFalse(control_crop["resized"])
            self.assertEqual(
                control_crop["working_dimensions"],
                control_crop["source_dimensions"],
            )
            self.assertTrue(candidate_crop["resized"])
            for axis in ("width", "height"):
                source_dimension = candidate_crop["source_dimensions"][axis]
                working_dimension = candidate_crop["working_dimensions"][axis]
                self.assertGreater(working_dimension, source_dimension)
                self.assertEqual(
                    working_dimension,
                    max(1, math.floor(source_dimension * candidate_scale + 0.5)),
                )
            self.assertLessEqual(
                candidate_crop["working_dimensions"]["width"]
                * candidate_crop["working_dimensions"]["height"],
                40_000_000,
            )
        comparison = json.loads(
            (output / "ocr-resolution-comparison.json").read_text(encoding="utf-8")
        )
        self.assertEqual(comparison["schema_version"], "1.0")
        self.assertEqual(comparison["decision"], "regressed")
        self.assertGreater(
            comparison["metrics"]["text_character_accuracy"]["candidate"],
            comparison["metrics"]["text_character_accuracy"]["control"],
        )
        self.assertGreater(
            comparison["metrics"]["logical_block_coverage"]["candidate"],
            comparison["metrics"]["logical_block_coverage"]["control"],
        )
        self.assertGreater(
            comparison["metrics"]["essential_anchor_recall"]["candidate"],
            comparison["metrics"]["essential_anchor_recall"]["control"],
        )
        self.assertFalse(comparison["recovery"]["anchors"]["lost"])
        self.assertEqual(
            comparison["recovery"]["logical_blocks"]["lost"],
            ["request-language-label"],
        )
        self.assertLessEqual(
            comparison["recovery"]["essential_blocks"]["unrecovered_count_delta"],
            0,
        )
        self.assertTrue(
            all(check["status"] == "pass" for check in comparison["checks"].values())
        )
        self.assertEqual(
            summary["layers"]["ocr_input_resolution_comparison"]["decision"],
            "regressed",
        )
        self.assertEqual(
            summary["layers"]["ocr_input_resolution_comparison"]["selected_input"],
            "control",
        )
        self.assertFalse(
            summary["layers"]["ocr_input_resolution_comparison"]["candidate_adopted"]
        )
        self.assertFalse(
            summary["layers"]["ocr_input_resolution_comparison"]["candidate_eligible"]
        )
        self.assertEqual(
            summary["layers"]["ocr_input_resolution_comparison"]["reasons"],
            comparison["reasons"],
        )
        self.assertEqual(
            comparison["reasons"],
            ["regression:lost_logical_block:request-language-label"],
        )
        self.assertEqual(
            summary["layers"]["candidate_300_dpi_experiment"]["adopted"],
            False,
        )
        self.assertEqual(
            summary["layers"]["candidate_300_dpi_experiment"]["eligible"],
            False,
        )
        self.assertEqual(
            summary["layers"]["candidate_ir_to_docx"]["selected_input"],
            "two-pixel-padding",
        )
        self.assertEqual(
            summary["layers"]["candidate_ir_to_docx"]["selected_profile"],
            "jpn",
        )
        self.assertEqual(
            summary["layers"]["candidate_ir_to_docx"]["selected_grouping"],
            expected_grouping,
        )
        padding_control = json.loads(
            (
                output / "ocr-padding" / "control" / "ocr-quality-evaluation.json"
            ).read_text(encoding="utf-8")
        )
        padding_candidate = json.loads(
            (
                output / "ocr-padding" / "candidate" / "ocr-quality-evaluation.json"
            ).read_text(encoding="utf-8")
        )
        padding_evidence = json.loads(
            (output / "ocr-padding" / "crop-padding-evidence.json").read_text(
                encoding="utf-8"
            )
        )
        padding_comparison = json.loads(
            (output / "ocr-padding" / "comparison.json").read_text(encoding="utf-8")
        )
        self.assertEqual(padding_comparison["decision"], "supported")
        self.assertGreaterEqual(
            padding_comparison["metrics"]["text_character_accuracy"][
                "delta_percentage_points"
            ],
            1.0,
        )
        self.assertGreaterEqual(
            padding_comparison["metrics"]["logical_block_coverage"]["candidate"],
            padding_comparison["metrics"]["logical_block_coverage"]["control"],
        )
        self.assertGreaterEqual(
            padding_comparison["metrics"]["essential_anchor_recall"]["candidate"],
            padding_comparison["metrics"]["essential_anchor_recall"]["control"],
        )
        self.assertFalse(padding_comparison["recovery"]["anchors"]["lost"])
        self.assertFalse(padding_comparison["recovery"]["logical_blocks"]["lost"])
        self.assertLessEqual(
            padding_comparison["recovery"]["essential_blocks"][
                "unrecovered_count_delta"
            ],
            0,
        )
        self.assertTrue(
            all(
                check["status"] == "pass"
                for check in padding_comparison["checks"].values()
            )
        )
        self.assertEqual(
            padding_control["source_digest"], padding_candidate["source_digest"]
        )
        self.assertEqual(
            padding_control["runtime"],
            padding_candidate["runtime"],
        )
        self.assertFalse(padding_evidence["control"]["enabled"])
        self.assertEqual(
            padding_evidence["control"]["configured_padding_pixels"],
            0,
        )
        self.assertTrue(padding_evidence["candidate"]["enabled"])
        self.assertEqual(
            padding_evidence["candidate"]["configured_padding_pixels"],
            2,
        )
        self.assertEqual(
            len(padding_evidence["control"]["crops"]),
            len(padding_evidence["candidate"]["crops"]),
        )
        self.assertTrue(padding_evidence["candidate"]["crops"])
        for control_crop, candidate_crop in zip(
            padding_evidence["control"]["crops"],
            padding_evidence["candidate"]["crops"],
        ):
            self.assertEqual(control_crop["region_ref"], candidate_crop["region_ref"])
            self.assertEqual(control_crop["source_bbox"], candidate_crop["source_bbox"])
            self.assertEqual(
                control_crop["source_dimensions"],
                candidate_crop["source_dimensions"],
            )
            self.assertEqual(control_crop["padding_pixels"], 0)
            self.assertFalse(control_crop["applied"])
            self.assertEqual(candidate_crop["padding_pixels"], 2)
            self.assertTrue(candidate_crop["applied"])
            self.assertEqual(
                candidate_crop["working_dimensions"]["width"],
                candidate_crop["source_dimensions"]["width"] + 4,
            )
            self.assertEqual(
                candidate_crop["working_dimensions"]["height"],
                candidate_crop["source_dimensions"]["height"] + 4,
            )
        self.assertEqual(
            summary["layers"]["ocr_crop_padding_comparison"]["decision"],
            "supported",
        )
        self.assertTrue(
            summary["layers"]["ocr_crop_padding_comparison"]["candidate_adopted"]
        )
        self.assertTrue(
            summary["layers"]["ocr_crop_padding_comparison"]["candidate_eligible"]
        )
        language_control = json.loads(
            (
                output / "ocr-language" / "control" / "ocr-quality-evaluation.json"
            ).read_text(encoding="utf-8")
        )
        language_candidate = json.loads(
            (
                output / "ocr-language" / "candidate" / "ocr-quality-evaluation.json"
            ).read_text(encoding="utf-8")
        )
        language_control_evidence = json.loads(
            (
                output / "ocr-language" / "control" / "runtime-config-evidence.json"
            ).read_text(encoding="utf-8")
        )
        language_candidate_evidence = json.loads(
            (
                output / "ocr-language" / "candidate" / "runtime-config-evidence.json"
            ).read_text(encoding="utf-8")
        )
        language_comparison = json.loads(
            (output / "ocr-language" / "comparison.json").read_text(encoding="utf-8")
        )
        multilingual_smoke = json.loads(
            (output / "ocr-language" / "multilingual-smoke.json").read_text(
                encoding="utf-8"
            )
        )
        self.assertEqual(language_comparison["decision"], "supported")
        self.assertGreaterEqual(
            language_comparison["metrics"]["text_character_accuracy"][
                "delta_percentage_points"
            ],
            1.0,
        )
        self.assertGreaterEqual(
            language_comparison["metrics"]["logical_block_coverage"]["candidate"],
            language_comparison["metrics"]["logical_block_coverage"]["control"],
        )
        self.assertGreaterEqual(
            language_comparison["metrics"]["essential_anchor_recall"]["candidate"],
            language_comparison["metrics"]["essential_anchor_recall"]["control"],
        )
        self.assertFalse(language_comparison["recovery"]["anchors"]["lost"])
        self.assertFalse(language_comparison["recovery"]["logical_blocks"]["lost"])
        self.assertFalse(language_comparison["recovery"]["protected_literals"]["lost"])
        self.assertLessEqual(
            language_comparison["recovery"]["essential_blocks"][
                "unrecovered_count_delta"
            ],
            0,
        )
        self.assertTrue(
            all(
                check["status"] == "pass"
                for check in language_comparison["checks"].values()
            )
        )
        self.assertEqual(
            language_control["runtime"]["configuration"]["languages"],
            ["jpn", "eng"],
        )
        self.assertEqual(
            language_candidate["runtime"]["configuration"]["languages"],
            ["jpn"],
        )
        self.assertEqual(
            [item["language"] for item in language_control_evidence["traineddata"]],
            ["jpn", "eng"],
        )
        self.assertEqual(
            [item["language"] for item in language_candidate_evidence["traineddata"]],
            ["jpn"],
        )
        self.assertEqual(
            language_control_evidence["traineddata"][0],
            language_candidate_evidence["traineddata"][0],
        )
        self.assertNotEqual(
            language_control_evidence["parameters_digest"],
            language_candidate_evidence["parameters_digest"],
        )
        for evidence in (
            language_control_evidence,
            language_candidate_evidence,
        ):
            self.assertEqual(evidence["configuration"]["region_padding_px"], 2)
            self.assertIsNone(evidence["configuration"]["target_dpi"])
            self.assertEqual(evidence["configuration"]["page_segmentation_mode"], 6)
            self.assertEqual(evidence["configuration"]["engine_mode"], 3)
            self.assertTrue(evidence["crops"])
            self.assertTrue(
                all(crop["padding_pixels"] == 2 for crop in evidence["crops"])
            )
        self.assertEqual(multilingual_smoke["status"], "pass")
        self.assertTrue(multilingual_smoke["required_literals"]["AITEQNO"])
        self.assertTrue(multilingual_smoke["required_literals"]["2026"])
        self.assertTrue(
            all(item["passed"] for item in multilingual_smoke["required_any_groups"])
        )
        self.assertEqual(
            summary["layers"]["ocr_language_profile_comparison"]["decision"],
            "supported",
        )
        self.assertEqual(
            summary["layers"]["ocr_language_profile_comparison"]["selected_profile"],
            "jpn",
        )
        self.assertTrue(
            summary["layers"]["ocr_language_profile_comparison"]["candidate_adopted"]
        )
        grouping_control = json.loads(
            (
                output
                / "ocr-region-grouping"
                / "control"
                / "ocr-quality-evaluation.json"
            ).read_text(encoding="utf-8")
        )
        grouping_candidate = json.loads(
            (
                output
                / "ocr-region-grouping"
                / "candidate"
                / "ocr-quality-evaluation.json"
            ).read_text(encoding="utf-8")
        )
        grouping_plans = json.loads(
            (output / "ocr-region-grouping" / "region-plan-evidence.json").read_text(
                encoding="utf-8"
            )
        )
        singleton_observations = json.loads(
            (output / "ocr-region-grouping" / "singleton-observations.json").read_text(
                encoding="utf-8"
            )
        )
        self.assertTrue(
            all(
                check["status"] == "pass"
                for check in grouping_comparison["checks"].values()
            )
        )
        self.assertEqual(
            grouping_comparison["adoption_policy"]["allowed_geometry_differences"],
            ["region_plan"],
        )
        self.assertFalse(grouping_comparison["recovery"]["anchors"]["lost"])
        self.assertFalse(grouping_comparison["recovery"]["logical_blocks"]["lost"])
        self.assertFalse(grouping_comparison["recovery"]["protected_literals"]["lost"])
        self.assertIn(
            "title",
            grouping_comparison["recovery"]["grouping_targets"]["newly_recovered"],
        )
        self.assertEqual(singleton_observations["status"], "pass")
        self.assertEqual(singleton_observations["changed_region_refs"], [])
        self.assertEqual(grouping_comparison["multilingual_smoke"]["status"], "pass")
        self.assertEqual(
            grouping_comparison["metrics"],
            {
                "essential_anchor_recall": {
                    "candidate": 75.0,
                    "control": 66.666667,
                    "delta_percentage_points": 8.333333,
                },
                "logical_block_coverage": {
                    "candidate": 75.0,
                    "control": 70.833333,
                    "delta_percentage_points": 4.166667,
                },
                "text_character_accuracy": {
                    "candidate": 76.893939,
                    "control": 76.136364,
                    "delta_percentage_points": 0.757575,
                },
            },
        )
        self.assertEqual(grouping_plans["control"]["counts"]["source_regions"], 79)
        self.assertEqual(grouping_plans["control"]["counts"]["planned_regions"], 79)
        self.assertEqual(grouping_plans["control"]["counts"]["groups"], 0)
        self.assertEqual(grouping_plans["candidate"]["counts"]["source_regions"], 79)
        self.assertEqual(grouping_plans["candidate"]["counts"]["planned_regions"], 50)
        self.assertEqual(grouping_plans["candidate"]["counts"]["groups"], 11)
        self.assertEqual(grouping_plans["candidate"]["counts"]["singletons"], 39)
        control_blocks = {
            value["reference_id"]: value for value in grouping_control["blocks"]
        }
        candidate_blocks = {
            value["reference_id"]: value for value in grouping_candidate["blocks"]
        }
        for field_name in ("observed_text", "character_accuracy", "recovered"):
            self.assertEqual(
                control_blocks["phone-label"][field_name],
                candidate_blocks["phone-label"][field_name],
            )
        self.assertFalse(control_blocks["title"]["recovered"])
        self.assertTrue(candidate_blocks["title"]["recovered"])
        self.assertEqual(
            candidate_blocks["title"]["observed_text"],
            "文書解析評価シート",
        )
        self.assertEqual(
            control_blocks["content-structure"]["character_accuracy"], 33.333333
        )
        self.assertEqual(
            candidate_blocks["content-structure"]["character_accuracy"], 26.315789
        )
        self.assertFalse(candidate_blocks["content-structure"]["recovered"])
        self.assertEqual(
            summary["layers"]["ocr_region_grouping_comparison"]["decision"],
            grouping_comparison["decision"],
        )
        self.assertEqual(
            summary["layers"]["ocr_region_grouping_comparison"]["selected_grouping"],
            expected_grouping,
        )
        self.assertEqual(
            summary["layers"]["ocr_region_grouping_comparison"]["candidate_adopted"],
            False,
        )
        self.assertFalse(
            summary["layers"]["ocr_region_grouping_comparison"]["candidate_eligible"]
        )
        self.assertTrue((output / "ocr-quality-control-evaluation.json").is_file())
        self.assertTrue((output / "ocr-quality-evaluation.json").is_file())
        self.assertTrue((output / "ocr-input-transform.json").is_file())
        self.assertTrue((output / "ocr-resolution-comparison.json").is_file())
        self.assertTrue((output / "control-bundle" / "document.ir.json").is_file())
        self.assertTrue((output / "candidate-bundle" / "document.ir.json").is_file())
        self.assertTrue(
            (
                output / "ocr-padding" / "candidate" / "bundle" / "document.ir.json"
            ).is_file()
        )
        self.assertTrue(
            (
                output / "ocr-language" / "control" / "bundle" / "document.ir.json"
            ).is_file()
        )
        self.assertTrue(
            (
                output / "ocr-language" / "candidate" / "bundle" / "document.ir.json"
            ).is_file()
        )
        for side in ("control", "candidate"):
            self.assertTrue(
                (
                    output
                    / "ocr-region-grouping"
                    / side
                    / "bundle"
                    / "document.ir.json"
                ).is_file()
            )
        control_document = DocumentIR.from_json(
            (output / "control-bundle" / "document.ir.json").read_bytes()
        )
        selected_ir_bytes = (output / "bundle" / "document.ir.json").read_bytes()
        self.assertEqual(
            hashlib.sha256(selected_ir_bytes).hexdigest(),
            "494ebee340dfa3cbd3bd20aacf363f97d8a7e972a8627f1247d2f5b12980a32c",
        )
        selected_document = DocumentIR.from_json(selected_ir_bytes)
        padding_document = DocumentIR.from_json(
            (
                output / "ocr-padding" / "candidate" / "bundle" / "document.ir.json"
            ).read_bytes()
        )
        language_candidate_document = DocumentIR.from_json(
            (
                output / "ocr-language" / "candidate" / "bundle" / "document.ir.json"
            ).read_bytes()
        )
        grouping_control_document = DocumentIR.from_json(
            (
                output
                / "ocr-region-grouping"
                / "control"
                / "bundle"
                / "document.ir.json"
            ).read_bytes()
        )
        grouping_candidate_document = DocumentIR.from_json(
            (
                output
                / "ocr-region-grouping"
                / "candidate"
                / "bundle"
                / "document.ir.json"
            ).read_bytes()
        )
        self.assertNotIn(
            TABLE_TOPOLOGY_EXTENSION_KEY,
            control_document.pages[0].extensions,
        )
        topology = read_page_table_topology(selected_document.pages[0])
        self.assertIsNotNone(topology)
        assert topology is not None
        raw_type_counts = {
            element_type: sum(
                element.type.value == element_type
                for element in selected_document.pages[0].elements
            )
            for element_type in ("text", "line", "rectangle", "image")
        }
        self.assertEqual(raw_type_counts["text"], 374)
        self.assertEqual(raw_type_counts["line"], 50)
        self.assertEqual(raw_type_counts["rectangle"], 51)
        self.assertEqual(raw_type_counts["image"], 0)
        self.assertEqual(len(topology.tables), 5)
        self.assertEqual(
            tuple(
                (table.logical_rows, table.logical_columns, len(table.cells))
                for table in topology.tables
            ),
            ((4, 4, 14), (7, 2, 14), (4, 2, 8), (2, 3, 6), (3, 1, 3)),
        )
        self.assertEqual(sum(len(table.cells) for table in topology.tables), 45)
        assigned_text_ids = tuple(
            text_id
            for table in topology.tables
            for cell in table.cells
            for text_id in cell.text_element_ids
        )
        self.assertTrue(assigned_text_ids)
        self.assertEqual(len(assigned_text_ids), len(set(assigned_text_ids)))
        self.assertEqual(
            tuple(
                (cell.row_index, cell.column_index, cell.rowspan, cell.colspan)
                for cell in topology.tables[0].cells
                if cell.rowspan > 1 or cell.colspan > 1
            ),
            ((2, 2, 1, 2), (3, 2, 1, 2)),
        )
        self.assertEqual(topology.diagnostics.ambiguous_text_element_ids, ())
        self.assertEqual(topology.diagnostics.ambiguous_primitive_element_ids, ())
        self.assertEqual(topology.diagnostics.unassigned_primitive_element_ids, ())
        self.assertEqual(
            len(assigned_text_ids)
            + len(topology.diagnostics.unassigned_text_element_ids),
            raw_type_counts["text"],
        )
        role_counts = {
            role: sum(
                assignment.role is role for assignment in topology.primitive_roles
            )
            for role in TablePrimitiveRole
        }
        self.assertEqual(role_counts[TablePrimitiveRole.PAGE_FRAME], 5)
        self.assertEqual(role_counts[TablePrimitiveRole.PAGE_DECORATION], 4)
        self.assertEqual(role_counts[TablePrimitiveRole.TABLE_OUTER_BORDER], 5)
        self.assertEqual(role_counts[TablePrimitiveRole.CELL_RECTANGLE], 45)
        self.assertEqual(
            role_counts[TablePrimitiveRole.DUPLICATED_SUPPORTING_PRIMITIVE],
            20,
        )
        self.assertEqual(role_counts[TablePrimitiveRole.ROW_BOUNDARY], 15)
        self.assertEqual(role_counts[TablePrimitiveRole.COLUMN_BOUNDARY], 7)

        padding_data = padding_document.to_dict()
        language_candidate_data = language_candidate_document.to_dict()
        grouping_control_data = grouping_control_document.to_dict()
        grouping_candidate_data = grouping_candidate_document.to_dict()
        selected_data = selected_document.to_dict()
        selected_page_extensions = selected_data["pages"][0]["extensions"]
        del selected_page_extensions[TABLE_TOPOLOGY_EXTENSION_KEY]
        if not selected_page_extensions:
            del selected_data["pages"][0]["extensions"]
        self.assertEqual(language_candidate_data, grouping_control_data)
        self.assertEqual(
            selected_data,
            grouping_candidate_data if grouping_supported else grouping_control_data,
        )
        self.assertNotEqual(selected_data, padding_data)
        self.assertNotEqual(
            (output / "bundle" / "document.ir.json").read_bytes(),
            (output / "control-bundle" / "document.ir.json").read_bytes(),
        )
        self.assertNotEqual(
            (output / "bundle" / "document.ir.json").read_bytes(),
            (output / "candidate-bundle" / "document.ir.json").read_bytes(),
        )
        self.assertNotEqual(
            (output / "bundle" / "document.ir.json").read_bytes(),
            (
                output / "ocr-padding" / "candidate" / "bundle" / "document.ir.json"
            ).read_bytes(),
        )
        self.assertNotEqual(
            (output / "bundle" / "document.ir.json").read_bytes(),
            (
                output / "ocr-language" / "candidate" / "bundle" / "document.ir.json"
            ).read_bytes(),
        )
        preview_report = json.loads(
            (output / "preview-render-report.json").read_text(encoding="utf-8")
        )
        selected_element_count = len(selected_document.pages[0].elements)
        self.assertEqual(
            len(preview_report["rendered_element_ids"]),
            selected_element_count,
        )
        self.assertEqual(preview_report["omitted_element_ids"], [])
        self.assertEqual(preview_report["fallback_element_ids"], [])
        self.assertEqual(preview_report["warnings"], [])
        docx_report = json.loads(
            (output / "docx-render-report.json").read_text(encoding="utf-8")
        )
        self.assertEqual(
            docx_report["native_table_ids"],
            [table.id for table in topology.tables],
        )
        self.assertTrue(
            set(assigned_text_ids).issubset(
                docx_report["native_table_consumed_element_ids"]
            )
        )
        self.assertEqual(
            len(docx_report["rendered_element_ids"]),
            selected_element_count,
        )
        self.assertEqual(docx_report["font_substitutions"], [])
        self.assertNotIn(
            "font_substituted",
            {warning["code"] for warning in docx_report["warnings"]},
        )
        consumed = set(docx_report["native_table_consumed_element_ids"])
        self.assertFalse(
            any(
                warning["code"]
                in {"vertical_position_approximated", "z_order_approximated"}
                and warning["element_id"] in consumed
                for warning in docx_report["warnings"]
            )
        )
        reconstructed = open_docx(output / "bundle" / "reconstructed.docx")
        self.assertEqual(len(reconstructed.tables), 5)
        source_contents = reconstructed._element.body.xpath(".//w:sdtContent")
        self.assertEqual(len(source_contents), raw_type_counts["text"])
        for content in source_contents:
            run = content.find(qn("w:r"))
            self.assertIsNotNone(run)
            run_properties = run.find(qn("w:rPr"))
            self.assertIsNotNone(run_properties)
            run_fonts = run_properties.find(qn("w:rFonts"))
            self.assertIsNotNone(run_fonts)
            for channel in ("ascii", "hAnsi", "eastAsia", "cs"):
                self.assertEqual(
                    run_fonts.get(qn(f"w:{channel}")),
                    "Noto Sans CJK JP",
                )
        self.assertEqual(
            [
                table._tbl.tblPr.find(qn("w:tblCaption")).get(qn("w:val"))
                for table in reconstructed.tables
            ],
            [f"aiteqno-table:{table.id}" for table in topology.tables],
        )
        self.assertEqual(
            [
                node.get(qn("w:val"))
                for node in reconstructed.tables[0]._tbl.xpath(".//w:gridSpan")
            ],
            ["2", "2"],
        )
        observation = json.loads(
            (output / "docx-observation.json").read_text(encoding="utf-8")
        )
        self.assertEqual(observation["errors"], [])
        self.assertEqual(
            sum(
                element["source_element_id"] is not None
                for element in observation["elements"]
                if element["type"] == "text"
            ),
            raw_type_counts["text"],
        )
        self.assertTrue((output / "source-quality-evaluation.json").is_file())
        self.assertTrue((output / "ir-to-docx-restoration-evaluation.json").is_file())
        source_components = source_evaluation["components"]
        self.assertGreater(source_evaluation["overall_score"], 45.22)
        self.assertGreater(
            source_components["text_accuracy"]["score"],
            21.212121,
        )
        self.assertGreaterEqual(
            source_components["logical_block_coverage"]["score"],
            70.833333,
        )
        self.assertGreaterEqual(
            source_components["structure_similarity"]["score"],
            46.969697,
        )
        self.assertGreaterEqual(
            source_components["geometry_similarity"]["score"],
            80.754604,
        )
        address_block = next(
            block
            for block in source_evaluation["logical_blocks"]
            if block["reference_id"] == "address-label"
        )
        self.assertTrue(address_block["essential"])
        self.assertTrue(address_block["covered"])
        self.assertEqual(address_block["observed_text"], "住所")
        restoration = json.loads(
            (output / "ir-to-docx-restoration-evaluation.json").read_text(
                encoding="utf-8"
            )
        )
        restoration_components = {
            name: component["score"]
            for name, component in restoration["components"].items()
        }
        self.assertEqual(restoration["state"], "pass")
        self.assertGreaterEqual(restoration["overall_score"], 78.73)
        self.assertEqual(restoration_components["text_similarity"], 100.0)
        self.assertGreaterEqual(
            restoration_components["element_coverage"],
            94.209354,
        )
        self.assertGreaterEqual(
            restoration_components["structure_similarity"],
            74.434957,
        )
        self.assertEqual(restoration_components["geometry_similarity"], 0.0)
        self.assertEqual(
            summary["layers"]["candidate_ir_to_docx"]["state"],
            "pass",
        )
        snapshot = output / "actual-docx-snapshot"
        self.assertTrue((snapshot / "snapshot.pdf").is_file())
        self.assertTrue((snapshot / "page-001.png").is_file())
        self.assertTrue((snapshot / "snapshot-evidence.json").is_file())
        self.assertTrue((snapshot / "visible-ocr.json").is_file())
        visible_ocr = json.loads(
            (snapshot / "visible-ocr.json").read_text(encoding="utf-8")
        )
        self.assertEqual(visible_ocr["languages"], ["jpn", "eng"])
        snapshot_evidence = json.loads(
            (snapshot / "snapshot-evidence.json").read_text(encoding="utf-8")
        )
        self.assertEqual(snapshot_evidence["page_count"], 1)
        pdffonts = shutil.which("pdffonts")
        self.assertIsNotNone(pdffonts)
        font_inventory = subprocess.run(
            [pdffonts, str(snapshot / "snapshot.pdf")],
            cwd=REPOSITORY_ROOT,
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=30,
        )
        self.assertEqual(font_inventory.returncode, 0, msg=font_inventory.stderr)
        self.assertIn("NotoSansCJKjp-Regular", font_inventory.stdout)
        self.assertNotIn("LiberationSans", font_inventory.stdout)


if __name__ == "__main__":
    unittest.main()
