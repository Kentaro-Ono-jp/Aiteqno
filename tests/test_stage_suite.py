import hashlib
import json
import shutil
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from scripts.run_stage_suite import (
    DEFAULT_SUITE_PATH,
    _docx_media_evidence,
    _hidden_text_check,
    _integrity_report,
    _load_suite,
    _read_fixture,
    run,
)


Q01_DIRECTORY = (
    Path(__file__).parent
    / "fixtures"
    / "generalization"
    / "japanese-questionnaires-v1"
)
Q01_MANIFEST = Q01_DIRECTORY / "questionnaire-01-general-medicine.manifest.json"
Q01_ID = "questionnaire-01-general-medicine"
Q01_SOURCE_SHA256 = "e6aadded4a7ca5d92358c93d87679b65f5f81f9ebf886a13871968a1dd96a734"
Q01_REFERENCE_SHA256 = "b3c5670dce98dedfced0d1508ba95583801ec430e69bb6391a20945a10fa82cb"
Q02_MANIFEST = Q01_DIRECTORY / "questionnaire-02-fever-respiratory.manifest.json"
Q02_ID = "questionnaire-02-fever-respiratory"
Q02_SOURCE_SHA256 = "6c27901390f4a1b43729d681884aae144886d2725ac3d585d9570f4499137ba2"
Q02_REFERENCE_SHA256 = "94b073eb3d7cb6df5001054d61f212c99b476c7af0fc08df55385dfce1d12b0a"
Q03_MANIFEST = Q01_DIRECTORY / "questionnaire-03-gastroenterology.manifest.json"
Q03_ID = "questionnaire-03-gastroenterology"
Q03_SOURCE_SHA256 = "825bddca8853986288cb5762bb26e80143762120ffe5859793f4ec5a171f83a7"
Q03_REFERENCE_SHA256 = "d358c1b7f92ef52ca1b59b3667f3d7922e81538fdd3429ac02e964a229baf1df"
Q04_MANIFEST = Q01_DIRECTORY / "questionnaire-04-orthopedics.manifest.json"
Q04_ID = "questionnaire-04-orthopedics"
Q04_SOURCE_SHA256 = "0e322bc9b5e8593d5a0fda959bd314cb6dc2c46de79fc3c07467527dba6dc4cd"
Q04_REFERENCE_SHA256 = "21d5279d9da146f4170f3adc197fd0d0927d6a50c920e7896d4eb6e4d9ce09c2"
Q05_MANIFEST = Q01_DIRECTORY / "questionnaire-05-dermatology.manifest.json"
Q05_ID = "questionnaire-05-dermatology"
Q05_SOURCE_SHA256 = "7db96f0118193e93685f4a80472bb64f4656bebe77ccc55d34fbfe447d158e01"
Q05_REFERENCE_SHA256 = "00813b3ee0d1a35ba4a572821cff6fa123ff281bb189e52e9370949b6d1bd852"
Q06_MANIFEST = Q01_DIRECTORY / "questionnaire-06-otorhinolaryngology.manifest.json"
Q06_ID = "questionnaire-06-otorhinolaryngology"
Q06_SOURCE_SHA256 = "2825b68275a631574eb9616e136e45bc0ee2f1739195eab4ca16c4d572e9f4b1"
Q06_REFERENCE_SHA256 = "a826e86eff13172f823a0f70d6f1599165aab717a341b90bc3b73c0eb3562afa"
Q07_MANIFEST = Q01_DIRECTORY / "questionnaire-07-ophthalmology.manifest.json"
Q07_ID = "questionnaire-07-ophthalmology"
Q07_SOURCE_SHA256 = "670ec15d9df677482b7122af306e7f6cedc0d3f2befea50a02b1907dcc0d6b66"
Q07_REFERENCE_SHA256 = "6efc969985c532cb6f2ddfbb9a4f29724ae18c7d2e1fdc9e12433ff051dd8f15"
Q08_MANIFEST = Q01_DIRECTORY / "questionnaire-08-pediatrics.manifest.json"
Q08_ID = "questionnaire-08-pediatrics"
Q08_SOURCE_SHA256 = "9c5679915a4e479a6fc806918832df8f0db23f325a228bec9d16fd76ba948bb9"
Q08_REFERENCE_SHA256 = "f2109a1cd0c8be2ff5556e3f9431a3efa4bb52bf4f5e0fc663de5307e6ec2095"
Q09_MANIFEST = Q01_DIRECTORY / "questionnaire-09-gynecology.manifest.json"
Q09_ID = "questionnaire-09-gynecology"
Q09_SOURCE_SHA256 = "b68368e751d971cd9869f4a47009c98d57e569bc6881f43a6853761931fa7089"
Q09_REFERENCE_SHA256 = "e2101d144bfadfb7a77d9fb229b31d41a137f426e1afb72869d41a0935217462"
Q10_MANIFEST = Q01_DIRECTORY / "questionnaire-10-health-lifestyle.manifest.json"
Q10_ID = "questionnaire-10-health-lifestyle"
Q10_SOURCE_SHA256 = "502ce57a54c3fc9982994da1aa64c2c810f3aca1e16d8a13f379b79cbb96fc7d"
Q10_REFERENCE_SHA256 = "18ac7dbc981cd5bdb52d89f6bcd23c62bddd5fa62ff04de3e36bee2fc1d68354"
STAGE_TWO_SUITE_PATH = (
    Path(__file__).parent / "fixtures" / "stages" / "questionnaire-stage-2.json"
)
STAGE_THREE_SUITE_PATH = (
    Path(__file__).parent / "fixtures" / "stages" / "questionnaire-stage-3.json"
)
STAGE_FOUR_SUITE_PATH = (
    Path(__file__).parent / "fixtures" / "stages" / "questionnaire-stage-4.json"
)
STAGE_FIVE_SUITE_PATH = (
    Path(__file__).parent / "fixtures" / "stages" / "questionnaire-stage-5.json"
)
STAGE_SIX_SUITE_PATH = (
    Path(__file__).parent / "fixtures" / "stages" / "questionnaire-stage-6.json"
)
STAGE_SEVEN_SUITE_PATH = (
    Path(__file__).parent / "fixtures" / "stages" / "questionnaire-stage-7.json"
)
STAGE_EIGHT_SUITE_PATH = (
    Path(__file__).parent / "fixtures" / "stages" / "questionnaire-stage-8.json"
)
STAGE_NINE_SUITE_PATH = (
    Path(__file__).parent / "fixtures" / "stages" / "questionnaire-stage-9.json"
)
STAGE_TEN_SUITE_PATH = (
    Path(__file__).parent / "fixtures" / "stages" / "questionnaire-stage-10.json"
)


class StageSuiteContractTest(unittest.TestCase):
    def test_stage_one_has_exactly_the_two_pinned_active_fixtures(self):
        suite = _load_suite(DEFAULT_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            ["synthetic-dense-japanese-form-v1", Q01_ID],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            ["base64", "raw"],
        )
        self.assertEqual([item.source_dpi for item in suite.fixtures], [96, 150])
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_stage_two_adds_only_q02_in_the_pinned_order(self):
        suite = _load_suite(STAGE_TWO_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            ["synthetic-dense-japanese-form-v1", Q01_ID, Q02_ID],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            ["base64", "raw", "raw"],
        )
        self.assertEqual([item.source_dpi for item in suite.fixtures], [96, 150, 200])
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_stage_three_adds_only_q03_in_the_pinned_order(self):
        suite = _load_suite(STAGE_THREE_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            ["synthetic-dense-japanese-form-v1", Q01_ID, Q02_ID, Q03_ID],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            ["base64", "raw", "raw", "raw"],
        )
        self.assertEqual(
            [item.source_dpi for item in suite.fixtures],
            [96, 150, 200, 150],
        )
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_stage_four_adds_only_q04_in_the_pinned_order(self):
        suite = _load_suite(STAGE_FOUR_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            [
                "synthetic-dense-japanese-form-v1",
                Q01_ID,
                Q02_ID,
                Q03_ID,
                Q04_ID,
            ],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            ["base64", "raw", "raw", "raw", "raw"],
        )
        self.assertEqual(
            [item.source_dpi for item in suite.fixtures],
            [96, 150, 200, 150, 150],
        )
        self.assertEqual(
            [(item.source_width, item.source_height) for item in suite.fixtures],
            [
                (700, 991),
                (1240, 1754),
                (1654, 2339),
                (1240, 1754),
                (1754, 1240),
            ],
        )
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_stage_five_adds_only_q05_in_the_pinned_order(self):
        suite = _load_suite(STAGE_FIVE_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            [
                "synthetic-dense-japanese-form-v1",
                Q01_ID,
                Q02_ID,
                Q03_ID,
                Q04_ID,
                Q05_ID,
            ],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            ["base64", "raw", "raw", "raw", "raw", "raw"],
        )
        self.assertEqual(
            [item.source_dpi for item in suite.fixtures],
            [96, 150, 200, 150, 150, 200],
        )
        self.assertEqual(
            [(item.source_width, item.source_height) for item in suite.fixtures],
            [
                (700, 991),
                (1240, 1754),
                (1654, 2339),
                (1240, 1754),
                (1754, 1240),
                (1654, 2339),
            ],
        )
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_stage_six_adds_only_q06_in_the_pinned_order(self):
        suite = _load_suite(STAGE_SIX_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            [
                "synthetic-dense-japanese-form-v1",
                Q01_ID,
                Q02_ID,
                Q03_ID,
                Q04_ID,
                Q05_ID,
                Q06_ID,
            ],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            ["base64", "raw", "raw", "raw", "raw", "raw", "raw"],
        )
        self.assertEqual(
            [item.source_dpi for item in suite.fixtures],
            [96, 150, 200, 150, 150, 200, 150],
        )
        self.assertEqual(
            [(item.source_width, item.source_height) for item in suite.fixtures],
            [
                (700, 991),
                (1240, 1754),
                (1654, 2339),
                (1240, 1754),
                (1754, 1240),
                (1654, 2339),
                (1240, 1754),
            ],
        )
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_stage_seven_adds_only_q07_in_the_pinned_order(self):
        suite = _load_suite(STAGE_SEVEN_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            [
                "synthetic-dense-japanese-form-v1",
                Q01_ID,
                Q02_ID,
                Q03_ID,
                Q04_ID,
                Q05_ID,
                Q06_ID,
                Q07_ID,
            ],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            ["base64", "raw", "raw", "raw", "raw", "raw", "raw", "raw"],
        )
        self.assertEqual(
            [item.source_dpi for item in suite.fixtures],
            [96, 150, 200, 150, 150, 200, 150, 200],
        )
        self.assertEqual(
            [(item.source_width, item.source_height) for item in suite.fixtures],
            [
                (700, 991),
                (1240, 1754),
                (1654, 2339),
                (1240, 1754),
                (1754, 1240),
                (1654, 2339),
                (1240, 1754),
                (2339, 1654),
            ],
        )
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_stage_eight_adds_only_q08_in_the_pinned_order(self):
        suite = _load_suite(STAGE_EIGHT_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            [
                "synthetic-dense-japanese-form-v1",
                Q01_ID,
                Q02_ID,
                Q03_ID,
                Q04_ID,
                Q05_ID,
                Q06_ID,
                Q07_ID,
                Q08_ID,
            ],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            ["base64", "raw", "raw", "raw", "raw", "raw", "raw", "raw", "raw"],
        )
        self.assertEqual(
            [item.source_dpi for item in suite.fixtures],
            [96, 150, 200, 150, 150, 200, 150, 200, 150],
        )
        self.assertEqual(
            [(item.source_width, item.source_height) for item in suite.fixtures],
            [
                (700, 991),
                (1240, 1754),
                (1654, 2339),
                (1240, 1754),
                (1754, 1240),
                (1654, 2339),
                (1240, 1754),
                (2339, 1654),
                (1240, 1754),
            ],
        )
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_stage_nine_adds_only_q09_in_the_pinned_order(self):
        suite = _load_suite(STAGE_NINE_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            [
                "synthetic-dense-japanese-form-v1",
                Q01_ID,
                Q02_ID,
                Q03_ID,
                Q04_ID,
                Q05_ID,
                Q06_ID,
                Q07_ID,
                Q08_ID,
                Q09_ID,
            ],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            [
                "base64",
                "raw",
                "raw",
                "raw",
                "raw",
                "raw",
                "raw",
                "raw",
                "raw",
                "raw",
            ],
        )
        self.assertEqual(
            [item.source_dpi for item in suite.fixtures],
            [96, 150, 200, 150, 150, 200, 150, 200, 150, 200],
        )
        self.assertEqual(
            [(item.source_width, item.source_height) for item in suite.fixtures],
            [
                (700, 991),
                (1240, 1754),
                (1654, 2339),
                (1240, 1754),
                (1754, 1240),
                (1654, 2339),
                (1240, 1754),
                (2339, 1654),
                (1240, 1754),
                (1654, 2339),
            ],
        )
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_stage_ten_adds_only_q10_in_the_pinned_order(self):
        suite = _load_suite(STAGE_TEN_SUITE_PATH)

        self.assertEqual(
            [item.fixture_id for item in suite.fixtures],
            [
                "synthetic-dense-japanese-form-v1",
                Q01_ID,
                Q02_ID,
                Q03_ID,
                Q04_ID,
                Q05_ID,
                Q06_ID,
                Q07_ID,
                Q08_ID,
                Q09_ID,
                Q10_ID,
            ],
        )
        self.assertEqual(
            [item.source_encoding for item in suite.fixtures],
            ["base64", *(["raw"] * 10)],
        )
        self.assertEqual(
            [item.source_dpi for item in suite.fixtures],
            [96, 150, 200, 150, 150, 200, 150, 200, 150, 200, 300],
        )
        self.assertEqual(
            [(item.source_width, item.source_height) for item in suite.fixtures],
            [
                (700, 991),
                (1240, 1754),
                (1654, 2339),
                (1240, 1754),
                (1754, 1240),
                (1654, 2339),
                (1240, 1754),
                (2339, 1654),
                (1240, 1754),
                (1654, 2339),
                (2480, 3508),
            ],
        )
        self.assertEqual(suite.threshold, 70)
        self.assertEqual(suite.production_languages, ("jpn",))
        self.assertEqual(suite.visible_languages, suite.production_languages)
        self.assertEqual(suite.snapshot_dpi, 300)

    def test_q01_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q01_MANIFEST, Q01_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q01_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q01_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (1240, 1754))
        self.assertTrue(reference.reviewed)
        self.assertGreaterEqual(len(reference.text_regions), 35)
        self.assertGreaterEqual(len(reference.structural_items), 35)
        self.assertGreaterEqual(len(reference.relationships), 14)
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "内科 初診問診票",
            "本日はどのような症状で受診されましたか。",
            "現在飲んでいる薬やサプリメントをご記入ください。",
            "医師へ伝えておきたいこと",
            "ご記入後、受付へお渡しください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])

    def test_q02_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q02_MANIFEST, Q02_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q02_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q02_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (1654, 2339))
        self.assertEqual(fixture.source_dpi, 200)
        self.assertTrue(reference.reviewed)
        self.assertGreaterEqual(len(reference.text_regions), 38)
        self.assertGreaterEqual(len(reference.structural_items), 51)
        self.assertGreaterEqual(len(reference.relationships), 21)
        self.assertGreaterEqual(
            sum("checkbox" in item.id for item in reference.structural_items),
            25,
        )
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "発熱・呼吸器症状 問診票",
            "発熱外来・呼吸器内科",
            "発熱に気づいた日時をご記入ください。",
            "これまでの最高体温は何度でしたか。",
            "せき、たん、のどの痛み、鼻水、息苦しさはありますか。",
            "最近、感染症と診断された方との接触はありましたか。",
            "補足（発症の順序、使用した解熱薬など）",
            "受付の案内に従い、待機場所でお待ちください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])
        self.assertTrue(fixture.manifest["redistribution_allowed"])
        self.assertFalse(fixture.manifest["contains_personal_data"])

    def test_q03_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q03_MANIFEST, Q03_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q03_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q03_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (1240, 1754))
        self.assertEqual(fixture.source_dpi, 150)
        self.assertTrue(reference.reviewed)
        self.assertGreaterEqual(len(reference.text_regions), 35)
        self.assertGreaterEqual(len(reference.structural_items), 66)
        self.assertGreaterEqual(len(reference.relationships), 28)
        self.assertGreaterEqual(
            sum("checkbox" in item.id for item in reference.structural_items),
            30,
        )
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "消化器内科 問診票",
            "腹部症状と便の状態",
            "痛む場所と、痛みが始まった時期をご記入ください。",
            "食事をすると症状は強くなりますか、軽くなりますか。",
            "便の回数や硬さに変化はありますか。",
            "便に血が混じったことや、黒い便が出たことはありますか。",
            "伴う症状",
            "吐き気、嘔吐、胸やけ、食欲低下はありますか。",
            "既往歴・服薬・アレルギー・手術歴",
            "症状の経過や気になる食べ物",
            "ご記入後、受付へお渡しください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])
        self.assertTrue(fixture.manifest["redistribution_allowed"])
        self.assertFalse(fixture.manifest["contains_personal_data"])

    def test_q04_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q04_MANIFEST, Q04_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q04_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q04_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (1754, 1240))
        self.assertEqual(fixture.source_dpi, 150)
        self.assertTrue(reference.reviewed)
        self.assertGreaterEqual(len(reference.text_regions), 33)
        self.assertGreaterEqual(len(reference.structural_items), 64)
        self.assertGreaterEqual(len(reference.relationships), 42)
        self.assertGreaterEqual(
            sum("checkbox" in item.id for item in reference.structural_items),
            10,
        )
        self.assertEqual(
            sum("pain-scale-tick" in item.id for item in reference.structural_items),
            11,
        )
        self.assertGreaterEqual(
            sum(
                item.id.startswith(("front-", "back-"))
                for item in reference.structural_items
            ),
            16,
        )
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "整形外科 初診問診票",
            "症状の場所と経過",
            "痛みやしびれのある場所を図に記してください。",
            "前面",
            "背面",
            "現在の痛みを0から10で表すと、どの程度ですか。",
            "歩行、階段、着替えなどで困っている動作はありますか。",
            "痛みの尺度",
            "0　1　2　3　4　5　6　7　8　9　10",
            "今回の症状について、ほかの医療機関で治療を受けましたか。",
            "医療機関・治療内容",
            "ご記入後、受付へお渡しください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])
        self.assertEqual(
            reference_json["source_dimensions"]["orientation"],
            "landscape",
        )
        self.assertTrue(fixture.manifest["redistribution_allowed"])
        self.assertFalse(fixture.manifest["contains_personal_data"])

    def test_q05_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q05_MANIFEST, Q05_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q05_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q05_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (1654, 2339))
        self.assertEqual(fixture.source_dpi, 200)
        self.assertTrue(reference.reviewed)
        self.assertGreaterEqual(len(reference.text_regions), 36)
        self.assertGreaterEqual(len(reference.structural_items), 65)
        self.assertGreaterEqual(len(reference.relationships), 60)
        self.assertGreaterEqual(
            sum("checkbox" in item.id for item in reference.structural_items),
            18,
        )
        self.assertGreaterEqual(
            sum(
                item.id.startswith(("front-", "back-"))
                for item in reference.structural_items
            ),
            16,
        )
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "皮膚科 問診票",
            "症状について",
            "いつ頃から症状がありますか。",
            "同じ症状を繰り返したことがありますか。",
            "かゆみ、痛み、出血、じゅくじゅくした感じはありますか。",
            "症状が出る前に、薬、化粧品、洗剤などを変えましたか。",
            "症状の場所",
            "症状が出ている場所を図に記してください。",
            "前面",
            "背面",
            "これまでの治療・アレルギー",
            "現在使用中の塗り薬・飲み薬",
            "薬・食べ物などのアレルギー",
            "診察時に伝えておきたいこと",
            "ご記入後、受付へお渡しください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])
        self.assertEqual(
            reference_json["source_dimensions"]["orientation"],
            "portrait",
        )
        self.assertTrue(fixture.manifest["redistribution_allowed"])
        self.assertFalse(fixture.manifest["contains_personal_data"])

    def test_q06_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q06_MANIFEST, Q06_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q06_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q06_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (1240, 1754))
        self.assertEqual(fixture.source_dpi, 150)
        self.assertTrue(reference.reviewed)
        self.assertGreaterEqual(len(reference.text_regions), 38)
        self.assertGreaterEqual(len(reference.structural_items), 58)
        self.assertGreaterEqual(len(reference.relationships), 67)
        self.assertEqual(
            sum("checkbox" in item.id for item in reference.structural_items),
            28,
        )
        structure_ids = {item.id for item in reference.structural_items}
        self.assertEqual(
            {
                item_id
                for item_id in structure_ids
                if item_id.endswith("grid-outer")
            },
            {
                "identity-grid-outer",
                "ear-grid-outer",
                "nose-grid-outer",
                "throat-grid-outer",
            },
        )
        self.assertIn("notes-box", structure_ids)
        self.assertIn("notes-writing-line", structure_ids)
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "耳鼻咽喉科 問診票",
            "症状がある側をお選びください。 右・左・両方",
            "症状",
            "右耳",
            "左耳",
            "両方・不明",
            "耳の痛み、耳鳴り、聞こえにくさ、耳だれはありますか。",
            "鼻の症状",
            "鼻づまり、鼻水、くしゃみ、においの分かりにくさはありますか。",
            "鼻水の状態",
            "のど・声の症状",
            "のどの痛み、声のかすれ、飲み込みにくさはありますか。",
            "めまいやふらつきを感じることはありますか。",
            "治療中の病気・薬",
            "症状の経過や気になること",
            "ご記入後、受付へお渡しください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])
        self.assertEqual(
            reference_json["source_dimensions"]["orientation"],
            "portrait",
        )
        self.assertTrue(fixture.manifest["redistribution_allowed"])
        self.assertFalse(fixture.manifest["contains_personal_data"])

    def test_q07_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q07_MANIFEST, Q07_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q07_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q07_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (2339, 1654))
        self.assertEqual(fixture.source_dpi, 200)
        self.assertTrue(reference.reviewed)
        self.assertGreaterEqual(len(reference.text_regions), 39)
        self.assertGreaterEqual(len(reference.structural_items), 55)
        self.assertGreaterEqual(len(reference.relationships), 68)
        self.assertEqual(
            sum("checkbox" in item.id for item in reference.structural_items),
            24,
        )
        structure_ids = {item.id for item in reference.structural_items}
        self.assertEqual(
            {
                item_id
                for item_id in structure_ids
                if item_id.endswith("grid-outer")
            },
            {
                "identity-grid-outer",
                "symptom-grid-outer",
                "lens-grid-outer",
                "history-grid-outer",
            },
        )
        self.assertIn("surgery-detail-line", structure_ids)
        self.assertIn("treatment-detail-line", structure_ids)
        self.assertIn("footer-rule", structure_ids)
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "眼科 初診問診票",
            "現在の眼の症状",
            "症状がある眼をお選びください。 右眼・左眼・両眼",
            "右眼",
            "左眼",
            "両眼・不明",
            "見えにくい・かすむ",
            "まぶしい・二重に見える",
            "痛み・かゆみ",
            "充血・目やに",
            "眼鏡・コンタクトレンズ",
            "眼鏡またはコンタクトレンズを使用していますか。",
            "コンタクトの種類",
            "最終交換・作成時期",
            "装用時間",
            "眼科の治療歴・全身状態",
            "これまでに眼の手術やレーザー治療を受けたことがありますか。",
            "治療中の病気・使用中の薬",
            "ご記入後、受付へお渡しください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])
        self.assertEqual(
            reference_json["source_dimensions"]["orientation"],
            "landscape",
        )
        self.assertTrue(fixture.manifest["redistribution_allowed"])
        self.assertFalse(fixture.manifest["contains_personal_data"])

    def test_q08_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q08_MANIFEST, Q08_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q08_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q08_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (1240, 1754))
        self.assertEqual(fixture.source_dpi, 150)
        self.assertTrue(reference.reviewed)
        self.assertEqual(len(reference.text_regions), 46)
        self.assertEqual(len(reference.structural_items), 52)
        self.assertEqual(len(reference.relationships), 105)
        self.assertEqual(
            sum("checkbox" in item.id for item in reference.structural_items),
            19,
        )
        structure_ids = {item.id for item in reference.structural_items}
        self.assertEqual(
            {
                item_id
                for item_id in structure_ids
                if item_id.endswith("grid-outer")
            },
            {
                "identity-grid-outer",
                "current-symptoms-grid-outer",
                "birth-grid-outer",
                "vaccination-grid-outer",
                "history-grid-outer",
            },
        )
        self.assertIn("guardian-message-outer", structure_ids)
        self.assertNotIn("guardian-message-grid-outer", structure_ids)
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "小児科 初診問診票",
            "お子さまの氏名",
            "現在の症状",
            "お子さまの症状と、いつから始まったかをご記入ください。",
            "食事や水分は普段どおり取れていますか。",
            "尿や便の回数に変化はありますか。",
            "出生時の情報",
            "出生時の週数と体重をご記入ください。",
            "予防接種・母子健康手帳",
            "母子健康手帳をお持ちですか。",
            "病歴・アレルギー・薬",
            "保護者の方から医師へ伝えておきたいこと",
            "ご記入後、受付へお渡しください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])
        identity_reading_order = [
            (item["source"], item["target"])
            for item in reference_json["relationships"]
            if item["kind"] == "reading_order"
        ][:16]
        self.assertEqual(
            identity_reading_order,
            [
                ("title", "child-name-label"),
                ("child-name-label", "furigana-label"),
                ("furigana-label", "birth-date-label"),
                ("birth-date-label", "age-label"),
                ("age-label", "gender-label"),
                ("gender-label", "guardian-name-label"),
                ("guardian-name-label", "relationship-label"),
                ("relationship-label", "entry-date-label"),
                ("entry-date-label", "contact-label"),
                ("contact-label", "height-label"),
                ("height-label", "height-unit"),
                ("height-unit", "weight-label"),
                ("weight-label", "weight-unit"),
                ("weight-unit", "temperature-label"),
                ("temperature-label", "temperature-unit"),
                ("temperature-unit", "current-symptoms-section-title"),
            ],
        )
        self.assertEqual(
            reference_json["source_dimensions"]["orientation"],
            "portrait",
        )
        self.assertTrue(fixture.manifest["redistribution_allowed"])
        self.assertFalse(fixture.manifest["contains_personal_data"])

    def test_q09_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q09_MANIFEST, Q09_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q09_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q09_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (1654, 2339))
        self.assertEqual(fixture.source_dpi, 200)
        self.assertTrue(reference.reviewed)
        self.assertEqual(len(reference.text_regions), 32)
        self.assertEqual(len(reference.structural_items), 46)
        self.assertEqual(len(reference.relationships), 68)
        self.assertEqual(
            sum("checkbox" in item.id for item in reference.structural_items),
            15,
        )
        structure_ids = {item.id for item in reference.structural_items}
        self.assertEqual(
            {
                item_id
                for item_id in structure_ids
                if item_id.endswith("grid-outer")
            },
            {
                "identity-grid-outer",
                "menstruation-grid-outer",
                "visit-history-grid-outer",
            },
        )
        self.assertIn("doctor-message-outer", structure_ids)
        self.assertNotIn("doctor-message-grid-outer", structure_ids)
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "婦人科 初診問診票",
            "回答しにくい項目は空欄でもかまいません。",
            "月経について",
            "最終月経の開始日をご記入ください。",
            "月経周期はおおむね規則的ですか。",
            "現在、妊娠している可能性はありますか。",
            "月経に伴う症状",
            "受診理由・妊娠出産歴",
            "本日はどのような症状やご相談で来院されましたか。",
            "妊娠・出産・流産の経験について、差し支えない範囲でご記入ください。",
            "治療・手術歴",
            "服薬・アレルギー",
            "医師へ伝えておきたいこと",
            "ご記入後、受付へお渡しください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])
        identity_reading_order = [
            (item["source"], item["target"])
            for item in reference_json["relationships"]
            if item["kind"] == "reading_order"
        ][:10]
        self.assertEqual(
            identity_reading_order,
            [
                ("title", "name-label"),
                ("name-label", "furigana-label"),
                ("furigana-label", "entry-date-label"),
                ("entry-date-label", "birth-date-label"),
                ("birth-date-label", "age-label"),
                ("age-label", "age-unit"),
                ("age-unit", "gender-label"),
                ("gender-label", "gender-options"),
                ("gender-options", "optional-note"),
                ("optional-note", "menstruation-section-title"),
            ],
        )
        self.assertEqual(
            reference_json["source_dimensions"]["orientation"],
            "portrait",
        )
        self.assertTrue(fixture.manifest["redistribution_allowed"])
        self.assertFalse(fixture.manifest["contains_personal_data"])

    def test_q10_reference_is_reviewed_source_grounded_and_complete(self):
        fixture = _read_fixture(Q10_MANIFEST, Q10_ID)
        reference = fixture.reference

        self.assertEqual(fixture.source_sha256, Q10_SOURCE_SHA256)
        self.assertEqual(fixture.reference_sha256, Q10_REFERENCE_SHA256)
        self.assertEqual((fixture.source_width, fixture.source_height), (2480, 3508))
        self.assertEqual(fixture.source_dpi, 300)
        self.assertTrue(reference.reviewed)
        self.assertEqual(len(reference.text_regions), 46)
        self.assertEqual(len(reference.structural_items), 77)
        self.assertEqual(len(reference.relationships), 107)
        self.assertEqual(
            sum("checkbox" in item.id for item in reference.structural_items),
            37,
        )
        structure_ids = {item.id for item in reference.structural_items}
        self.assertEqual(
            {
                item_id
                for item_id in structure_ids
                if item_id.endswith("grid-outer")
            },
            {
                "identity-grid-outer",
                "smoking-grid-outer",
                "drinking-grid-outer",
                "exercise-sleep-grid-outer",
                "food-grid-outer",
                "disease-history-grid-outer",
            },
        )
        self.assertEqual(
            {
                item_id
                for item_id in structure_ids
                if item_id.endswith("section-band")
            },
            {
                "smoking-section-band",
                "drinking-section-band",
                "exercise-sleep-section-band",
                "food-section-band",
                "disease-history-section-band",
            },
        )
        expected_text = "".join(item.text for item in reference.text_regions)
        for phrase in (
            "健康診断・生活習慣 問診票",
            "喫煙",
            "たばこを吸いますか。吸う場合は本数と年数をご記入ください。",
            "受動喫煙",
            "禁煙希望",
            "飲酒",
            "お酒を飲む頻度と、1回あたりのおおよその量をご記入ください。",
            "休肝日",
            "運動・睡眠",
            "週にどのくらい運動していますか。",
            "睡眠で十分に休養が取れていますか。",
            "食事",
            "治療中の病気・家族歴",
            "現在、治療中または経過観察中の病気はありますか。",
            "服薬・サプリメント",
            "健診時に相談したいこと",
            "ご記入後、受付へお渡しください。",
        ):
            self.assertIn(phrase, expected_text)
        reference_json = json.loads(fixture.reference_path.read_text(encoding="utf-8"))
        self.assertEqual(reference_json["review"]["exclusions"], [])
        identity_reading_order = [
            (item["source"], item["target"])
            for item in reference_json["relationships"]
            if item["kind"] == "reading_order"
        ][:10]
        self.assertEqual(
            identity_reading_order,
            [
                ("title", "category"),
                ("category", "name-label"),
                ("name-label", "furigana-label"),
                ("furigana-label", "entry-date-label"),
                ("entry-date-label", "birth-date-label"),
                ("birth-date-label", "age-label"),
                ("age-label", "age-unit"),
                ("age-unit", "gender-label"),
                ("gender-label", "gender-options"),
                ("gender-options", "smoking-section-title"),
            ],
        )
        self.assertEqual(
            reference_json["source_dimensions"]["orientation"],
            "portrait",
        )
        self.assertTrue(fixture.manifest["redistribution_allowed"])
        self.assertFalse(fixture.manifest["contains_personal_data"])

    def _copied_q01_fixture(self, root: Path) -> Path:
        destination = root / "fixture"
        destination.mkdir()
        for name in (
            "questionnaire-01-general-medicine.png",
            "questionnaire-01-general-medicine.manifest.json",
            "questionnaire-01-general-medicine.reference.json",
        ):
            shutil.copy2(Q01_DIRECTORY / name, destination / name)
        return destination / "questionnaire-01-general-medicine.manifest.json"

    def test_source_hash_mismatch_is_rejected(self):
        with tempfile.TemporaryDirectory() as root:
            manifest_path = self._copied_q01_fixture(Path(root))
            source_path = manifest_path.parent / "questionnaire-01-general-medicine.png"
            source_path.write_bytes(source_path.read_bytes() + b"tamper")

            with self.assertRaisesRegex(RuntimeError, "source SHA-256 mismatch"):
                _read_fixture(manifest_path, Q01_ID)

    def test_reference_hash_mismatch_is_rejected(self):
        with tempfile.TemporaryDirectory() as root:
            manifest_path = self._copied_q01_fixture(Path(root))
            reference_path = (
                manifest_path.parent
                / "questionnaire-01-general-medicine.reference.json"
            )
            reference_path.write_bytes(reference_path.read_bytes() + b" ")

            with self.assertRaisesRegex(RuntimeError, "reference SHA-256 mismatch"):
                _read_fixture(manifest_path, Q01_ID)

    def test_unreviewed_reference_is_rejected_even_with_a_matching_hash(self):
        with tempfile.TemporaryDirectory() as root:
            manifest_path = self._copied_q01_fixture(Path(root))
            reference_path = (
                manifest_path.parent
                / "questionnaire-01-general-medicine.reference.json"
            )
            reference = json.loads(reference_path.read_text(encoding="utf-8"))
            reference["reviewed"] = False
            reference_path.write_text(
                json.dumps(reference, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            manifest["reference"]["sha256"] = hashlib.sha256(
                reference_path.read_bytes()
            ).hexdigest()
            manifest_path.write_text(
                json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )

            with self.assertRaisesRegex(RuntimeError, "not human-reviewed"):
                _read_fixture(manifest_path, Q01_ID)

    def test_path_traversal_is_rejected(self):
        with tempfile.TemporaryDirectory() as root:
            manifest_path = self._copied_q01_fixture(Path(root))
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            manifest["source"]["path"] = "../outside.png"
            manifest_path.write_text(
                json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )

            with self.assertRaisesRegex(RuntimeError, "without traversal"):
                _read_fixture(manifest_path, Q01_ID)

    def test_existing_output_directory_is_never_overwritten(self):
        with tempfile.TemporaryDirectory() as root:
            output = Path(root) / "already-there"
            output.mkdir()

            with self.assertRaises(FileExistsError):
                run(DEFAULT_SUITE_PATH, output)

    def test_source_page_image_and_hidden_text_are_rejected(self):
        fixture = _read_fixture(Q01_MANIFEST, Q01_ID)
        with tempfile.TemporaryDirectory() as root:
            root_path = Path(root)
            image_docx = root_path / "image.docx"
            document = Document()
            document.add_picture(str(fixture.source_path))
            document.save(image_docx)
            passed, reason, _ = _docx_media_evidence(
                image_docx,
                source_sha256=fixture.source_sha256,
                source_dimensions=(fixture.source_width, fixture.source_height),
            )
            self.assertFalse(passed)
            self.assertIn("source-sized", reason)

            hidden_docx = root_path / "hidden.docx"
            hidden = Document()
            run_element = hidden.add_paragraph().add_run("invisible truth")
            run_element.font.hidden = True
            hidden.save(hidden_docx)
            hidden_passed, hidden_reason = _hidden_text_check(hidden_docx)
            self.assertFalse(hidden_passed)
            self.assertIn("hidden-text", hidden_reason)

            layout_docx = root_path / "layout-spacer.docx"
            layout = Document()
            spacer = layout.add_paragraph().add_run("\u200b")
            spacer.font.hidden = True
            layout.save(layout_docx)
            layout_passed, layout_reason = _hidden_text_check(layout_docx)
            self.assertTrue(layout_passed)
            self.assertIn("no hidden semantic text", layout_reason)

    def test_external_relationship_makes_integrity_fail(self):
        fixture = _read_fixture(Q01_MANIFEST, Q01_ID)
        with tempfile.TemporaryDirectory() as root:
            docx_path = Path(root) / "plain.docx"
            document = Document()
            document.add_paragraph("visible")
            document.save(docx_path)
            observation = SimpleNamespace(
                package_readable=True,
                python_docx_reopenable=True,
                errors=(),
                external_relationships=("https://example.invalid/",),
            )
            snapshot = SimpleNamespace(
                page_count=1,
                pages=(object(),),
                renderer_name="test",
                renderer_version="1",
                rasterizer_name="test",
                rasterizer_version="1",
            )

            result = _integrity_report(
                fixture,
                docx_path=docx_path,
                observation=observation,
                snapshot=snapshot,
                visible_text="visible",
            )

            self.assertFalse(result["passed"])
            external = next(
                item for item in result["checks"] if item["name"] == "external_relationships"
            )
            self.assertFalse(external["passed"])


if __name__ == "__main__":
    unittest.main()
