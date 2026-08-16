import base64
import hashlib
import json
import tempfile
import unittest
from dataclasses import replace
from pathlib import Path
from zipfile import ZipFile

from docx import Document as open_docx
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from aiteqno.adapters import (
    DEFAULT_PAGE_MARGIN_PT,
    DEFAULT_SUPPORTED_FONTS,
    BundleAssetResolver,
    PythonDocxRenderer,
)
from aiteqno.application import render_docx
from aiteqno.domain import (
    BoundingBox,
    DocumentIR,
    FontStyle,
    ImageElement,
    ImageFit,
    PageSize,
    Point,
    TextAlign,
    TextElement,
)
from aiteqno.ports import DocxRenderError, RenderPolicy


FIXTURE_DIRECTORY = Path(__file__).parent / "fixtures" / "document_ir"
FIXTURE_PATH = FIXTURE_DIRECTORY / "canonical.document.ir.json"
ASSET_B64_PATH = FIXTURE_DIRECTORY / "canonical-logo.png.b64"


def load_canonical_document() -> DocumentIR:
    return DocumentIR.from_json(FIXTURE_PATH.read_text(encoding="utf-8"))


def canonical_asset_bytes() -> bytes:
    return base64.b64decode(ASSET_B64_PATH.read_text(encoding="ascii"))


def materialize_canonical_asset(bundle_root: Path, document: DocumentIR) -> Path:
    asset_path = bundle_root / document.assets[0].path
    asset_path.parent.mkdir(parents=True, exist_ok=True)
    asset_path.write_bytes(canonical_asset_bytes())
    return asset_path


def text_only_document(*elements: TextElement) -> DocumentIR:
    canonical = load_canonical_document()
    page = replace(canonical.pages[0], elements=elements)
    return replace(canonical, pages=(page,), assets=())


def visible_paragraph_text(document) -> list[str]:
    return [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text.replace("\u200b", "").strip()
    ]


class DocxRendererTest(unittest.TestCase):
    def test_canonical_visual_ir_generates_strict_reopenable_docx(self):
        document = load_canonical_document()
        with tempfile.TemporaryDirectory() as temporary_directory:
            bundle_root = Path(temporary_directory)
            materialize_canonical_asset(bundle_root, document)
            output_path = bundle_root / "reconstructed.docx"
            result = render_docx(
                document,
                output_path,
                renderer=PythonDocxRenderer(
                    supported_fonts=(*DEFAULT_SUPPORTED_FONTS, "Noto Sans CJK JP"),
                    asset_resolver=BundleAssetResolver(bundle_root),
                ),
                policy=RenderPolicy.STRICT,
            )

            self.assertEqual(result.output_path, output_path.resolve())
            with ZipFile(output_path) as package:
                package_files = set(package.namelist())
                document_xml = package.read("word/document.xml")
                relationships_xml = package.read("word/_rels/document.xml.rels")
            media_files = {
                name for name in package_files if name.startswith("word/media/")
            }
            self.assertEqual(len(media_files), 1)
            self.assertIn(b"<wp:inline", document_xml)
            self.assertNotIn(b"<wp:anchor", document_xml)
            self.assertIn(b'descr="Fixture image"', document_xml)
            self.assertNotIn(b'TargetMode="External"', relationships_xml)

            reopened = open_docx(output_path)
            output_digest = hashlib.sha256(output_path.read_bytes()).hexdigest()

        self.assertEqual(visible_paragraph_text(reopened), ["問診票"])
        section = reopened.sections[0]
        self.assertAlmostEqual(section.page_width.pt, 595.28, places=1)
        self.assertAlmostEqual(section.page_height.pt, 841.89, places=1)
        self.assertEqual(section.orientation, WD_ORIENT.PORTRAIT)
        for margin in (
            section.left_margin,
            section.right_margin,
            section.top_margin,
            section.bottom_margin,
        ):
            self.assertAlmostEqual(margin.pt, DEFAULT_PAGE_MARGIN_PT, places=1)

        title_paragraph = reopened.paragraphs[0]
        self.assertAlmostEqual(title_paragraph.paragraph_format.left_indent.pt, 12.0)
        self.assertAlmostEqual(title_paragraph.paragraph_format.space_before.pt, 6.0)
        title_run = title_paragraph.runs[0]
        self.assertEqual(title_run.font.name, "Noto Sans CJK JP")
        self.assertAlmostEqual(title_run.font.size.pt, 18.0)
        self.assertTrue(title_run.bold)
        self.assertEqual(
            title_run._element.rPr.rFonts.get(qn("w:eastAsia")),
            "Noto Sans CJK JP",
        )

        line_borders = reopened._element.body.xpath(".//w:pBdr/w:bottom")
        self.assertEqual(len(line_borders), 1)
        self.assertEqual(line_borders[0].get(qn("w:color")), "333333")
        self.assertEqual(line_borders[0].get(qn("w:sz")), "8")

        self.assertEqual(len(reopened.tables), 1)
        table = reopened.tables[0]
        grid_widths = [
            int(grid_column.get(qn("w:w"))) / 20
            for grid_column in table._tbl.tblGrid
        ]
        expected_widths = [12.0, 240.0, 24.0, 96.0, 151.28]
        self.assertEqual(len(grid_widths), len(expected_widths))
        for actual, expected in zip(grid_widths, expected_widths, strict=True):
            self.assertAlmostEqual(actual, expected, places=1)

        rectangle_borders = table.cell(0, 1)._tc.xpath("./w:tcPr/w:tcBorders/*")
        self.assertEqual(
            {border.tag.rsplit("}", 1)[-1] for border in rectangle_borders},
            {"top", "left", "bottom", "right"},
        )
        self.assertEqual(len(reopened.inline_shapes), 1)
        inline_shape = reopened.inline_shapes[0]
        self.assertAlmostEqual(inline_shape.width.pt, 96.0, places=1)
        self.assertAlmostEqual(inline_shape.height.pt, 72.0, places=1)

        report = result.report
        self.assertEqual(
            report.rendered_element_ids,
            (
                "p001-text-0000",
                "p001-line-0001",
                "p001-rectangle-0002",
                "p001-image-0003",
            ),
        )
        self.assertEqual(report.fallback_element_ids, ())
        self.assertEqual(report.omitted_element_ids, ())
        self.assertEqual(report.warnings, ())
        self.assertEqual(report.font_substitutions, ())
        self.assertEqual(
            report.output_sha256,
            output_digest,
        )
        json.dumps(report.to_dict())

    def test_missing_asset_uses_clear_placeholder_and_strict_rejects(self):
        document = load_canonical_document()
        with tempfile.TemporaryDirectory() as temporary_directory:
            bundle_root = Path(temporary_directory)
            renderer = PythonDocxRenderer(
                asset_resolver=BundleAssetResolver(bundle_root)
            )
            output_path = bundle_root / "best-effort.docx"
            result = render_docx(document, output_path, renderer=renderer)
            reopened = open_docx(output_path)
            strict_output = bundle_root / "strict.docx"
            with self.assertRaisesRegex(DocxRenderError, "strict rendering rejected"):
                render_docx(
                    document,
                    strict_output,
                    renderer=renderer,
                    policy=RenderPolicy.STRICT,
                )

        self.assertFalse(strict_output.exists())
        self.assertIn("Image unavailable", reopened.tables[0].cell(0, 3).text)
        warning_codes = {warning.code for warning in result.report.warnings}
        self.assertIn("asset_missing", warning_codes)
        self.assertIn("p001-image-0003", result.report.fallback_element_ids)
        self.assertIn("p001-image-0003", result.report.rendered_element_ids)
        self.assertEqual(result.report.omitted_element_ids, ())

    def test_no_resolver_never_reads_fixture_or_creates_external_relationship(self):
        document = load_canonical_document()
        asset_path = FIXTURE_PATH.parent / document.assets[0].path
        self.assertFalse(asset_path.exists())
        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = Path(temporary_directory) / "reconstructed.docx"
            result = render_docx(
                document,
                output_path,
                renderer=PythonDocxRenderer(),
            )
            with ZipFile(output_path) as package:
                package_files = set(package.namelist())
                relationships_xml = package.read("word/_rels/document.xml.rels")
            reopened = open_docx(output_path)

        self.assertFalse(any(name.startswith("word/media/") for name in package_files))
        self.assertNotIn(b'TargetMode="External"', relationships_xml)
        self.assertIn("Image unavailable", reopened.tables[0].cell(0, 3).text)
        warning_codes = {warning.code for warning in result.report.warnings}
        self.assertIn("asset_resolver_unavailable", warning_codes)
        self.assertEqual(result.report.omitted_element_ids, ())
        self.assertTrue(
            {
                "p001-text-0000",
                "p001-line-0001",
                "p001-rectangle-0002",
                "p001-image-0003",
            }.issubset(result.report.rendered_element_ids)
        )

    def test_reading_order_and_supported_styles_are_preserved(self):
        canonical_text = load_canonical_document().pages[0].elements[0]
        self.assertIsInstance(canonical_text, TextElement)
        first = replace(
            canonical_text,
            text="First",
            style=replace(canonical_text.style, font_family="Arial"),
        )
        second = replace(
            canonical_text,
            id="p001-text-0001",
            bbox=BoundingBox(x=72, y=96, width=240, height=18),
            text="Second",
            reading_order=1,
            style=replace(
                canonical_text.style,
                font_family="Times New Roman",
                font_size_pt=12,
                font_weight=400,
                font_style=FontStyle.ITALIC,
                color="#336699",
                align=TextAlign.RIGHT,
                line_height=1.5,
            ),
        )
        document = text_only_document(first, second)

        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = Path(temporary_directory) / "reconstructed.docx"
            result = render_docx(
                document,
                output_path,
                renderer=PythonDocxRenderer(),
                policy=RenderPolicy.STRICT,
            )
            reopened = open_docx(output_path)

        self.assertEqual(
            [paragraph.text for paragraph in reopened.paragraphs],
            ["First", "Second"],
        )
        second_paragraph = reopened.paragraphs[1]
        self.assertEqual(second_paragraph.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertAlmostEqual(second_paragraph.paragraph_format.line_spacing, 1.5)
        second_run = second_paragraph.runs[0]
        self.assertEqual(second_run.font.name, "Times New Roman")
        self.assertAlmostEqual(second_run.font.size.pt, 12.0)
        self.assertFalse(second_run.bold)
        self.assertTrue(second_run.italic)
        self.assertEqual(str(second_run.font.color.rgb), "336699")
        self.assertEqual(result.report.fallback_element_ids, ())
        self.assertEqual(result.report.omitted_element_ids, ())
        self.assertEqual(result.report.warnings, ())

    def test_unsupported_styles_have_explicit_fallbacks_and_strict_rejects_them(self):
        canonical_text = load_canonical_document().pages[0].elements[0]
        self.assertIsInstance(canonical_text, TextElement)
        styled_text = replace(
            canonical_text,
            style=replace(
                canonical_text.style,
                font_family="Imaginary Sans",
                font_weight=650,
                font_style=FontStyle.OBLIQUE,
                rotation_deg=15,
                opacity=0.5,
            ),
        )
        document = text_only_document(styled_text)

        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = Path(temporary_directory) / "best-effort.docx"
            result = render_docx(
                document,
                output_path,
                renderer=PythonDocxRenderer(),
            )
            strict_output = Path(temporary_directory) / "strict.docx"
            with self.assertRaisesRegex(DocxRenderError, "strict rendering rejected"):
                render_docx(
                    document,
                    strict_output,
                    renderer=PythonDocxRenderer(),
                    policy=RenderPolicy.STRICT,
                )
            self.assertFalse(strict_output.exists())

        warning_codes = {warning.code for warning in result.report.warnings}
        self.assertEqual(result.report.fallback_element_ids, (styled_text.id,))
        self.assertTrue(
            {
                "font_substituted",
                "font_weight_approximated",
                "oblique_mapped_to_italic",
                "rotation_omitted",
                "opacity_approximated",
            }.issubset(warning_codes)
        )

    def test_visual_approximations_are_reported_and_strict_rejects(self):
        canonical = load_canonical_document()
        line = canonical.pages[0].elements[1]
        rectangle = canonical.pages[0].elements[2]
        diagonal = replace(
            line,
            bbox=BoundingBox(x=48, y=100, width=100, height=20),
            start=Point(x=48, y=100),
            end=Point(x=148, y=120),
            style=replace(line.style, opacity=0.5),
        )
        rounded = replace(
            rectangle,
            style=replace(
                rectangle.style,
                fill_color="#ddeeff",
                corner_radius_pt=8,
                opacity=0.5,
            ),
        )
        page = replace(canonical.pages[0], elements=(diagonal, rounded))
        document = replace(canonical, pages=(page,), assets=())

        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = Path(temporary_directory) / "best-effort.docx"
            result = render_docx(
                document,
                output_path,
                renderer=PythonDocxRenderer(),
            )
            strict_output = Path(temporary_directory) / "strict.docx"
            with self.assertRaisesRegex(DocxRenderError, "strict rendering rejected"):
                render_docx(
                    document,
                    strict_output,
                    renderer=PythonDocxRenderer(),
                    policy=RenderPolicy.STRICT,
                )

        self.assertFalse(strict_output.exists())
        warning_codes = {warning.code for warning in result.report.warnings}
        self.assertTrue(
            {
                "diagonal_line_approximated",
                "rounded_rectangle_squared",
                "opacity_approximated",
            }.issubset(warning_codes)
        )
        self.assertEqual(
            set(result.report.fallback_element_ids),
            {diagonal.id, rounded.id},
        )
        self.assertEqual(result.report.omitted_element_ids, ())

    def test_vertical_line_uses_cell_border_without_fallback(self):
        canonical = load_canonical_document()
        line = canonical.pages[0].elements[1]
        vertical = replace(
            line,
            bbox=BoundingBox(x=120, y=100, width=0, height=96),
            start=Point(x=120, y=100),
            end=Point(x=120, y=196),
            style=replace(line.style, dash="dashed"),
        )
        page = replace(canonical.pages[0], elements=(vertical,))
        document = replace(canonical, pages=(page,), assets=())

        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = Path(temporary_directory) / "vertical.docx"
            result = render_docx(
                document,
                output_path,
                renderer=PythonDocxRenderer(),
                policy=RenderPolicy.STRICT,
            )
            reopened = open_docx(output_path)

        borders = reopened.tables[0]._tbl.xpath(".//w:tcBorders/w:left")
        self.assertEqual(len(borders), 1)
        self.assertEqual(borders[0].get(qn("w:val")), "dashed")
        self.assertEqual(borders[0].get(qn("w:color")), "333333")
        self.assertEqual(result.report.fallback_element_ids, ())

    def test_overlapping_elements_have_explicit_z_order_approximation(self):
        canonical = load_canonical_document()
        rectangle = canonical.pages[0].elements[2]
        image = canonical.pages[0].elements[3]
        self.assertIsInstance(image, ImageElement)
        overlapping_image = replace(image, bbox=rectangle.bbox)
        page = replace(
            canonical.pages[0],
            elements=(rectangle, overlapping_image),
        )
        document = replace(canonical, pages=(page,))

        with tempfile.TemporaryDirectory() as temporary_directory:
            bundle_root = Path(temporary_directory)
            materialize_canonical_asset(bundle_root, document)
            renderer = PythonDocxRenderer(
                asset_resolver=BundleAssetResolver(bundle_root)
            )
            output_path = bundle_root / "overlap.docx"
            result = render_docx(document, output_path, renderer=renderer)
            strict_output = bundle_root / "strict.docx"
            with self.assertRaisesRegex(DocxRenderError, "strict rendering rejected"):
                render_docx(
                    document,
                    strict_output,
                    renderer=renderer,
                    policy=RenderPolicy.STRICT,
                )

        z_order_warnings = [
            warning
            for warning in result.report.warnings
            if warning.code == "z_order_approximated"
        ]
        self.assertEqual(
            {warning.element_id for warning in z_order_warnings},
            {rectangle.id, overlapping_image.id},
        )
        self.assertEqual(
            set(result.report.fallback_element_ids),
            {rectangle.id, overlapping_image.id},
        )
        self.assertEqual(result.report.omitted_element_ids, ())

    def test_full_page_image_is_never_used_as_background(self):
        canonical = load_canonical_document()
        image = canonical.pages[0].elements[3]
        self.assertIsInstance(image, ImageElement)
        page = canonical.pages[0]
        page_image = replace(
            image,
            bbox=BoundingBox(
                x=0,
                y=0,
                width=page.size.width,
                height=page.size.height,
            ),
        )
        document = replace(
            canonical,
            pages=(replace(page, elements=(page_image,)),),
        )
        with tempfile.TemporaryDirectory() as temporary_directory:
            bundle_root = Path(temporary_directory)
            materialize_canonical_asset(bundle_root, document)
            output_path = bundle_root / "guarded.docx"
            result = render_docx(
                document,
                output_path,
                renderer=PythonDocxRenderer(
                    asset_resolver=BundleAssetResolver(bundle_root)
                ),
            )
            with ZipFile(output_path) as package:
                media_files = [
                    name
                    for name in package.namelist()
                    if name.startswith("word/media/")
                ]
            reopened = open_docx(output_path)

        self.assertEqual(media_files, [])
        table_text = " ".join(
            cell.text for row in reopened.tables[0].rows for cell in row.cells
        )
        self.assertIn("Image unavailable", table_text)
        self.assertIn(
            "source_page_background_rejected",
            {warning.code for warning in result.report.warnings},
        )

    def test_cover_and_stretch_fit_use_requested_inline_dimensions(self):
        canonical = load_canonical_document()
        image = canonical.pages[0].elements[3]
        self.assertIsInstance(image, ImageElement)
        cover = replace(
            image,
            id="p001-image-0004",
            bbox=BoundingBox(x=48, y=124, width=60, height=72),
            fit=ImageFit.COVER,
        )
        stretch = replace(
            image,
            id="p001-image-0005",
            bbox=BoundingBox(x=132, y=124, width=96, height=50),
            fit=ImageFit.STRETCH,
        )
        document = replace(
            canonical,
            pages=(replace(canonical.pages[0], elements=(cover, stretch)),),
        )
        with tempfile.TemporaryDirectory() as temporary_directory:
            bundle_root = Path(temporary_directory)
            materialize_canonical_asset(bundle_root, document)
            output_path = bundle_root / "fits.docx"
            result = render_docx(
                document,
                output_path,
                renderer=PythonDocxRenderer(
                    asset_resolver=BundleAssetResolver(bundle_root)
                ),
                policy=RenderPolicy.STRICT,
            )
            reopened = open_docx(output_path)

        self.assertEqual(len(reopened.inline_shapes), 2)
        dimensions = [
            (round(shape.width.pt), round(shape.height.pt))
            for shape in reopened.inline_shapes
        ]
        self.assertEqual(dimensions, [(60, 72), (96, 50)])
        self.assertEqual(result.report.fallback_element_ids, ())
        self.assertEqual(result.report.omitted_element_ids, ())

    def test_landscape_page_size_sets_section_orientation(self):
        canonical_text = load_canonical_document().pages[0].elements[0]
        self.assertIsInstance(canonical_text, TextElement)
        supported_text = replace(
            canonical_text,
            style=replace(canonical_text.style, font_family="Arial"),
        )
        document = text_only_document(supported_text)
        landscape_page = replace(
            document.pages[0],
            size=PageSize(width=841.89, height=595.28),
        )
        document = replace(document, pages=(landscape_page,))

        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = Path(temporary_directory) / "landscape.docx"
            render_docx(
                document,
                output_path,
                renderer=PythonDocxRenderer(),
                policy=RenderPolicy.STRICT,
            )
            reopened = open_docx(output_path)

        section = reopened.sections[0]
        self.assertEqual(section.orientation, WD_ORIENT.LANDSCAPE)
        self.assertAlmostEqual(section.page_width.pt, 841.89, places=1)
        self.assertAlmostEqual(section.page_height.pt, 595.28, places=1)


if __name__ == "__main__":
    unittest.main()
