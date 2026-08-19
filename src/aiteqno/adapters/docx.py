"""Flow-first DOCX reconstruction backed by ``python-docx``.

The adapter treats Document IR as the visual authority. It maps page settings,
text, lines, rectangles, and verified bundle-local images to interoperable
WordprocessingML. Coordinate-perfect layering is intentionally replaced by
deterministic flow layout. Pages carrying the validated table-topology
extension use real editable Word tables; other pages retain the legacy
horizontal-band path. Every approximation is reported against the affected
element ID.
"""

from __future__ import annotations

import hashlib
import math
import os
import tempfile
import unicodedata
from collections.abc import Iterable
from dataclasses import dataclass, field, replace
from io import BytesIO
from os import PathLike
from pathlib import Path

from docx import Document as open_docx
from docx.document import Document as WordDocument
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.exceptions import UnrecognizedImageError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Pt, RGBColor
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from PIL import Image

from aiteqno._version import __version__
from aiteqno.domain import (
    GEOMETRY_TOLERANCE_PT,
    Asset,
    DocumentElement,
    DocumentIR,
    FontStyle,
    ImageElement,
    ImageFit,
    LineDash,
    LineElement,
    Page,
    PageTableTopology,
    RectangleElement,
    TableCellTopology,
    TablePrimitiveRole,
    TableTopology,
    TextAlign,
    TextElement,
    TextStyle,
    read_page_table_topology,
    validate_document,
)
from aiteqno.ports import (
    AssetResolutionError,
    AssetResolver,
    DocxRenderError,
    DocxRenderReport,
    DocxRenderResult,
    FontSubstitution,
    RenderPolicy,
    RenderWarning,
    ResolvedAsset,
)


DEFAULT_PAGE_MARGIN_PT = 36.0
DEFAULT_FALLBACK_FONT = "Arial"
DEFAULT_SUPPORTED_FONTS = (
    "Arial",
    "Calibri",
    "Courier New",
    "Noto Sans CJK JP",
    "Times New Roman",
    "Yu Gothic",
)

_BAND_TOLERANCE_PT = 3.0
_MIN_LAYOUT_COLUMN_PT = 1.0
_MAX_WORD_BORDER_PT = 12.0
_MIN_WORD_BORDER_PT = 0.25
_SOURCE_PAGE_COVERAGE_LIMIT = 0.90
_TOPOLOGY_PAGE_MARGIN_PT = 18.0
_TOPOLOGY_ROW_HEIGHT_SCALE = 0.85
_TOPOLOGY_GAP_SCALE = 0.75
_SOURCE_TAG_PREFIX = "aiteqno-source:"
_TABLE_CAPTION_PREFIX = "aiteqno-table:"
_TEXT_LAYOUT_TAB_MINIMUM_PT = 24.0
_TEXT_LAYOUT_TAB_FONT_MULTIPLIER = 4.0
_TEXT_NARROW_ADVANCE_UNITS = 0.55
_TEXT_SHORT_CELL_MAX_ADVANCE_UNITS = 2.0
_TEXT_SHORT_CELL_MINIMUM_FONT_PT = 10.5
_TEXT_SHORT_CELL_HEIGHT_RATIO = 0.5
_TEXT_SHORT_CELL_MINIMUM_GLYPH_HEIGHT_PT = 8.5

_ALIGNMENT_MAP = {
    TextAlign.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    TextAlign.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    TextAlign.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
    TextAlign.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_LINE_DASH_MAP = {
    LineDash.SOLID: "single",
    LineDash.DASHED: "dashed",
    LineDash.DOTTED: "dotted",
    LineDash.DASH_DOT: "dotDash",
}


@dataclass(slots=True)
class _RenderState:
    resolver: AssetResolver | None
    rendered_ids: list[str] = field(default_factory=list)
    rendered_seen: set[str] = field(default_factory=set)
    fallback_ids: list[str] = field(default_factory=list)
    fallback_seen: set[str] = field(default_factory=set)
    omitted_ids: list[str] = field(default_factory=list)
    warnings: list[RenderWarning] = field(default_factory=list)
    warning_seen: set[tuple[str, str | None, str | None]] = field(default_factory=set)
    substitutions: list[FontSubstitution] = field(default_factory=list)
    resolved_assets: dict[str, ResolvedAsset] = field(default_factory=dict)
    unavailable_images: dict[str, tuple[str, str]] = field(default_factory=dict)
    native_table_ids: list[str] = field(default_factory=list)
    native_table_seen: set[str] = field(default_factory=set)
    native_table_consumed_ids: list[str] = field(default_factory=list)
    native_table_consumed_seen: set[str] = field(default_factory=set)

    def record_rendered(self, element_id: str) -> None:
        if element_id not in self.rendered_seen:
            self.rendered_seen.add(element_id)
            self.rendered_ids.append(element_id)

    def record_native_table(
        self,
        table_id: str,
        consumed_element_ids: Iterable[str],
    ) -> None:
        if table_id not in self.native_table_seen:
            self.native_table_seen.add(table_id)
            self.native_table_ids.append(table_id)
        for element_id in consumed_element_ids:
            if element_id not in self.native_table_consumed_seen:
                self.native_table_consumed_seen.add(element_id)
                self.native_table_consumed_ids.append(element_id)
            self.record_rendered(element_id)

    def warn_fallback(
        self,
        *,
        page_id: str,
        element_id: str,
        code: str,
        message: str,
    ) -> None:
        if element_id not in self.fallback_seen:
            self.fallback_seen.add(element_id)
            self.fallback_ids.append(element_id)
        warning_key = (code, page_id, element_id)
        if warning_key not in self.warning_seen:
            self.warning_seen.add(warning_key)
            self.warnings.append(
                RenderWarning(
                    code=code,
                    message=message,
                    page_id=page_id,
                    element_id=element_id,
                )
            )


@dataclass(frozen=True, slots=True)
class _LayoutBand:
    elements: tuple[DocumentElement, ...]
    top: float
    bottom: float


@dataclass(slots=True)
class _LayoutSlot:
    elements: list[DocumentElement]
    left: float
    right: float

    @property
    def width(self) -> float:
        return max(_MIN_LAYOUT_COLUMN_PT, self.right - self.left)


@dataclass(frozen=True, slots=True)
class _TableSegment:
    width: float
    slot: _LayoutSlot | None


@dataclass(frozen=True, slots=True)
class _TextLine:
    elements: tuple[TextElement, ...]
    top: float
    bottom: float
    left: float
    right: float


class PythonDocxRenderer:
    """Render validated Document IR without consulting its source page image."""

    renderer_name = "aiteqno-python-docx"
    renderer_version = __version__

    def __init__(
        self,
        *,
        fallback_font: str = DEFAULT_FALLBACK_FONT,
        supported_fonts: Iterable[str] | None = None,
        asset_resolver: AssetResolver | None = None,
    ) -> None:
        if not fallback_font.strip():
            raise ValueError("fallback_font must not be empty")
        font_names = tuple(
            DEFAULT_SUPPORTED_FONTS if supported_fonts is None else supported_fonts
        )
        if not font_names or any(
            not isinstance(name, str) or not name.strip() for name in font_names
        ):
            raise ValueError("supported_fonts must contain non-empty names")
        if fallback_font.casefold() not in {name.casefold() for name in font_names}:
            font_names = (*font_names, fallback_font)
        self._fallback_font = fallback_font
        self._supported_fonts = {name.casefold(): name for name in font_names}
        self._asset_resolver = asset_resolver

    def render(
        self,
        document: DocumentIR,
        output_path: str | PathLike[str],
        *,
        policy: RenderPolicy = RenderPolicy.BEST_EFFORT,
    ) -> DocxRenderResult:
        """Render a valid DOCX atomically and return a complete report."""

        if not isinstance(document, DocumentIR):
            raise TypeError("document must be a DocumentIR")
        validate_document(document)
        selected_policy = RenderPolicy(policy)
        target = Path(output_path)
        if target.suffix.lower() != ".docx":
            raise ValueError("output_path must use the .docx extension")

        state = _RenderState(resolver=self._asset_resolver)
        self._prepare_assets(document, state)
        self._reject_strict_fallbacks(selected_policy, state)

        word_document = open_docx()
        title = document.metadata.get("title")
        if isinstance(title, str):
            word_document.core_properties.title = title

        for page_index, page in enumerate(document.pages):
            if page_index == 0:
                section = word_document.sections[0]
            else:
                section = word_document.add_section(WD_SECTION.NEW_PAGE)
            topology = read_page_table_topology(page)
            horizontal_margin, vertical_margin = self._configure_section(
                section,
                page,
                topology_page=topology is not None and bool(topology.tables),
            )
            if topology is not None and topology.tables:
                self._render_topology_page(
                    word_document,
                    section,
                    page,
                    topology,
                    horizontal_margin=horizontal_margin,
                    vertical_margin=vertical_margin,
                    state=state,
                )
            else:
                self._render_page(
                    word_document,
                    page,
                    horizontal_margin=horizontal_margin,
                    vertical_margin=vertical_margin,
                    state=state,
                )

        self._reject_strict_fallbacks(selected_policy, state)
        self._save_atomically(word_document, target)
        resolved_target = target.resolve()
        output_sha256 = hashlib.sha256(resolved_target.read_bytes()).hexdigest()
        report = DocxRenderReport(
            renderer_name=self.renderer_name,
            renderer_version=self.renderer_version,
            ir_version=document.ir_version,
            output_path=str(resolved_target),
            output_sha256=output_sha256,
            rendered_element_ids=tuple(state.rendered_ids),
            fallback_element_ids=tuple(state.fallback_ids),
            omitted_element_ids=tuple(state.omitted_ids),
            warnings=tuple(state.warnings),
            errors=(),
            font_substitutions=tuple(state.substitutions),
            native_table_ids=tuple(state.native_table_ids),
            native_table_consumed_element_ids=tuple(state.native_table_consumed_ids),
        )
        return DocxRenderResult(output_path=resolved_target, report=report)

    @staticmethod
    def _reject_strict_fallbacks(
        policy: RenderPolicy,
        state: _RenderState,
    ) -> None:
        if policy is RenderPolicy.STRICT and (state.fallback_ids or state.omitted_ids):
            details = ", ".join((*state.fallback_ids, *state.omitted_ids))
            raise DocxRenderError(
                f"strict rendering rejected approximated or omitted elements: {details}"
            )

    def _prepare_assets(self, document: DocumentIR, state: _RenderState) -> None:
        assets_by_id = {asset.id: asset for asset in document.assets}
        resolution_errors: dict[str, AssetResolutionError] = {}
        for page in document.pages:
            for element in page.elements:
                if not isinstance(element, ImageElement):
                    continue
                asset = assets_by_id[element.asset_id]
                background_reason = self._source_page_background_reason(
                    page,
                    element,
                    asset,
                )
                if background_reason is not None:
                    code = "source_page_background_rejected"
                    state.unavailable_images[element.id] = (code, background_reason)
                    state.warn_fallback(
                        page_id=page.id,
                        element_id=element.id,
                        code=code,
                        message=background_reason,
                    )
                    continue
                if state.resolver is None:
                    code = "asset_resolver_unavailable"
                    message = (
                        f"image asset {asset.id!r} was replaced with a placeholder "
                        "because no bundle asset resolver was configured"
                    )
                    state.unavailable_images[element.id] = (code, message)
                    state.warn_fallback(
                        page_id=page.id,
                        element_id=element.id,
                        code=code,
                        message=message,
                    )
                    continue
                if (
                    asset.id not in state.resolved_assets
                    and asset.id not in resolution_errors
                ):
                    try:
                        resolved = state.resolver.resolve(asset)
                        if resolved.asset_id != asset.id:
                            raise AssetResolutionError(
                                "asset_identity_mismatch",
                                asset.id,
                                (
                                    f"resolver returned asset {resolved.asset_id!r} "
                                    f"for registry entry {asset.id!r}"
                                ),
                            )
                        state.resolved_assets[asset.id] = resolved
                    except AssetResolutionError as exc:
                        resolution_errors[asset.id] = exc
                error = resolution_errors.get(asset.id)
                if error is not None:
                    message = f"{error}; image was replaced with a placeholder"
                    state.unavailable_images[element.id] = (error.code, message)
                    state.warn_fallback(
                        page_id=page.id,
                        element_id=element.id,
                        code=error.code,
                        message=message,
                    )

    @staticmethod
    def _source_page_background_reason(
        page: Page,
        element: ImageElement,
        asset: Asset,
    ) -> str | None:
        width_coverage = element.bbox.width / page.size.width
        height_coverage = element.bbox.height / page.size.height
        area_coverage = (
            element.bbox.width
            * element.bbox.height
            / (page.size.width * page.size.height)
        )
        source_size_match = (
            page.source is not None
            and asset.pixel_width == page.source.pixel_width
            and asset.pixel_height == page.source.pixel_height
        )
        if area_coverage >= _SOURCE_PAGE_COVERAGE_LIMIT or (
            source_size_match and width_coverage >= 0.75 and height_coverage >= 0.75
        ):
            return (
                f"image {element.id!r} substantially covers page {page.id!r}; "
                "whole-source-page backgrounds are prohibited and a placeholder was used"
            )
        return None

    @staticmethod
    def _configure_section(
        section: Section,
        page: Page,
        *,
        topology_page: bool = False,
    ) -> tuple[float, float]:
        width = page.size.width
        height = page.size.height
        section.orientation = (
            WD_ORIENT.LANDSCAPE if width > height else WD_ORIENT.PORTRAIT
        )
        section.page_width = Pt(width)
        section.page_height = Pt(height)

        horizontal_requested = (
            _TOPOLOGY_PAGE_MARGIN_PT if topology_page else DEFAULT_PAGE_MARGIN_PT
        )
        vertical_requested = 0.0 if topology_page else DEFAULT_PAGE_MARGIN_PT
        horizontal_margin = min(horizontal_requested, width / 4)
        vertical_margin = min(vertical_requested, height / 4)
        section.left_margin = Pt(horizontal_margin)
        section.right_margin = Pt(horizontal_margin)
        section.top_margin = Pt(vertical_margin)
        section.bottom_margin = Pt(vertical_margin)
        section.gutter = Pt(0)
        return horizontal_margin, vertical_margin

    def _render_topology_page(
        self,
        word_document: WordDocument,
        section: Section,
        page: Page,
        topology: PageTableTopology,
        *,
        horizontal_margin: float,
        vertical_margin: float,
        state: _RenderState,
    ) -> None:
        """Render one validated topology page without legacy layout tables."""

        table_element_ids = {
            element_id
            for table in topology.tables
            for element_id in (
                *table.supporting_element_ids,
                *(text_id for cell in table.cells for text_id in cell.text_element_ids),
            )
        }
        page_frame_ids = {
            assignment.element_id
            for assignment in topology.primitive_roles
            if assignment.role is TablePrimitiveRole.PAGE_FRAME
        }
        self._render_page_frame(
            section,
            page,
            page_frame_ids=page_frame_ids,
            state=state,
        )

        remaining = tuple(
            element
            for element in page.elements
            if element.id not in table_element_ids and element.id not in page_frame_ids
        )
        blocks: list[tuple[float, int, str, _LayoutBand | TableTopology]] = [
            (band.top, 0, min(element.id for element in band.elements), band)
            for band in self._cluster_bands(remaining)
        ]
        blocks.extend((table.bbox.y, 1, table.id, table) for table in topology.tables)
        blocks.sort(key=lambda item: (item[0], item[1], item[2]))

        previous_bottom = vertical_margin
        for _top, _kind, _identifier, block in blocks:
            if isinstance(block, TableTopology):
                previous_bottom = self._render_native_table(
                    word_document,
                    page,
                    topology,
                    block,
                    horizontal_margin=horizontal_margin,
                    previous_bottom=previous_bottom,
                    state=state,
                )
                continue
            if all(isinstance(element, TextElement) for element in block.elements):
                previous_bottom = self._render_topology_text_band(
                    word_document,
                    page,
                    block,
                    horizontal_margin=horizontal_margin,
                    previous_bottom=previous_bottom,
                    state=state,
                )
                continue
            if all(
                isinstance(element, LineElement) and self._is_horizontal_line(element)
                for element in block.elements
            ):
                for element in block.elements:
                    assert isinstance(element, LineElement)
                    previous_bottom = self._render_topology_horizontal_line(
                        word_document,
                        page,
                        element,
                        horizontal_margin=horizontal_margin,
                        previous_bottom=previous_bottom,
                        state=state,
                    )
                continue
            if all(
                isinstance(element, TextElement)
                or (
                    isinstance(element, LineElement)
                    and self._is_horizontal_line(element)
                )
                for element in block.elements
            ):
                text_elements = tuple(
                    element
                    for element in block.elements
                    if isinstance(element, TextElement)
                )
                if text_elements:
                    text_band = _LayoutBand(
                        elements=text_elements,
                        top=min(element.bbox.y for element in text_elements),
                        bottom=max(element.bbox.bottom for element in text_elements),
                    )
                    previous_bottom = self._render_topology_text_band(
                        word_document,
                        page,
                        text_band,
                        horizontal_margin=horizontal_margin,
                        previous_bottom=previous_bottom,
                        state=state,
                    )
                for element in block.elements:
                    if not isinstance(element, LineElement):
                        continue
                    previous_bottom = self._render_topology_horizontal_line(
                        word_document,
                        page,
                        element,
                        horizontal_margin=horizontal_margin,
                        previous_bottom=previous_bottom,
                        state=state,
                    )
                continue

            # Non-grid decorations remain on the established best-effort path.
            previous_bottom = self._render_table_band(
                word_document,
                page,
                block,
                horizontal_margin=horizontal_margin,
                previous_bottom=previous_bottom,
                state=state,
                tag_source_elements=True,
            )

    def _render_page_frame(
        self,
        section: Section,
        page: Page,
        *,
        page_frame_ids: set[str],
        state: _RenderState,
    ) -> None:
        if not page_frame_ids:
            return
        elements = {
            element.id: element
            for element in page.elements
            if element.id in page_frame_ids
        }
        rectangle = next(
            (
                element
                for element in elements.values()
                if isinstance(element, RectangleElement)
            ),
            None,
        )
        line = next(
            (
                element
                for element in elements.values()
                if isinstance(element, LineElement)
            ),
            None,
        )
        if rectangle is not None:
            if (
                rectangle.style.stroke_width_pt <= 0
                or rectangle.style.stroke_color is None
            ):
                border = {"val": "nil"}
            else:
                border = self._visual_border(
                    page,
                    rectangle,
                    width_pt=rectangle.style.stroke_width_pt,
                    color=rectangle.style.stroke_color,
                    dash=LineDash.SOLID,
                    state=state,
                )
        elif line is not None:
            border = self._line_border(page, line, state)
        else:
            return
        self._set_section_page_borders(section, border)
        for element in page.elements:
            if element.id in page_frame_ids:
                state.record_rendered(element.id)

    def _render_topology_text_band(
        self,
        word_document: WordDocument,
        page: Page,
        band: _LayoutBand,
        *,
        horizontal_margin: float,
        previous_bottom: float,
        state: _RenderState,
    ) -> float:
        text_elements = tuple(
            element
            for element in band.elements
            if isinstance(element, TextElement)
        )
        content_right = page.size.width - horizontal_margin
        current_bottom = previous_bottom
        for line in self._plan_text_lines(text_elements):
            paragraph = word_document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Pt(
                max(0.0, line.left - horizontal_margin)
            )
            paragraph_format.right_indent = Pt(max(0.0, content_right - line.right))
            paragraph_format.space_before = Pt(
                max(0.0, line.top - current_bottom) * _TOPOLOGY_GAP_SCALE
            )
            paragraph_format.space_after = Pt(0)
            paragraph_format.keep_together = True
            previous: TextElement | None = None
            font_size_pt = self._line_font_size(line)
            for element in line.elements:
                self._append_text_separator(
                    paragraph,
                    previous,
                    element,
                    container_left=horizontal_margin,
                )
                self._apply_text_style(
                    paragraph,
                    page,
                    element,
                    state=state,
                    source_element_id=element.id,
                    font_size_pt=font_size_pt,
                )
                state.record_rendered(element.id)
                previous = element
            current_bottom = max(current_bottom, line.bottom)
        return max(current_bottom, band.bottom)

    def _render_topology_horizontal_line(
        self,
        word_document: WordDocument,
        page: Page,
        element: LineElement,
        *,
        horizontal_margin: float,
        previous_bottom: float,
        state: _RenderState,
    ) -> float:
        paragraph = word_document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        content_right = page.size.width - horizontal_margin
        left = min(element.start.x, element.end.x)
        right = max(element.start.x, element.end.x)
        paragraph_format.left_indent = Pt(max(0.0, left - horizontal_margin))
        paragraph_format.right_indent = Pt(max(0.0, content_right - right))
        paragraph_format.space_before = Pt(
            max(0.0, min(element.start.y, element.end.y) - previous_bottom)
            * _TOPOLOGY_GAP_SCALE
        )
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = Pt(1)
        run = paragraph.add_run("\u200b")
        run.font.size = Pt(1)
        run.font.hidden = True
        self._set_paragraph_bottom_border(
            paragraph,
            self._line_border(page, element, state),
        )
        state.record_rendered(element.id)
        return max(previous_bottom, max(element.start.y, element.end.y))

    def _render_native_table(
        self,
        word_document: WordDocument,
        page: Page,
        page_topology: PageTableTopology,
        table_topology: TableTopology,
        *,
        horizontal_margin: float,
        previous_bottom: float,
        state: _RenderState,
    ) -> float:
        gap = max(0.0, table_topology.bbox.y - previous_bottom) * _TOPOLOGY_GAP_SCALE
        self._add_flow_spacer(word_document, gap)

        column_widths = tuple(
            column.end - column.start for column in table_topology.columns
        )
        table = word_document.add_table(
            rows=table_topology.logical_rows,
            cols=table_topology.logical_columns,
        )
        self._configure_native_table(
            table,
            table_id=table_topology.id,
            column_widths=column_widths,
            table_width=table_topology.bbox.width,
            indent=max(0.0, table_topology.bbox.x - horizontal_margin),
        )
        for row, axis in zip(table.rows, table_topology.rows, strict=True):
            row.height = Pt((axis.end - axis.start) * _TOPOLOGY_ROW_HEIGHT_SCALE)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row_properties = row._tr.get_or_add_trPr()
            cannot_split = OxmlElement("w:cantSplit")
            row_properties.append(cannot_split)

        for cell_topology in table_topology.cells:
            if cell_topology.rowspan == 1 and cell_topology.colspan == 1:
                continue
            start = table.cell(
                cell_topology.row_index,
                cell_topology.column_index,
            )
            end = table.cell(
                cell_topology.row_index + cell_topology.rowspan - 1,
                cell_topology.column_index + cell_topology.colspan - 1,
            )
            start.merge(end)

        elements = {element.id: element for element in page.elements}
        cell_rectangles = {
            assignment.cell_id: assignment.element_id
            for assignment in page_topology.primitive_roles
            if assignment.role is TablePrimitiveRole.CELL_RECTANGLE
            and assignment.table_id == table_topology.id
            and assignment.cell_id is not None
        }
        for cell_topology in table_topology.cells:
            cell = table.cell(
                cell_topology.row_index,
                cell_topology.column_index,
            )
            width = sum(
                column_widths[
                    cell_topology.column_index : cell_topology.column_index
                    + cell_topology.colspan
                ]
            )
            self._configure_cell(cell, width)
            self._clear_cell(cell)
            rectangle_id = cell_rectangles[cell_topology.id]
            rectangle = elements[rectangle_id]
            assert isinstance(rectangle, RectangleElement)
            self._render_rectangle(cell, page, rectangle, state)
            self._render_native_cell_text(
                cell,
                page,
                cell_topology,
                elements=elements,
                state=state,
            )

        consumed_ids = {
            *table_topology.supporting_element_ids,
            *(
                text_id
                for cell in table_topology.cells
                for text_id in cell.text_element_ids
            ),
        }
        state.record_native_table(
            table_topology.id,
            (element.id for element in page.elements if element.id in consumed_ids),
        )
        return max(previous_bottom, table_topology.bbox.bottom)

    def _render_native_cell_text(
        self,
        cell: _Cell,
        page: Page,
        cell_topology: TableCellTopology,
        *,
        elements: dict[str, DocumentElement],
        state: _RenderState,
    ) -> None:
        text_elements = tuple(
            elements[element_id] for element_id in cell_topology.text_element_ids
        )
        if not text_elements:
            return
        assert all(isinstance(element, TextElement) for element in text_elements)
        lines = self._plan_text_lines(
            tuple(
                element for element in text_elements if isinstance(element, TextElement)
            )
        )
        for line_index, line in enumerate(lines):
            paragraph = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.keep_together = True
            paragraph_format.left_indent = Pt(
                max(0.0, line.left - cell_topology.bbox.x)
            )
            paragraph_format.right_indent = Pt(
                max(0.0, cell_topology.bbox.right - line.right)
            )
            previous: TextElement | None = None
            font_size_pt = self._cell_line_font_size(line, cell_topology)
            for element in line.elements:
                self._append_text_separator(
                    paragraph,
                    previous,
                    element,
                    container_left=cell_topology.bbox.x,
                )
                self._apply_text_style(
                    paragraph,
                    page,
                    element,
                    state=state,
                    source_element_id=element.id,
                    font_size_pt=font_size_pt,
                )
                previous = element

    @staticmethod
    def _plan_text_lines(
        elements: tuple[TextElement, ...],
    ) -> tuple[_TextLine, ...]:
        lines: list[list[TextElement]] = []
        for element in sorted(
            elements,
            key=lambda item: (
                item.bbox.y,
                item.bbox.x,
                item.reading_order,
                item.id,
            ),
        ):
            candidates: list[tuple[float, float, float, int]] = []
            element_center = (element.bbox.y + element.bbox.bottom) / 2.0
            for index, line in enumerate(lines):
                line_top = min(item.bbox.y for item in line)
                line_bottom = max(item.bbox.bottom for item in line)
                overlap = min(line_bottom, element.bbox.bottom) - max(
                    line_top,
                    element.bbox.y,
                )
                if overlap < -_BAND_TOLERANCE_PT:
                    continue
                smaller_height = max(
                    1.0,
                    min(line_bottom - line_top, element.bbox.height),
                )
                overlap_ratio = max(0.0, overlap) / smaller_height
                line_center = (line_top + line_bottom) / 2.0
                candidates.append(
                    (
                        overlap_ratio,
                        -abs(element_center - line_center),
                        -line_top,
                        -index,
                    )
                )
            if candidates:
                candidate_index = -max(candidates)[3]
                lines[candidate_index].append(element)
            else:
                lines.append([element])

        planned = [
            _TextLine(
                elements=tuple(
                    sorted(
                        line,
                        key=lambda item: (
                            item.bbox.x,
                            item.reading_order,
                            item.bbox.y,
                            item.id,
                        ),
                    )
                ),
                top=min(item.bbox.y for item in line),
                bottom=max(item.bbox.bottom for item in line),
                left=min(item.bbox.x for item in line),
                right=max(item.bbox.right for item in line),
            )
            for line in lines
        ]
        return tuple(
            sorted(
                planned,
                key=lambda line: (
                    line.top,
                    line.left,
                    min(element.reading_order for element in line.elements),
                    min(element.id for element in line.elements),
                ),
            )
        )

    @staticmethod
    def _append_text_separator(
        paragraph: Paragraph,
        previous: TextElement | None,
        current: TextElement,
        *,
        container_left: float,
    ) -> None:
        if previous is None:
            return
        gap = max(0.0, current.bbox.x - previous.bbox.right)
        if gap <= GEOMETRY_TOLERANCE_PT:
            return
        font_size = min(previous.style.font_size_pt, current.style.font_size_pt)
        layout_tab_threshold = max(
            _TEXT_LAYOUT_TAB_MINIMUM_PT,
            max(previous.style.font_size_pt, current.style.font_size_pt)
            * _TEXT_LAYOUT_TAB_FONT_MULTIPLIER,
        )
        if gap >= layout_tab_threshold:
            tab_position = max(0.0, current.bbox.x - container_left)
            paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(tab_position))
            separator = paragraph.add_run()
            separator.add_tab()
            separator.font.size = Pt(font_size)
            return
        if PythonDocxRenderer._requires_word_space(previous.text, current.text):
            separator = paragraph.add_run(" ")
            separator.font.size = Pt(font_size)

    @staticmethod
    def _line_font_size(line: _TextLine) -> float:
        """Stabilize OCR-fragment sizes without exceeding their source span."""

        source_maximum = max(element.style.font_size_pt for element in line.elements)
        advance_units = sum(
            PythonDocxRenderer._text_advance_units(element.text)
            for element in line.elements
        )
        advance_units += _TEXT_NARROW_ADVANCE_UNITS * sum(
            PythonDocxRenderer._requires_word_space(previous.text, current.text)
            for previous, current in zip(
                line.elements,
                line.elements[1:],
                strict=False,
            )
        )
        if advance_units <= 0:
            return source_maximum
        geometry_maximum = line.right - line.left
        if geometry_maximum <= 0:
            return source_maximum
        return min(source_maximum, max(0.5, geometry_maximum / advance_units))

    @staticmethod
    def _cell_line_font_size(
        line: _TextLine,
        cell_topology: TableCellTopology,
    ) -> float:
        font_size = PythonDocxRenderer._line_font_size(line)
        advance_units = PythonDocxRenderer._line_advance_units(line)
        if not PythonDocxRenderer._is_short_cell_line(line):
            return font_size
        available_width = max(0.0, cell_topology.bbox.right - line.left)
        readable_minimum = min(
            _TEXT_SHORT_CELL_MINIMUM_FONT_PT,
            cell_topology.bbox.height * _TEXT_SHORT_CELL_HEIGHT_RATIO,
            available_width / advance_units,
        )
        return max(font_size, readable_minimum)

    @staticmethod
    def _is_short_cell_line(line: _TextLine) -> bool:
        advance_units = PythonDocxRenderer._line_advance_units(line)
        return (
            len(line.elements) == 1
            and line.elements[0].bbox.height
            >= _TEXT_SHORT_CELL_MINIMUM_GLYPH_HEIGHT_PT
            and 0 < advance_units <= _TEXT_SHORT_CELL_MAX_ADVANCE_UNITS
        )

    @staticmethod
    def _line_advance_units(line: _TextLine) -> float:
        return sum(
            PythonDocxRenderer._text_advance_units(element.text)
            for element in line.elements
        )

    @staticmethod
    def _text_advance_units(text: str) -> float:
        return sum(
            1.0
            if PythonDocxRenderer._is_cjk_character(character)
            else _TEXT_NARROW_ADVANCE_UNITS
            for character in text
            if not character.isspace()
        )

    @staticmethod
    def _requires_word_space(previous_text: str, current_text: str) -> bool:
        previous_char = next(
            (character for character in reversed(previous_text) if not character.isspace()),
            "",
        )
        current_char = next(
            (character for character in current_text if not character.isspace()),
            "",
        )
        if not previous_char or not current_char:
            return False
        if PythonDocxRenderer._is_numbered_prefix(previous_text):
            return PythonDocxRenderer._is_word_character(current_char)
        return (
            PythonDocxRenderer._is_word_character(previous_char)
            and PythonDocxRenderer._is_word_character(current_char)
            and not PythonDocxRenderer._is_cjk_character(previous_char)
            and not PythonDocxRenderer._is_cjk_character(current_char)
        )

    @staticmethod
    def _is_numbered_prefix(text: str) -> bool:
        candidate = text.strip()
        if len(candidate) < 2 or candidate[-1] not in {".", "．", ")", "）"}:
            return False
        return all(unicodedata.category(character) == "Nd" for character in candidate[:-1])

    @staticmethod
    def _is_word_character(character: str) -> bool:
        return unicodedata.category(character)[0] in {"L", "N"}

    @staticmethod
    def _is_cjk_character(character: str) -> bool:
        codepoint = ord(character)
        return any(
            start <= codepoint <= end
            for start, end in (
                (0x2E80, 0x2FFF),
                (0x3000, 0x303F),
                (0x3040, 0x30FF),
                (0x31F0, 0x31FF),
                (0x3400, 0x4DBF),
                (0x4E00, 0x9FFF),
                (0xAC00, 0xD7AF),
                (0xF900, 0xFAFF),
                (0x20000, 0x2FA1F),
            )
        )

    @staticmethod
    def _add_flow_spacer(word_document: WordDocument, gap: float) -> None:
        if gap <= GEOMETRY_TOLERANCE_PT:
            return
        spacer = word_document.add_paragraph()
        spacer.paragraph_format.space_before = Pt(max(0.0, gap - 1.0))
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.line_spacing = Pt(1)
        run = spacer.add_run("\u200b")
        run.font.size = Pt(1)
        run.font.hidden = True

    @staticmethod
    def _configure_native_table(
        table: Table,
        *,
        table_id: str,
        column_widths: tuple[float, ...],
        table_width: float,
        indent: float,
    ) -> None:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        properties = table._tbl.tblPr

        width = properties.find(qn("w:tblW"))
        if width is None:
            width = OxmlElement("w:tblW")
            properties.insert(0, width)
        width.set(qn("w:type"), "dxa")
        width.set(qn("w:w"), str(_points_to_twips(table_width)))

        table_indent = properties.find(qn("w:tblInd"))
        if table_indent is None:
            table_indent = OxmlElement("w:tblInd")
            properties.append(table_indent)
        table_indent.set(qn("w:type"), "dxa")
        table_indent.set(qn("w:w"), str(_points_to_twips(indent)))

        layout = properties.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            properties.append(layout)
        layout.set(qn("w:type"), "fixed")

        caption = properties.find(qn("w:tblCaption"))
        if caption is None:
            caption = OxmlElement("w:tblCaption")
            properties.append(caption)
        caption.set(qn("w:val"), f"{_TABLE_CAPTION_PREFIX}{table_id}")

        description = properties.find(qn("w:tblDescription"))
        if description is None:
            description = OxmlElement("w:tblDescription")
            properties.append(description)
        description.set(
            qn("w:val"),
            f"Editable table reconstructed from Document IR topology {table_id}",
        )

        borders = properties.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            properties.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            node.set(qn("w:val"), "nil")

        widths_twips = _distributed_twips(column_widths, table_width)
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width_twips in widths_twips:
            grid_column = OxmlElement("w:gridCol")
            grid_column.set(qn("w:w"), str(width_twips))
            grid.append(grid_column)

    @staticmethod
    def _clear_cell(cell: _Cell) -> None:
        for paragraph in tuple(cell._tc.findall(qn("w:p"))):
            cell._tc.remove(paragraph)
        cell.add_paragraph()

    @staticmethod
    def _set_section_page_borders(
        section: Section,
        border: dict[str, str],
    ) -> None:
        section_properties = section._sectPr
        borders = section_properties.find(qn("w:pgBorders"))
        if borders is None:
            borders = OxmlElement("w:pgBorders")
            section_properties.append(borders)
        borders.set(qn("w:offsetFrom"), "page")
        for edge in ("top", "left", "bottom", "right"):
            node = borders.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                borders.append(node)
            for key, value in border.items():
                node.set(qn(f"w:{key}"), value)
            node.set(qn("w:space"), "0")

    def _render_page(
        self,
        word_document: WordDocument,
        page: Page,
        *,
        horizontal_margin: float,
        vertical_margin: float,
        state: _RenderState,
    ) -> None:
        previous_bottom = vertical_margin
        for band in self._cluster_bands(page.elements):
            if len(band.elements) == 1 and isinstance(band.elements[0], TextElement):
                element = band.elements[0]
                paragraph = word_document.add_paragraph()
                previous_bottom = self._apply_text_position(
                    paragraph,
                    page,
                    element,
                    horizontal_margin=horizontal_margin,
                    previous_bottom=previous_bottom,
                    state=state,
                )
                self._apply_text_style(paragraph, page, element, state=state)
                state.record_rendered(element.id)
                continue
            if (
                len(band.elements) == 1
                and isinstance(band.elements[0], LineElement)
                and self._is_horizontal_line(band.elements[0])
            ):
                previous_bottom = self._render_horizontal_line(
                    word_document,
                    page,
                    band.elements[0],
                    horizontal_margin=horizontal_margin,
                    previous_bottom=previous_bottom,
                    state=state,
                )
                continue
            previous_bottom = self._render_table_band(
                word_document,
                page,
                band,
                horizontal_margin=horizontal_margin,
                previous_bottom=previous_bottom,
                state=state,
            )

    @staticmethod
    def _cluster_bands(
        elements: tuple[DocumentElement, ...],
    ) -> tuple[_LayoutBand, ...]:
        def sort_key(element: DocumentElement) -> tuple[float, int, float, int, str]:
            reading_order = (
                element.reading_order if isinstance(element, TextElement) else 2**31 - 1
            )
            return (
                element.bbox.y,
                reading_order,
                element.bbox.x,
                element.z_index,
                element.id,
            )

        bands: list[_LayoutBand] = []
        for element in sorted(elements, key=sort_key):
            top = element.bbox.y
            bottom = element.bbox.bottom
            if isinstance(element, LineElement):
                top = min(element.start.y, element.end.y)
                bottom = max(element.start.y, element.end.y)
            if bands and top <= bands[-1].bottom + _BAND_TOLERANCE_PT:
                previous = bands[-1]
                bands[-1] = _LayoutBand(
                    elements=(*previous.elements, element),
                    top=min(previous.top, top),
                    bottom=max(previous.bottom, bottom),
                )
            else:
                bands.append(_LayoutBand(elements=(element,), top=top, bottom=bottom))
        return tuple(bands)

    @staticmethod
    def _apply_text_position(
        paragraph: Paragraph,
        page: Page,
        element: TextElement,
        *,
        horizontal_margin: float,
        previous_bottom: float,
        state: _RenderState,
    ) -> float:
        paragraph_format = paragraph.paragraph_format
        content_right = page.size.width - horizontal_margin
        left_indent = element.bbox.x - horizontal_margin
        right_indent = content_right - element.bbox.right
        space_before = element.bbox.y - previous_bottom

        if (
            left_indent < -GEOMETRY_TOLERANCE_PT
            or right_indent < -GEOMETRY_TOLERANCE_PT
        ):
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="text_outside_layout_margins",
                message="text was moved inside the DOCX compatibility margins",
            )
        if space_before < -GEOMETRY_TOLERANCE_PT:
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="vertical_position_approximated",
                message=(
                    "reading order moves upward or overlaps; flow layout used zero "
                    "additional spacing"
                ),
            )

        paragraph_format.left_indent = Pt(max(0.0, left_indent))
        paragraph_format.right_indent = Pt(max(0.0, right_indent))
        paragraph_format.space_before = Pt(max(0.0, space_before))
        paragraph_format.space_after = Pt(0)
        paragraph_format.keep_together = True
        return max(previous_bottom, element.bbox.bottom)

    def _render_horizontal_line(
        self,
        word_document: WordDocument,
        page: Page,
        element: LineElement,
        *,
        horizontal_margin: float,
        previous_bottom: float,
        state: _RenderState,
    ) -> float:
        paragraph = word_document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        content_right = page.size.width - horizontal_margin
        left_indent = min(element.start.x, element.end.x) - horizontal_margin
        right_indent = content_right - max(element.start.x, element.end.x)
        space_before = min(element.start.y, element.end.y) - previous_bottom
        if (
            left_indent < -GEOMETRY_TOLERANCE_PT
            or right_indent < -GEOMETRY_TOLERANCE_PT
        ):
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="visual_outside_layout_margins",
                message="line was clipped to the DOCX compatibility margins",
            )
        if space_before < -GEOMETRY_TOLERANCE_PT:
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="vertical_position_approximated",
                message="overlapping line was placed at the next available flow position",
            )
        paragraph_format.left_indent = Pt(max(0.0, left_indent))
        paragraph_format.right_indent = Pt(max(0.0, right_indent))
        paragraph_format.space_before = Pt(max(0.0, space_before))
        paragraph_format.space_after = Pt(0)
        paragraph_format.keep_with_next = True
        border = self._line_border(page, element, state)
        self._set_paragraph_bottom_border(paragraph, border)
        state.record_rendered(element.id)
        return max(previous_bottom, max(element.start.y, element.end.y))

    def _render_table_band(
        self,
        word_document: WordDocument,
        page: Page,
        band: _LayoutBand,
        *,
        horizontal_margin: float,
        previous_bottom: float,
        state: _RenderState,
        tag_source_elements: bool = False,
    ) -> float:
        if band.top < previous_bottom - GEOMETRY_TOLERANCE_PT:
            for element in band.elements:
                state.warn_fallback(
                    page_id=page.id,
                    element_id=element.id,
                    code="vertical_position_approximated",
                    message=(
                        "overlapping visual band was placed at the next available "
                        "flow position"
                    ),
                )
        gap = max(0.0, band.top - previous_bottom)
        if gap > GEOMETRY_TOLERANCE_PT:
            spacer = word_document.add_paragraph()
            spacer.paragraph_format.space_before = Pt(max(0.0, gap - 1.0))
            spacer.paragraph_format.space_after = Pt(0)
            spacer_run = spacer.add_run("\u200b")
            spacer_run.font.size = Pt(1)
            spacer_run.font.hidden = True

        content_width = page.size.width - 2 * horizontal_margin
        segments = self._derive_segments(
            page,
            band,
            horizontal_margin=horizontal_margin,
            content_width=content_width,
            state=state,
        )
        table = word_document.add_table(rows=1, cols=len(segments))
        self._configure_fixed_table(table, segments, content_width)
        row = table.rows[0]
        row.height = Pt(max(_MIN_LAYOUT_COLUMN_PT, band.bottom - band.top))
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        for cell, segment in zip(row.cells, segments, strict=True):
            self._configure_cell(cell, segment.width)
            if segment.slot is None:
                continue
            self._render_slot(
                cell,
                page,
                segment.slot,
                state,
                tag_source_elements=tag_source_elements,
            )
        return max(previous_bottom, band.bottom)

    def _derive_segments(
        self,
        page: Page,
        band: _LayoutBand,
        *,
        horizontal_margin: float,
        content_width: float,
        state: _RenderState,
    ) -> tuple[_TableSegment, ...]:
        slots: list[_LayoutSlot] = []
        content_right = page.size.width - horizontal_margin
        for element in sorted(
            band.elements,
            key=lambda item: (item.bbox.x, item.z_index, item.id),
        ):
            raw_left = element.bbox.x
            raw_right = element.bbox.right
            if isinstance(element, LineElement):
                raw_left = min(element.start.x, element.end.x)
                raw_right = max(element.start.x, element.end.x)
            if (
                raw_left < horizontal_margin - GEOMETRY_TOLERANCE_PT
                or raw_right > content_right + GEOMETRY_TOLERANCE_PT
            ):
                state.warn_fallback(
                    page_id=page.id,
                    element_id=element.id,
                    code="visual_outside_layout_margins",
                    message=(
                        f"{element.type.value} was clipped to the DOCX compatibility "
                        "margins"
                    ),
                )
            left = min(content_width, max(0.0, raw_left - horizontal_margin))
            right = min(content_width, max(0.0, raw_right - horizontal_margin))
            if right - left < _MIN_LAYOUT_COLUMN_PT:
                right = min(content_width, left + _MIN_LAYOUT_COLUMN_PT)
                if right <= left:
                    left = max(0.0, content_width - _MIN_LAYOUT_COLUMN_PT)
                    right = content_width

            if slots and left < slots[-1].right - GEOMETRY_TOLERANCE_PT:
                slot = slots[-1]
                slot.elements.append(element)
                slot.left = min(slot.left, left)
                slot.right = max(slot.right, right)
            else:
                slots.append(_LayoutSlot(elements=[element], left=left, right=right))

        for slot in slots:
            if len(slot.elements) > 1:
                for element in slot.elements:
                    state.warn_fallback(
                        page_id=page.id,
                        element_id=element.id,
                        code="z_order_approximated",
                        message=(
                            "overlapping elements were layered sequentially inside one "
                            "flow-layout cell"
                        ),
                    )

        segments: list[_TableSegment] = []
        cursor = 0.0
        for slot in slots:
            if slot.left > cursor + 0.01:
                segments.append(_TableSegment(width=slot.left - cursor, slot=None))
            if slot.left < cursor:
                slot.left = cursor
            width = min(content_width - slot.left, slot.width)
            segments.append(
                _TableSegment(width=max(_MIN_LAYOUT_COLUMN_PT, width), slot=slot)
            )
            cursor = min(content_width, slot.left + max(_MIN_LAYOUT_COLUMN_PT, width))
        if cursor < content_width - 0.01:
            segments.append(_TableSegment(width=content_width - cursor, slot=None))
        if not segments:
            segments.append(_TableSegment(width=content_width, slot=None))
        return tuple(segments)

    @staticmethod
    def _configure_fixed_table(
        table: Table,
        segments: tuple[_TableSegment, ...],
        content_width: float,
    ) -> None:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table_properties = table._tbl.tblPr
        table_width = table_properties.find(qn("w:tblW"))
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            table_properties.insert(0, table_width)
        table_width.set(qn("w:type"), "dxa")
        table_width.set(qn("w:w"), str(_points_to_twips(content_width)))

        table_layout = table_properties.find(qn("w:tblLayout"))
        if table_layout is None:
            table_layout = OxmlElement("w:tblLayout")
            table_properties.append(table_layout)
        table_layout.set(qn("w:type"), "fixed")

        table_borders = table_properties.find(qn("w:tblBorders"))
        if table_borders is None:
            table_borders = OxmlElement("w:tblBorders")
            table_properties.append(table_borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            edge_element = table_borders.find(qn(f"w:{edge}"))
            if edge_element is None:
                edge_element = OxmlElement(f"w:{edge}")
                table_borders.append(edge_element)
            edge_element.set(qn("w:val"), "nil")

        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for segment in segments:
            grid_column = OxmlElement("w:gridCol")
            grid_column.set(qn("w:w"), str(_points_to_twips(segment.width)))
            grid.append(grid_column)

    @staticmethod
    def _configure_cell(cell: _Cell, width: float) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_properties = cell._tc.get_or_add_tcPr()
        cell_width = cell_properties.find(qn("w:tcW"))
        if cell_width is None:
            cell_width = OxmlElement("w:tcW")
            cell_properties.append(cell_width)
        cell_width.set(qn("w:type"), "dxa")
        cell_width.set(qn("w:w"), str(_points_to_twips(width)))

        margins = cell_properties.find(qn("w:tcMar"))
        if margins is None:
            margins = OxmlElement("w:tcMar")
            cell_properties.append(margins)
        for edge in ("top", "left", "bottom", "right"):
            margin = margins.find(qn(f"w:{edge}"))
            if margin is None:
                margin = OxmlElement(f"w:{edge}")
                margins.append(margin)
            margin.set(qn("w:w"), "0")
            margin.set(qn("w:type"), "dxa")

        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

    def _render_slot(
        self,
        cell: _Cell,
        page: Page,
        slot: _LayoutSlot,
        state: _RenderState,
        *,
        tag_source_elements: bool = False,
    ) -> None:
        ordered_elements = sorted(
            slot.elements,
            key=lambda element: (element.z_index, element.id),
        )
        for element in ordered_elements:
            if isinstance(element, RectangleElement):
                self._render_rectangle(cell, page, element, state)
            elif isinstance(element, ImageElement):
                self._render_image(cell, page, element, state)
            elif isinstance(element, LineElement):
                self._render_line_in_cell(cell, page, element, state)
            elif isinstance(element, TextElement):
                paragraph = self._available_cell_paragraph(cell)
                self._apply_text_style(
                    paragraph,
                    page,
                    element,
                    state=state,
                    source_element_id=(element.id if tag_source_elements else None),
                )
            state.record_rendered(element.id)

    def _render_rectangle(
        self,
        cell: _Cell,
        page: Page,
        element: RectangleElement,
        state: _RenderState,
    ) -> None:
        style = element.style
        if style.corner_radius_pt > GEOMETRY_TOLERANCE_PT:
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="rounded_rectangle_squared",
                message="rounded rectangle corners were approximated as square corners",
            )
        if not math.isclose(style.opacity, 1.0, abs_tol=1e-9):
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="opacity_approximated",
                message="rectangle opacity was rendered as fully opaque",
            )
        if style.stroke_width_pt <= 0 or style.stroke_color is None:
            border = {"val": "nil"}
        else:
            border = self._visual_border(
                page,
                element,
                width_pt=style.stroke_width_pt,
                color=style.stroke_color,
                dash=LineDash.SOLID,
                state=state,
            )
        self._set_cell_borders(
            cell,
            top=border,
            left=border,
            bottom=border,
            right=border,
        )
        if style.fill_color is not None:
            cell_properties = cell._tc.get_or_add_tcPr()
            shading = cell_properties.find(qn("w:shd"))
            if shading is None:
                shading = OxmlElement("w:shd")
                cell_properties.append(shading)
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), style.fill_color[1:].upper())
        if style.stroke_width_pt <= 0 and style.fill_color is None:
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="invisible_rectangle_placeholder",
                message="invisible rectangle was retained as an empty layout cell",
            )

    def _render_image(
        self,
        cell: _Cell,
        page: Page,
        element: ImageElement,
        state: _RenderState,
    ) -> None:
        resolved = state.resolved_assets.get(element.asset_id)
        if resolved is None or element.id in state.unavailable_images:
            self._render_image_placeholder(cell, element)
            return
        paragraph = self._available_cell_paragraph(cell)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        image_data, width_pt, height_pt = self._fit_image(resolved, element)
        try:
            run = paragraph.add_run()
            inline_shape = run.add_picture(
                BytesIO(image_data),
                width=Pt(width_pt),
                height=Pt(height_pt),
            )
            alt_text = element.alt_text or f"Image asset {element.asset_id}"
            inline_shape._inline.docPr.set("descr", alt_text)
            inline_shape._inline.docPr.set("title", alt_text)
        except (UnrecognizedImageError, OSError, ValueError) as exc:
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="image_embedding_failed",
                message=f"verified image could not be embedded: {exc}",
            )
            self._render_image_placeholder(cell, element)

    @staticmethod
    def _fit_image(
        resolved: ResolvedAsset,
        element: ImageElement,
    ) -> tuple[bytes, float, float]:
        with Image.open(BytesIO(resolved.data)) as image:
            pixel_width, pixel_height = image.size
            if element.fit is ImageFit.STRETCH:
                return resolved.data, element.bbox.width, element.bbox.height
            if element.fit is ImageFit.CONTAIN:
                scale = min(
                    element.bbox.width / pixel_width,
                    element.bbox.height / pixel_height,
                )
                return resolved.data, pixel_width * scale, pixel_height * scale

            target_ratio = element.bbox.width / element.bbox.height
            source_ratio = pixel_width / pixel_height
            if source_ratio > target_ratio:
                crop_width = max(1, round(pixel_height * target_ratio))
                left = (pixel_width - crop_width) // 2
                crop_box = (left, 0, left + crop_width, pixel_height)
            else:
                crop_height = max(1, round(pixel_width / target_ratio))
                top = (pixel_height - crop_height) // 2
                crop_box = (0, top, pixel_width, top + crop_height)
            cropped = image.crop(crop_box)
            buffer = BytesIO()
            cropped.save(buffer, format="PNG", optimize=False, compress_level=9)
            return buffer.getvalue(), element.bbox.width, element.bbox.height

    @staticmethod
    def _render_image_placeholder(cell: _Cell, element: ImageElement) -> None:
        paragraph = PythonDocxRenderer._available_cell_paragraph(cell)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(
            f"[Image unavailable: {element.alt_text or element.asset_id}]"
        )
        run.font.name = DEFAULT_FALLBACK_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    def _render_line_in_cell(
        self,
        cell: _Cell,
        page: Page,
        element: LineElement,
        state: _RenderState,
    ) -> None:
        border = self._line_border(page, element, state)
        if self._is_vertical_line(element):
            self._set_cell_borders(cell, left=border)
            return
        paragraph = self._available_cell_paragraph(cell)
        self._set_paragraph_bottom_border(paragraph, border)
        if not self._is_horizontal_line(element):
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="diagonal_line_approximated",
                message="diagonal line was approximated as a horizontal border",
            )

    def _line_border(
        self,
        page: Page,
        element: LineElement,
        state: _RenderState,
    ) -> dict[str, str]:
        if not math.isclose(element.style.opacity, 1.0, abs_tol=1e-9):
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="opacity_approximated",
                message="line opacity was rendered as fully opaque",
            )
        return self._visual_border(
            page,
            element,
            width_pt=element.style.width_pt,
            color=element.style.color,
            dash=element.style.dash,
            state=state,
        )

    @staticmethod
    def _visual_border(
        page: Page,
        element: DocumentElement,
        *,
        width_pt: float,
        color: str | None,
        dash: LineDash,
        state: _RenderState,
    ) -> dict[str, str]:
        clamped_width = min(
            _MAX_WORD_BORDER_PT,
            max(_MIN_WORD_BORDER_PT, width_pt),
        )
        if not math.isclose(clamped_width, width_pt, abs_tol=1e-9):
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="border_width_clamped",
                message=(
                    f"border width {width_pt:g}pt was clamped to "
                    f"{clamped_width:g}pt for Word compatibility"
                ),
            )
        return {
            "val": _LINE_DASH_MAP[dash],
            "sz": str(max(2, min(96, round(clamped_width * 8)))),
            "space": "0",
            "color": color[1:].upper() if color is not None else "auto",
        }

    @staticmethod
    def _is_horizontal_line(element: LineElement) -> bool:
        return math.isclose(
            element.start.y,
            element.end.y,
            abs_tol=GEOMETRY_TOLERANCE_PT,
        )

    @staticmethod
    def _is_vertical_line(element: LineElement) -> bool:
        return math.isclose(
            element.start.x,
            element.end.x,
            abs_tol=GEOMETRY_TOLERANCE_PT,
        )

    @staticmethod
    def _available_cell_paragraph(cell: _Cell) -> Paragraph:
        paragraph = cell.paragraphs[-1]
        if paragraph.text or paragraph._p.xpath(".//w:drawing"):
            paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        return paragraph

    def _apply_text_style(
        self,
        paragraph: Paragraph,
        page: Page,
        element: TextElement,
        *,
        state: _RenderState,
        source_element_id: str | None = None,
        font_size_pt: float | None = None,
    ) -> None:
        style = (
            element.style
            if font_size_pt is None
            else replace(element.style, font_size_pt=font_size_pt)
        )
        paragraph.alignment = _ALIGNMENT_MAP[style.align]
        paragraph.paragraph_format.line_spacing = style.line_height
        run = paragraph.add_run(element.text)

        resolved_font = self._supported_fonts.get(style.font_family.casefold())
        if resolved_font is None:
            resolved_font = self._fallback_font
            state.substitutions.append(
                FontSubstitution(
                    element_id=element.id,
                    requested=style.font_family,
                    replacement=resolved_font,
                )
            )
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="font_substituted",
                message=(
                    f"font {style.font_family!r} was replaced with {resolved_font!r}"
                ),
            )

        self._set_run_font(run, resolved_font, style)
        if source_element_id is not None:
            self._wrap_run_with_source_tag(run, source_element_id)
        if style.font_weight not in {400, 700}:
            mapped_weight = 700 if style.font_weight >= 600 else 400
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="font_weight_approximated",
                message=f"font weight {style.font_weight} was mapped to {mapped_weight}",
            )
        if style.font_style is FontStyle.OBLIQUE:
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="oblique_mapped_to_italic",
                message="oblique text was mapped to DOCX italic",
            )
        if not math.isclose(style.rotation_deg % 360, 0.0, abs_tol=1e-9):
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="rotation_omitted",
                message="text rotation is unsupported and was rendered unrotated",
            )
        if not math.isclose(style.opacity, 1.0, abs_tol=1e-9):
            state.warn_fallback(
                page_id=page.id,
                element_id=element.id,
                code="opacity_approximated",
                message="text opacity is unsupported and was rendered opaque",
            )

    @staticmethod
    def _wrap_run_with_source_tag(run: Run, source_element_id: str) -> None:
        run_element = run._r
        paragraph = run_element.getparent()
        if paragraph is None:
            raise DocxRenderError("text run has no OOXML paragraph parent")
        index = paragraph.index(run_element)
        paragraph.remove(run_element)

        control = OxmlElement("w:sdt")
        properties = OxmlElement("w:sdtPr")
        tag = OxmlElement("w:tag")
        tag.set(qn("w:val"), f"{_SOURCE_TAG_PREFIX}{source_element_id}")
        properties.append(tag)
        alias = OxmlElement("w:alias")
        alias.set(qn("w:val"), source_element_id)
        properties.append(alias)
        content = OxmlElement("w:sdtContent")
        content.append(run_element)
        control.append(properties)
        control.append(content)
        paragraph.insert(index, control)

    @staticmethod
    def _set_run_font(run: Run, font_name: str, style: TextStyle) -> None:
        run.font.name = font_name
        run.font.size = Pt(style.font_size_pt)
        run.bold = style.font_weight >= 600
        run.italic = style.font_style in {FontStyle.ITALIC, FontStyle.OBLIQUE}
        if style.color is not None:
            run.font.color.rgb = RGBColor.from_string(style.color[1:].upper())

        run_properties = run._element.get_or_add_rPr()
        run_fonts = run_properties.get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            run_fonts.set(qn(f"w:{attribute}"), font_name)

    @staticmethod
    def _set_paragraph_bottom_border(
        paragraph: Paragraph,
        border: dict[str, str],
    ) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        borders = paragraph_properties.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            paragraph_properties.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        for key, value in border.items():
            bottom.set(qn(f"w:{key}"), value)

    @staticmethod
    def _set_cell_borders(
        cell: _Cell,
        **edges: dict[str, str],
    ) -> None:
        cell_properties = cell._tc.get_or_add_tcPr()
        borders = cell_properties.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            cell_properties.append(borders)
        for edge, settings in edges.items():
            border = borders.find(qn(f"w:{edge}"))
            if border is None:
                border = OxmlElement(f"w:{edge}")
                borders.append(border)
            for key, value in settings.items():
                border.set(qn(f"w:{key}"), value)

    @staticmethod
    def _save_atomically(word_document: WordDocument, target: Path) -> None:
        target.parent.mkdir(parents=True, exist_ok=True)
        file_descriptor, temporary_name = tempfile.mkstemp(
            dir=target.parent,
            prefix=f".{target.stem}-",
            suffix=".tmp.docx",
        )
        os.close(file_descriptor)
        temporary_path = Path(temporary_name)
        try:
            word_document.save(temporary_path)
            open_docx(temporary_path)
            os.replace(temporary_path, target)
        except Exception as exc:
            raise DocxRenderError(f"failed to create DOCX at {target}: {exc}") from exc
        finally:
            temporary_path.unlink(missing_ok=True)


def _points_to_twips(points: float) -> int:
    return max(1, round(points * 20))


def _distributed_twips(
    widths: tuple[float, ...],
    total_width: float,
) -> tuple[int, ...]:
    """Round grid widths while preserving the exact rounded table width."""

    if not widths:
        raise ValueError("native table requires at least one column")
    rounded = [max(1, round(width * 20)) for width in widths]
    target = _points_to_twips(total_width)
    rounded[-1] += target - sum(rounded)
    if rounded[-1] < 1:
        raise DocxRenderError("native table column widths collapse after rounding")
    return tuple(rounded)
